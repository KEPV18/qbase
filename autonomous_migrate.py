#!/usr/bin/env python3
"""
Autonomous Migration — UPDATE mode for all existing 249 records.
Re-extracts from DOCX files and PATCHES existing QBase records.
Creates missing ones if any.
Logs to /tmp/autonomous_migration.log
"""
import json, os, sys, re, urllib.request, urllib.error
from docx import Document
from pathlib import Path

RECORDS_DIR = Path('/home/Kepv/Downloads/03.Records')
LOG_FILE = '/tmp/autonomous_migration.log'

def load_env():
    env = {}
    with open('/mnt/ahmed/Projects/qbase/.env.local') as f:
        for line in f:
            line = line.strip()
            if line and not line.startswith('#') and '=' in line:
                k,v=line.split('=',1)
                env[k]=v
    return env

ENV = load_env()
SUPABASE_URL = ENV['SUPABASE_URL']
SERVICE_KEY = ENV['SUPABASE_SERVICE_ROLE_KEY']

FORM_CONFIG = {
    'F/08': ('Order Form', 1, 'Sales \u0026 Customer Service', 'On event'),
    'F/09': ('Customer Complaint', 1, 'Sales \u0026 Customer Service', 'On event'),
    'F/10': ('Customer Feedback', 1, 'Sales \u0026 Customer Service', 'On event'),
    'F/11': ('Production Plan', 2, 'Operations \u0026 Production', 'Monthly'),
    'F/12': ('Non-Conforming', 3, 'Quality \u0026 Audit', 'On event'),
    'F/13': ('Purchase Order', 4, 'Procurement \u0026 Vendors', 'On event'),
    'F/14': ('Incoming Inspection', 4, 'Procurement \u0026 Vendors', 'On event'),
    'F/15': ('Approved Vendor List', 4, 'Procurement \u0026 Vendors', 'Annual'),
    'F/16': ('Supplier Registration Form', 4, 'Procurement \u0026 Vendors', 'On event'),
    'F/17': ('QA Test Request', 3, 'Quality \u0026 Audit', 'On event'),
    'F/18': ('Product Re-Call', 3, 'Quality \u0026 Audit', 'On event'),
    'F/19': ('Product Description', 2, 'Operations \u0026 Production', 'Per-project'),
    'F/20': ('Review Agenda', 7, 'Management \u0026 Documentation', 'On event'),
    'F/21': ('Review Minutes', 7, 'Management \u0026 Documentation', 'On event'),
    'F/22': ('Corrective Action', 3, 'Quality \u0026 Audit', 'On event'),
    'F/23': ('Master List of Records', 7, 'Management \u0026 Documentation', 'On event'),
    'F/24': ('Objectives \u0026 Targets', 7, 'Management \u0026 Documentation', 'Annual'),
    'F/25': ('Audit Plan', 3, 'Quality \u0026 Audit', 'Annual'),
    'F/28': ('Training Attendance', 5, 'HR \u0026 Training', 'On event'),
    'F/29': ('Training Record', 5, 'HR \u0026 Training', 'On event'),
    'F/30': ('Performance Appraisal', 5, 'HR \u0026 Training', 'Annual'),
    'F/32': ('R\u0026D Request', 6, 'R\u0026D \u0026 Design', 'On event'),
    'F/34': ('Design Verification', 6, 'R\u0026D \u0026 Design', 'On event'),
    'F/35': ('Design Monitoring', 6, 'R\u0026D \u0026 Design', 'On event'),
    'F/37': ('Experiment Data', 6, 'R\u0026D \u0026 Design', 'On event'),
    'F/40': ('Competence Matrix', 5, 'HR \u0026 Training', 'Annual'),
    'F/41': ('Gap Analysis', 5, 'HR \u0026 Training', 'Annual'),
    'F/42': ('Annual Training Plan', 5, 'HR \u0026 Training', 'Annual'),
    'F/43': ('Induction Training Record', 5, 'HR \u0026 Training', 'On event'),
    'F/44': ('Job Description', 5, 'HR \u0026 Training', 'On event'),
    'F/45': ('Master List of Docs', 7, 'Management \u0026 Documentation', 'On event'),
    'F/46': ('Change Management', 7, 'Management \u0026 Documentation', 'On event'),
    'F/47': ('Audit Checklist', 3, 'Quality \u0026 Audit', 'Annual'),
    'F/48': ('Internal Audit Report', 3, 'Quality \u0026 Audit', 'On event'),
    'F/50': ('Customer Property', 1, 'Sales \u0026 Customer Service', 'On event'),
}

FORM_FOLDER_MAP = {
    'F-08': 'F/08', 'F-09': 'F/09', 'F-10': 'F/10', 'F-11': 'F/11', 'F-12': 'F/12',
    'F-13': 'F/13', 'F-14': 'F/14', 'F-15': 'F/15', 'F-16': 'F/16', 'F-17': 'F/17',
    'F-18': 'F/18', 'F-19': 'F/19', 'F-20': 'F/20', 'F-21': 'F/21', 'F-22': 'F/22',
    'F-23': 'F/23', 'F-24': 'F/24', 'F-25': 'F/25', 'F-28': 'F/28', 'F-29': 'F/29',
    'F-30': 'F/30', 'F-32': 'F/32', 'F-34': 'F/34', 'F-35': 'F/35', 'F-37': 'F/37',
    'F-40': 'F/40', 'F-41': 'F/41', 'F-42': 'F/42', 'F-43': 'F/43', 'F-44': 'F/44',
    'F-45': 'F/45', 'F-46': 'F/46', 'F-47': 'F/47', 'F-48': 'F/48', 'F-50': 'F/50',
}

def clean_cells(row):
    cells = [cell.text.strip().replace('\n',' ') for cell in row.cells]
    clean = []
    for c in cells:
        if clean and clean[-1] == c: continue
        clean.append(c)
    return clean

def extract_generic(filepath):
    doc = Document(filepath)
    data = {'_type': 'generic'}
    for table in doc.tables:
        for row in table.rows:
            cells = clean_cells(row)
            for i in range(len(cells)-1):
                label = cells[i].lower().strip()
                value = cells[i+1].strip()
                if value and label and len(label) < 80 and len(value) < 500:
                    if ':' in cells[i]:
                        parts = cells[i].split(':', 1)
                        l = parts[0].strip().lower()
                        v = parts[1].strip() if len(parts) > 1 else ''
                        if v: data[l] = v
                    elif len(label) > 3 and len(value) > 0:
                        data[label] = value
    return data

def extract_f08(filepath):
    doc = Document(filepath)
    data = {'_type': 'f08'}
    table = doc.tables[0] if doc.tables else None
    if not table: return data
    for row in table.rows:
        cells = clean_cells(row)
        row_text = '|'.join(cells)
        m = re.search(r'Date\s*[🡪:\t\s]+(\d{2}/\d{2}/\d{4})', row_text)
        if m: data['date'] = m.group(1)
        if 'Customer' in row_text and not data.get('client_name'):
            for c in cells:
                if 'Customer' not in c and len(c) > 3:
                    data['client_name'] = c; break
        if 'Receipt' in row_text and not data.get('mode_of_receipt'):
            for c in cells:
                if 'Mode' not in c and 'Receipt' not in c and len(c) > 5:
                    data['mode_of_receipt'] = c; break
        if 'Product Name' in row_text and 'Specifications' in row_text:
            data['_items_header_seen'] = '1'
            continue
        if data.get('_items_header_seen'):
            if any(k in row_text for k in ['Video Detection', 'Vocal AI', 'Tennis', 'Omniaz', 'ETH AI', 'Widget']):
                product = [c for c in cells if 'Service' in c or 'Analytics' in c or 'Annotation' in c or 'Testing' in c or 'Widget' in c]
                specs = [c for c in cells if 'labeling' in c or 'Conversational' in c or 'Match tagging' in c or 'Data annotation' in c or 'validation' in c or 'optimization' in c or 'evaluation' in c or 'Spec' in c]
                qty = [c for c in cells if 'client volume' in c.lower() or c.strip().isdigit()]
                if product:
                    item = {'product_name': product[0]}
                    if specs: item['specifications'] = specs[0]
                    if qty: item['qty'] = qty[0]
                    elif 'As per client volume' in row_text: item['qty'] = 'As per client volume'
                    if '_items' not in data: data['_items'] = []
                    data['_items'].append(item)
        if 'Requirement' in row_text and 'Test Certificate' in row_text:
            for c in cells:
                if c.strip() in ['No', 'Yes', 'N/A']:
                    data['test_certificate_required'] = 'Yes' if c == 'Yes' else 'No'
        if 'Delivery' in row_text and 'Schedule' in row_text:
            for c in cells:
                if 'Delivery' not in c and 'Schedule' not in c and len(c) > 10:
                    data['delivery_schedule'] = c; break
        if 'Statutory' in row_text or ('Regulatory' in row_text and 'Requirements' in row_text):
            for c in cells:
                if len(c) > 10 and 'Statutory' not in c and 'Regulatory' not in c and 'Requirements' not in c:
                    data['complies'] = c; break
        if 'Accepted' in row_text:
            data['order_status'] = 'Accepted'
        if 'Order' in row_text and 'Rejected' in row_text:
            data['order_status'] = 'Rejected'
        if 'Bill No' in row_text:
            for c in cells:
                if 'Bill' not in c and 'No' not in c and c.strip():
                    data['bill_no'] = c; break
        if 'Despatch' in row_text:
            for c in cells:
                if 'Despatch' not in c and c.strip():
                    data['despatch_date'] = c; break
        if 'Remarks' in row_text:
            for c in cells:
                if 'Remarks' not in c and 'Reviewed' not in c and len(c) > 20:
                    data['remarks'] = c; break
        if 'Reviewed' in row_text:
            for c in cells:
                if 'Reviewed' not in c and 'By' not in c and len(c) > 5:
                    data['reviewed_by'] = c; break
    if '_items' in data:
        data['items'] = json.dumps(data.pop('_items'), ensure_ascii=False)
    return data

EXTRACTORS = {
    'F/08': extract_f08,
    # TODO: add more custom extractors for F/10, F/19, F/28, F/29, F/30, F/35, F/44, F/47, F/48, F/50
}

def log(msg):
    with open(LOG_FILE, 'a') as f:
        f.write(msg + '\n')
    print(msg)

def get_serial_from_filename(filename):
    m = re.search(r'[Ff][_/\-](\d{2})[\-_](\d+)', filename)
    if m:
        return f"F/{m.group(1)}-{int(m.group(2)):03d}"
    return None

def sb_fetch(method, path, body=None):
    url = f"{SUPABASE_URL}/rest/v1/{path}"
    req = urllib.request.Request(url, data=body.encode() if body else None, method=method)
    req.add_header('apikey', SERVICE_KEY)
    req.add_header('Authorization', f'Bearer {SERVICE_KEY}')
    req.add_header('Content-Type', 'application/json')
    if method == 'POST':
        req.add_header('Prefer', 'return=minimal')
    try:
        resp = urllib.request.urlopen(req, timeout=30)
        return {'status': resp.status, 'data': resp.read().decode()}
    except urllib.error.HTTPError as e:
        return {'status': e.code, 'error': e.read().decode()[:200]}
    except Exception as e:
        return {'status': 0, 'error': str(e)}

def main():
    open(LOG_FILE, 'w').close()
    log(f"🚀 AUTONOMOUS MIGRATION (UPDATE MODE) STARTED")
    log(f"Supabase: {SUPABASE_URL}")

    # Fetch existing records
    existing = {}
    resp = sb_fetch('GET', "records?select=serial,form_code,id&limit=10000")
    if resp['status'] == 200:
        rows = json.loads(resp['data'])
        for r in rows:
            if r.get('serial'):
                existing[r['serial']] = {'form_code': r['form_code'], 'id': r['id']}
        log(f"📊 Existing QBase records: {len(existing)}")
    else:
        log(f"⚠️ Could not fetch existing: {resp}")
        return

    # Scan ALL DOCX files (records + templates + archives)
    all_files = []
    for section_dir in sorted(RECORDS_DIR.iterdir()):
        if not section_dir.is_dir(): continue
        for form_dir in sorted(section_dir.iterdir()):
            if not form_dir.is_dir(): continue
            fc = None
            for prefix, code in FORM_FOLDER_MAP.items():
                if prefix in form_dir.name:
                    fc = code; break
            if not fc: continue
            for f in sorted(form_dir.rglob('*.docx')):
                if 'Archive' in str(f): continue
                serial = get_serial_from_filename(f.name)
                if serial:
                    all_files.append((f, fc, serial))

    # Separate: existing vs missing
    to_update = [(f, fc, s) for f, fc, s in all_files if s in existing]
    to_create = [(f, fc, s) for f, fc, s in all_files if s not in existing]

    log(f"📄 DOCX files mapped: {len(all_files)}")
    log(f"   → Need UPDATE: {len(to_update)} (existing)")
    log(f"   → Need CREATE: {len(to_create)} (missing)")

    updated = 0
    created = 0
    errors = []

    # === UPDATE EXISTING ===
    for idx, (fpath, form_code, serial) in enumerate(to_update):
        if idx % 20 == 0:
            log(f"  [UPD {idx:3d}/{len(to_update)}] {form_code} {serial}...")

        extractor = EXTRACTORS.get(form_code, extract_generic)
        try:
            extracted = extractor(str(fpath))
        except Exception as e:
            errors.append(f"{serial}: EXTRACT ERROR {e}")
            continue

        form_data = {}
        for k, v in extracted.items():
            if k.startswith('_'): continue
            if isinstance(v, (list, dict)):
                form_data[k] = json.dumps(v, ensure_ascii=False)
            elif isinstance(v, str):
                form_data[k] = v
            else:
                form_data[k] = str(v) if v else ''

        # PATCH existing
        record_id = existing[serial]['id']
        patch_body = json.dumps({'form_data': form_data, 'last_modified_by': 'migration_v5'})
        resp = sb_fetch('PATCH', f'records?id=eq.{record_id}', patch_body)
        if resp['status'] in (200, 201, 204):
            updated += 1
        else:
            err = resp.get('error', 'unknown')
            errors.append(f"{serial}: HTTP {resp['status']} {err}")

    # === CREATE MISSING ===
    for idx, (fpath, form_code, serial) in enumerate(to_create):
        if idx % 5 == 0:
            log(f"  [NEW {idx:3d}/{len(to_create)}] {form_code} {serial}...")

        extractor = EXTRACTORS.get(form_code, extract_generic)
        try:
            extracted = extractor(str(fpath))
        except Exception as e:
            errors.append(f"{serial}: EXTRACT ERROR {e}")
            continue

        form_data = {}
        for k, v in extracted.items():
            if k.startswith('_'): continue
            if isinstance(v, (list, dict)):
                form_data[k] = json.dumps(v, ensure_ascii=False)
            elif isinstance(v, str):
                form_data[k] = v
            else:
                form_data[k] = str(v) if v else ''

        cfg = FORM_CONFIG.get(form_code, (form_code, 0, '', 'Unknown'))
        row = {
            'form_code': form_code,
            'serial': serial,
            'form_name': cfg[0],
            'status': 'approved',
            'form_data': form_data,
            'section': cfg[1],
            'section_name': cfg[2],
            'frequency': cfg[3] if len(cfg) > 3 else 'Unknown',
            'created_by': 'migration_v5',
            'last_modified_by': 'migration_v5',
            'edit_count': 0,
            'modification_reason': 'Autonomous migration v5 from DOCX files',
            'deleted_at': None,
        }
        resp = sb_fetch('POST', 'records', json.dumps(row))
        if resp['status'] in (200, 201, 204):
            created += 1
        else:
            err = resp.get('error', 'unknown')
            errors.append(f"{serial}: HTTP {resp['status']} {err}")

    log(f"\n{'='*60}")
    log(f"📊 AUTONOMOUS MIGRATION COMPLETE")
    log(f"   Total DOCX files: {len(all_files)}")
    log(f"   Updated: {updated}")
    log(f"   Created: {created}")
    log(f"   Errors: {len(errors)}")
    if errors:
        for e in errors[:30]:
            log(f"     ⚠️ {e}")
    log(f"   Log saved: {LOG_FILE}")

if __name__ == '__main__':
    main()
