#!/usr/bin/env python3
"""
QMS → QBase Migration — Phase 3: Extract & Insert all records
Uses cached Drive listing. Downloads files not yet in cache.
Usage:
  python3 migrate_phase3.py analyze   # Analyze only, don't insert
  python3 migrate_phase3.py migrate  # Full migration to Supabase
"""
import json, os, sys, re, urllib.request, urllib.parse, time
from datetime import datetime
from docx import Document
from pathlib import Path

# ============================================================================
# CONFIG
# ============================================================================
DRIVE_ROOT = '1zA3ZqbFwxsa75DBl55oXybB2jyxvagYS'
TOKEN_PATH=Path.home()/'.hermes'/'google_token.json'
CACHE_DIR=Path('/tmp/qms_drive_cache')
SAMPLES_DIR=Path('/tmp/qms_samples')
MIGRATION_LOG=Path('/tmp/qms_migration_log.json')

def load_env():
    env={}
    with open('/mnt/ahmed/Projects/qbase/.env.local') as f:
        for line in f:
            line=line.strip()
            if line and not line.startswith('#') and '=' in line:
                k,v=line.split('=',1)
                env[k]=v
    return env

ENV=load_env()
SUPABASE_URL=ENV['SUPABASE_URL']
SERVICE_KEY=ENV['SUPABASE_SERVICE_ROLE_KEY']

FORM_CONFIG={
    'F/08 - Order Form':{'code':'F/08','section':1,'name':'Order Form'},
    'F/09 - Customer Complaint':{'code':'F/09','section':1,'name':'Customer Complaint'},
    'F/10 - Customer Feedback':{'code':'F/10','section':1,'name':'Customer Feedback'},
    'F/50 - Customer Property':{'code':'F/50','section':1,'name':'Customer Property'},
    'F/11 - Production Plan':{'code':'F/11','section':2,'name':'Production Plan'},
    'F/19 - Product Description':{'code':'F/19','section':2,'name':'Product Description'},
    'F/12 - Non-Conforming':{'code':'F/12','section':3,'name':'Non-Conforming'},
    'F/17 - QA Test Request':{'code':'F/17','section':3,'name':'QA Test Request'},
    'F/18 - Product Re-Call':{'code':'F/18','section':3,'name':'Product Re-Call'},
    'F/22 - Corrective Action':{'code':'F/22','section':3,'name':'Corrective Action'},
    'F/25 - Audit Plan':{'code':'F/25','section':3,'name':'Audit Plan'},
    'F/47 - Audit Checklist':{'code':'F/47','section':3,'name':'Audit Checklist'},
    'F/48 - Internal Audit Report':{'code':'F/48','section':3,'name':'Internal Audit Report'},
    'F-16 Supplier Registration Form':{'code':'F/16','section':4,'name':'Supplier Registration Form'},
    'F/13 - Purchase Order':{'code':'F/13','section':4,'name':'Purchase Order'},
    'F/14 - Incoming Inspection':{'code':'F/14','section':4,'name':'Incoming Inspection'},
    'F/15 - Approved Vendor List':{'code':'F/15','section':4,'name':'Approved Vendor List'},
    'F/28 - Training Attendance':{'code':'F/28','section':5,'name':'Training Attendance'},
    'F/29 - Training Record':{'code':'F/29','section':5,'name':'Training Record'},
    'F/30 - Performance Appraisal':{'code':'F/30','section':5,'name':'Performance Appraisal'},
    'F/40 - Competence Matrix':{'code':'F/40','section':5,'name':'Competence Matrix'},
    'F/41 - Gap Analysis':{'code':'F/41','section':5,'name':'Gap Analysis'},
    'F/42 - Annual Training Plan':{'code':'F/42','section':5,'name':'Annual Training Plan'},
    'F/43 - Induction Training Record':{'code':'F/43','section':5,'name':'Induction Training Record'},
    'F/44 - Job Description':{'code':'F/44','section':5,'name':'Job Description'},
    'F/32 - R&D Request':{'code':'F/32','section':6,'name':'R&D Request'},
    'F/34 - Design Verification':{'code':'F/34','section':6,'name':'Design Verification'},
    'F/35 - Design Monitoring':{'code':'F/35','section':6,'name':'Design Monitoring'},
    'F/37 - Experiment Data':{'code':'F/37','section':6,'name':'Experiment Data'},
    'F/20 - Review Agenda':{'code':'F/20','section':7,'name':'Review Agenda'},
    'F/21 - Review Minutes':{'code':'F/21','section':7,'name':'Review Minutes'},
    'F/23 - Master List of Records':{'code':'F/23','section':7,'name':'Master List of Records'},
    'F/24 - Objectives & Targets':{'code':'F/24','section':7,'name':'Objectives & Targets'},
    'F/45 - Master List of Docs':{'code':'F/45','section':7,'name':'Master List of Docs'},
    'F/46 - Change Management':{'code':'F/46','section':7,'name':'Change Management'},
}

SECTION_NAMES={1:'Sales & Customer Service',2:'Operations & Production',3:'Quality & Audit',4:'Procurement & Vendors',5:'HR & Training',6:'R&D & Design',7:'Management & Documentation'}

# ============================================================================
# GOOGLE DRIVE API
# ============================================================================
def get_access_token():
    with open(TOKEN_PATH) as f: data=json.load(f)
    post_data=urllib.parse.urlencode({'client_id':data['client_id'],'client_secret':data['client_secret'],'refresh_token':data['refresh_token'],'grant_type':'refresh_token'}).encode()
    req=urllib.request.Request('https://oauth2.googleapis.com/token',data=post_data)
    result=json.loads(urllib.request.urlopen(req).read())
    data['access_token']=result['access_token']
    with open(TOKEN_PATH,'w') as f: json.dump(data,f,indent=2)
    return result['access_token']

def drive_api(path, token):
    url=f"https://www.googleapis.com{path}"
    req=urllib.request.Request(url)
    req.add_header('Authorization',f'Bearer {token}')
    return json.loads(urllib.request.urlopen(req, timeout=30).read())

def drive_list_all(folder_id, token):
    all_files=[]; page_token=None
    while True:
        q=f"'{folder_id}' in parents and trashed=false"
        url=f"https://www.googleapis.com/drive/v3/files?q={urllib.parse.quote(q)}&fields=nextPageToken,files(id,name,mimeType)&pageSize=500"
        if page_token: url+=f"&pageToken={page_token}"
        req=urllib.request.Request(url)
        req.add_header('Authorization',f'Bearer {token}')
        data=json.loads(urllib.request.urlopen(req, timeout=30).read())
        all_files.extend(data.get('files',[]))
        page_token=data.get('nextPageToken')
        if not page_token: break
    return all_files

def download_docx(file_id, token, outpath):
    meta=drive_api(f"/drive/v3/files/{file_id}?fields=mimeType,name", token)
    mime=meta.get('mimeType','')
    os.makedirs(os.path.dirname(outpath), exist_ok=True)
    if mime=='application/vnd.google-apps.document':
        url=f"https://www.googleapis.com/drive/v3/files/{file_id}/export?mimeType=application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif mime.startswith('application/vnd.google-apps.spreadsheet'):
        return None
    else:
        url=f"https://www.googleapis.com/drive/v3/files/{file_id}?alt=media"
    req=urllib.request.Request(url)
    req.add_header('Authorization',f'Bearer {token}')
    with open(outpath,'wb') as f: f.write(urllib.request.urlopen(req, timeout=30).read())
    return outpath

# ============================================================================
# DOCX PARSER
# ============================================================================
def parse_qms_docx(filepath):
    doc=Document(filepath)
    all_paras=[p.text.strip() for p in doc.paragraphs if p.text.strip()]
    if not doc.tables:
        return _extract_paragraphs(all_paras)
    all_rows=[]
    for table in doc.tables:
        for row in table.rows:
            cells=[cell.text.strip().replace('\n',' | ').replace('\r','') for cell in row.cells]
            all_rows.append(cells)
    return _extract_rows(all_rows, all_paras)

def _extract_paragraphs(paragraphs):
    data={'_type':'paragraphs'}
    full_text='\n'.join(paragraphs)
    m=re.search(r'(F[/\-]\d+[-]\d+)',full_text)
    if m:
        s=m.group(1).replace('\\','/')
        m2=re.match(r'F(\d{2})[-/](\d+)',s)
        data['serial']=f"F/{m2.group(1)}-{m2.group(2)}" if m2 else s
    m=re.search(r'Date[:\s]*(\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4})',full_text)
    if m: data['date']=m.group(1)
    for p in paragraphs:
        for sep in ['🡪','\t']:
            if sep in p:
                label,_,value=p.partition(sep)
                label=label.strip().lower(); value=value.strip()
                if value: _map_label(label,value,data)
    data['_raw_text']=full_text[:3000]
    return data

def _extract_rows(rows, paragraphs=None):
    data={'_type':'table'}
    all_cells=[]
    for row in rows: all_cells.extend([c for c in row if c])
    full_text=' | '.join(all_cells)
    
    # Serial
    m=re.search(r'(F[/\-]\d+[-]\d+)',full_text)
    if m:
        s=m.group(1).replace('\\','/')
        m2=re.match(r'F(\d{2})[-/](\d+)',s)
        data['serial']=f"F/{m2.group(1)}-{m2.group(2)}" if m2 else s
    
    # Tab-separated label🡪value pairs
    for row in rows:
        for cell in row:
            for sep in ['🡪','\t']:
                if sep in cell:
                    parts=cell.split(sep)
                    if len(parts)>=2:
                        label=parts[0].strip().lower()
                        value=sep.join(parts[1:]).strip()
                        if value and label: _map_label(label,value,data)
    
    # Items from table headers
    header_idx=None
    for i,row in enumerate(rows):
        row_text=' '.join(c.lower() for c in row)
        if any(kw in row_text for kw in ['sr. no','sr no','sr | no','topic no']):
            header_idx=i; break
    
    if header_idx is not None:
        headers=[c.lower().strip().replace('|','').replace('\n',' ').strip() for c in rows[header_idx]]
        for row in rows[header_idx+1:]:
            if len(row)>=len(headers) and any(c.strip() for c in row):
                item={}
                for ci,h in enumerate(headers):
                    if ci<len(row) and h: item[h]=row[ci].strip()
                if item and any(v for v in item.values() if v):
                    if '_items' not in data: data['_items']=[]
                    data['_items'].append(item)
    
    # Simple label:value pairs
    for row in rows:
        if len(row)<2: continue
        for cell in row:
            for sep in [':']:
                if sep in cell and '🡪' not in cell and '\t' not in cell:
                    parts=cell.split(sep,1)
                    label=parts[0].strip().lower()
                    value=parts[1].strip() if len(parts)>1 else ''
                    if value: _map_label(label,value,data)
    
    data['_raw_text']=full_text[:3000]
    data['_num_rows']=len(rows)
    return data

def _map_label(label, value, data):
    label=re.sub(r'\s+',' ',label.replace('|',' ').replace('\n',' ')).strip()
    field_map={
        'serial':['serial no','serial','sr no','sr. no','ref no','ref. no','complaint sr. no','project number','project no','ref. no.'],
        'date':['date','date of','dated'],
        'client_name':['customer','client name','name of customer','name of the customer','name :','client','name of client'],
        'mode_of_receipt':['mode of receipt','mode','receipt'],
        'description':['description','details','desc','object and variables','experiment title','title'],
        'remarks':['remarks','remark','comments','comment'],
        'reviewed_by':['reviewed by','reviewed'],
        'prepared_by':['prepared by','prepared'],
        'employee_name':['name of employee','employee name','name','name of the participant','name & designation','name of employees','name of employee'],
        'department':['department','dept'],
        'project':['project','project name','name of product','product name'],
        'supplier':['supplier','name of supplier','to,'],
        'product_type':['product type','product'],
        'address':['address'],
        'year':['year'],
        'month':['month'],
        'period':['period','quarter'],
        'designation':['designation','job title'],
        'qualification':['qualification','qualifications','qualifications required'],
        'experience':['experience'],
        'employee_id':['id no','id no.','employee id'],
        'training_date':['training date'],
        'course_name':['title of training','course name','topic','training'],
        'score':['score'],
        'result':['result'],
        'trainer':['trainer','incharge','trainer name'],
        'reason':['reason','reason for','reason of development'],
        'scope':['scope'],
        'objective':['objective','objectives'],
        'findings':['findings','output observed'],
        'conclusion':['conclusion'],
        'request_for':['request for'],
        'name_of_customer':['name of the customer','name of customer'],
        'change_type':['type of proposed change','change type'],
        'delivery_schedule':['delivery schedule','delivery'],
        'order_status':['order status','status'],
    }
    for field, patterns in field_map.items():
        if field not in data:
            for p in patterns:
                if p in label and len(p)>=3:
                    data[field]=value; return

# ============================================================================
# SUPABASE
# ============================================================================
def supabase_insert(records):
    url=f"{SUPABASE_URL}/rest/v1/records"
    data=json.dumps(records,ensure_ascii=False).encode('utf-8')
    req=urllib.request.Request(url,data=data,method='POST')
    req.add_header('apikey',SERVICE_KEY)
    req.add_header('Authorization',f'Bearer {SERVICE_KEY}')
    req.add_header('Content-Type','application/json')
    req.add_header('Prefer','return=representation')
    try:
        resp=urllib.request.urlopen(req, timeout=30)
        return json.loads(resp.read())
    except urllib.error.HTTPError as e:
        return {'_error':f'HTTP {e.code}: {e.read().decode()[:300]}'}

def supabase_get_existing():
    url=f"{SUPABASE_URL}/rest/v1/records?select=id,form_code,serial,status"
    req=urllib.request.Request(url)
    req.add_header('apikey',SERVICE_KEY)
    req.add_header('Authorization',f'Bearer {SERVICE_KEY}')
    req.add_header('Range','0-99999')
    return json.loads(urllib.request.urlopen(req, timeout=30).read())

def build_record(form_code, config, extracted):
    serial=extracted.get('serial','')
    if not serial: return None
    form_data={}
    for k,v in extracted.items():
        if k.startswith('_') or k=='serial': continue
        if isinstance(v,(list,dict)): form_data[k]=json.dumps(v,ensure_ascii=False)
        else: form_data[k]=str(v) if v else ''
    record={
        'form_code':form_code,
        'serial':serial,
        'form_name':config['name'],
        'status':'migrated',
        'section':config['section'],
        'section_name':SECTION_NAMES.get(config['section'],''),
        'form_data':form_data,
    }
    if 'date' in extracted: record['date']=extracted['date']
    return record

# ============================================================================
# MAIN — Two-phase: cache Drive listing, then process
# ============================================================================
def main():
    cmd=sys.argv[1] if len(sys.argv)>1 else 'analyze'
    token=get_access_token()
    print(f"🔑 Token refreshed. Running in {cmd} mode.")
    
    # Phase 1: Get all existing records from QBase
    existing=supabase_get_existing()
    existing_serials={r['serial'] for r in existing if r.get('serial')}
    print(f"📊 QBase has {len(existing)} existing records ({len(existing_serials)} unique serials)")
    
    # Phase 2: Cache or load Drive file listing
    CACHE_DIR.mkdir(parents=True, exist_ok=True)
    listing_path=CACHE_DIR/'drive_listing.json'
    
    if listing_path.exists():
        print("📂 Using cached Drive listing...")
        with open(listing_path) as f: all_files=json.load(f)
    else:
        print("📂 Scanning Google Drive (this may take a minute)...")
        root_contents=drive_api(f"/drive/v3/files?q='{DRIVE_ROOT}'+in+parents+and+trashed=false&fields=files(id,name,mimeType)&pageSize=200",token)
        section_folders=[f for f in root_contents.get('files',[]) if f['mimeType']=='application/vnd.google-apps.folder' and 'Records' in f.get('name','')]
        
        all_files = []  # list of {file_id, file_name, form_code, form_folder}
        for sf in section_folders:
            form_folders=drive_list_all(sf['id'], token)
            for ff in form_folders:
                folder_name=ff['name']
                config=FORM_CONFIG.get(folder_name)
                if not config: continue
                files_in_folder=drive_list_all(ff['id'], token)
                for f in files_in_folder:
                    if f.get('mimeType','') in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document','application/msword','application/vnd.google-apps.document']:
                        all_files.append({
                            'file_id':f['id'],
                            'file_name':f['name'],
                            'form_code':config['code'],
                            'form_folder':folder_name,
                        })
        
        with open(listing_path,'w') as f: json.dump(all_files,f,indent=2,ensure_ascii=False)
        print(f"   Cached {len(all_files)} files from Drive")
    
    print(f"📄 Total DOCX files to process: {len(all_files)}")
    
    # Phase 3: Process each file
    total_files=len(all_files)
    total_extracted=0
    total_inserted=0
    total_skipped=0
    total_errors=[]
    migration_log=[]
    batch=[]
    
    for i, finfo in enumerate(all_files):
        file_name=finfo['file_name']
        form_code=finfo['form_code']
        file_id=finfo['file_id']
        
        # Extract serial from filename
        m=re.search(r'(F[/\-]\d+[-]\d+)',file_name)
        file_serial=m.group(1).replace('\\','/') if m else None
        if file_serial:
            m2=re.match(r'F(\d{2})[-/](\d+)',file_serial)
            if m2: file_serial=f"F/{m2.group(1)}-{m2.group(2)}"
        
        # Skip if already in QBase
        if file_serial and file_serial in existing_serials:
            total_skipped+=1
            continue
        
        # Download file
        safe_name=file_name.replace('/','-').replace('\\','-').replace(' ','_')
        cache_dir=CACHE_DIR/form_code
        cache_dir.mkdir(parents=True, exist_ok=True)
        outpath=str(cache_dir/safe_name)
        
        if os.path.exists(outpath):
            dl_path=outpath
        else:
            try:
                token=get_access_token()
                dl_path=download_docx(file_id, token, outpath)
                if not dl_path:
                    total_errors.append(f"{form_code}/{file_name}: skip (spreadsheet)")
                    continue
            except Exception as e:
                total_errors.append(f"{form_code}/{file_name}: download error — {str(e)[:80]}")
                continue
        
        # Parse DOCX
        try:
            extracted=parse_qms_docx(dl_path)
            extracted_serial=extracted.get('serial','')
            if not extracted_serial and file_serial:
                extracted_serial=file_serial
                extracted['serial']=file_serial
            if not extracted_serial:
                total_errors.append(f"{form_code}/{file_name}: no serial")
                continue
            
            total_extracted+=1
            config=FORM_CONFIG.get(finfo['form_folder'],{'code':form_code,'section':1,'name':form_code})
            record=build_record(form_code, config, extracted)
            if not record:
                total_errors.append(f"{form_code}/{file_name}: could not build record")
                continue
            
            # Print progress
            n_fields=len([k for k in extracted.keys() if not k.startswith('_')])
            print(f"  [{i+1}/{total_files}] {file_serial:12} → {n_fields} fields  ({form_code})")
            sys.stdout.flush()
            
            if cmd=='migrate':
                batch.append(record)
                if len(batch)>=10:
                    result=supabase_insert(batch)
                    if isinstance(result,dict) and '_error' in result:
                        total_errors.append(f"BATCH INSERT ERROR: {result['_error'][:100]}")
                    else:
                        total_inserted+=len(batch)
                    batch=[]
            
            migration_log.append({'form_code':form_code,'serial':extracted_serial,'file':file_name,'fields':n_fields,'status':'inserted' if cmd=='migrate' else 'would_insert'})
            
        except Exception as e:
            total_errors.append(f"{form_code}/{file_name}: parse error — {str(e)[:80]}")
    
    # Flush remaining batch
    if cmd=='migrate' and batch:
        result=supabase_insert(batch)
        if isinstance(result,dict) and '_error' in result:
            total_errors.append(f"FINAL BATCH ERROR: {result['_error'][:100]}")
        else:
            total_inserted+=len(batch)
    
    # Summary
    print(f"\n{'='*60}")
    print(f"📊 MIGRATION SUMMARY ({cmd.upper()} mode)")
    print(f"   Total files on Drive: {total_files}")
    print(f"   Extracted: {total_extracted}")
    print(f"   Skipped (already in QBase): {total_skipped}")
    if cmd=='migrate':
        print(f"   Inserted: {total_inserted}")
    print(f"   Errors: {len(total_errors)}")
    if total_errors:
        for e in total_errors[:10]:
            print(f"     ⚠️ {e}")
    
    # Save log
    with open(MIGRATION_LOG,'w') as f:
        json.dump(migration_log,f,indent=2,ensure_ascii=False)
    print(f"\n💾 Log saved to {MIGRATION_LOG}")

if __name__=='__main__':
    main()