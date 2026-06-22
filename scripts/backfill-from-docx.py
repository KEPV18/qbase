#!/usr/bin/env python3
from __future__ import annotations
"""
QBase DOCX ETL Backfill Script
================================
Scans /home/Kepv/Downloads/03.Records/ recursively, extracts data from .docx/.doc files,
maps to canonical template-expected JSONB schemas, generates SQL UPDATE statements.

Usage:
  python3 backfill-from-docx.py [--dry-run] [--form F/12]

Output:
  supabase/migrations/20260622_canonical_records_backfill.sql
"""
import os, sys, re, json, subprocess, tempfile
from pathlib import Path
from docx import Document

BASE_DIR = Path("/home/Kepv/Downloads/03.Records/")
OUTPUT_FILE = Path("/mnt/ahmed/Projects/qbase/supabase/migrations/20260622_canonical_records_backfill.sql")
CONVERT_DIR = Path("/tmp/docx_converted")  # .doc → .docx converted files
CONVERT_DIR.mkdir(exist_ok=True)

# ── Serial extraction from filename ──────────────────────────────
def extract_serial(filename: str) -> str | None:
    """Extract F/XX-NNN from filename like F_12-001.docx or F／08-002.docx"""
    m = re.match(r'F[_／](\d+)-(\d+)\.(?:docx?|doc)$', filename)
    if m:
        return f"F/{m.group(1)}-{m.group(2)}"
    return None

def extract_form_code(filename: str) -> str | None:
    s = extract_serial(filename)
    if s:
        return s.split("-")[0]  # F/12
    return None

# ── .doc → .docx conversion via LibreOffice ──────────────────────
_converted_cache = {}

def convert_doc_to_docx(doc_path: str) -> str | None:
    """Convert .doc to .docx using LibreOffice headless."""
    if doc_path in _converted_cache:
        return _converted_cache[doc_path]
    
    outdir = str(CONVERT_DIR)
    try:
        result = subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "docx", "--outdir", outdir, doc_path],
            capture_output=True, text=True, timeout=30
        )
        converted = os.path.join(outdir, Path(doc_path).stem + ".docx")
        if os.path.exists(converted):
            _converted_cache[doc_path] = converted
            return converted
    except Exception as e:
        print(f"  ⚠️ Convert failed: {doc_path}: {e}")
    return None

# ── Open DOCX (convert if needed) ────────────────────────────────
def open_document(filepath: str) -> Document | None:
    if filepath.endswith('.doc') and not filepath.endswith('.docx'):
        converted = convert_doc_to_docx(filepath)
        if not converted:
            return None
        filepath = converted
    try:
        return Document(filepath)
    except Exception as e:
        print(f"  ⚠️ Open failed: {filepath}: {e}")
        return None

# ── Table helpers ────────────────────────────────────────────────
def dedup_cells(row_cells) -> list[str]:
    """Get unique cell texts (handles merged cells)."""
    seen = set()
    result = []
    for cell in row_cells:
        txt = cell.text.strip()
        if txt and txt not in seen:
            result.append(txt)
            seen.add(txt)
    return result

def all_cells(row_cells) -> list[str]:
    """Get ALL cell texts (including merged duplicates)."""
    return [c.text.strip() for c in row_cells]

def find_row_by_label(table, label_patterns: list[str], search_col=0) -> dict | None:
    """Find a row where col 0 matches any of the label patterns, return {col: val} mapping."""
    for row in table.rows:
        cells = all_cells(row.cells)
        if len(cells) > search_col:
            cell_text = cells[search_col].strip().lower()
            for pattern in label_patterns:
                if pattern.lower() in cell_text:
                    return cells
    return None

def extract_kv_from_table(table) -> dict:
    """Extract key-value pairs from a 2-column table (label → value)."""
    result = {}
    for row in table.rows:
        cells = all_cells(row.cells)
        if len(cells) >= 2:
            label = cells[0].strip()
            value = cells[1].strip()
            if label and value:
                result[label] = value
    return result

def extract_row_with_label_value(table, label: str) -> str:
    """Find row where col 0 = label, return col 1 value."""
    row_cells = find_row_by_label(table, [label])
    if row_cells and len(row_cells) > 1:
        return row_cells[1].strip()
    return ""

def extract_after_arrow(text: str) -> str:
    """Extract value after '🡪' or '➜' or '→' arrow in a cell."""
    for arrow in ["🡪", "➜", "→", "=>"]:
        if arrow in text:
            parts = text.split(arrow, 1)
            if len(parts) > 1:
                return parts[1].strip()
    return text.strip()

def extract_label_value_arrow(cell_text: str) -> tuple[str, str]:
    """From 'Label\t🡪 Value' return ('Label', 'Value')."""
    text = cell_text.strip()
    label = text
    value = ""
    for arrow in ["🡪", "➜", "→", "=>"]:
        if arrow in text:
            parts = text.split(arrow, 1)
            label = parts[0].strip().replace("\t", "")
            value = parts[1].strip() if len(parts) > 1 else ""
            return label, value
    # Try tab-separated
    if "\t" in text:
        parts = text.split("\t", 1)
        return parts[0].strip(), parts[1].strip() if len(parts) > 1 else ""
    return label, ""

# ── Form-specific extractors ─────────────────────────────────────
# Each extractor takes a Document object and returns a dict matching template keys.

def extract_F08(doc: Document) -> dict:
    """F/08 Order Form — 1 table, header row + customer info + line items."""
    data = {}
    if not doc.tables:
        return data
    t = doc.tables[0]
    
    # Row 1: Sr.No + Date
    for row in t.rows:
        cells = all_cells(row.cells)
        for cell in cells:
            text = cell.strip()
            if "Sr. No" in text or "Sr.No" in text:
                _, serial = extract_label_value_arrow(text)
                if serial:
                    data["serial"] = serial
            if "Date" in text and "🡪" in text:
                _, date_val = extract_label_value_arrow(text)
                if date_val:
                    data["date"] = date_val
    
    # Find customer, mode of receipt, etc from labeled rows
    for row in t.rows:
        cells = all_cells(row.cells)
        if len(cells) >= 3:
            label = cells[0].strip().lower()
            if label == "customer" and len(cells) > 2:
                data["client_name"] = cells[2].strip()
            elif "mode of receipt" in label and len(cells) > 2:
                data["mode_of_receipt"] = cells[2].strip()
            elif "delivery schedule" in label and len(cells) > 2:
                data["delivery_schedule"] = cells[2].strip()
            elif "despatch" in label and len(cells) > 2:
                data["despatch_date"] = cells[2].strip()
            elif "order status" in label and len(cells) > 2:
                data["order_status"] = cells[2].strip()
            elif "test certificate" in label and len(cells) > 2:
                data["test_certificate_required"] = cells[2].strip()
            elif "bill no" in label and len(cells) > 2:
                data["bill_no"] = cells[2].strip()
            elif "complies" in label and len(cells) > 2:
                data["complies"] = cells[2].strip()
            elif "reviewed" in label and len(cells) > 2:
                data["reviewed_by"] = cells[2].strip()
            elif "remarks" in label and len(cells) > 2:
                data["remarks"] = cells[2].strip()
    
    return data

def extract_F09(doc: Document) -> dict:
    """F/09 Customer Complaint Report."""
    data = {}
    if not doc.tables:
        return data
    t = doc.tables[0]
    
    for row in t.rows:
        cells = all_cells(row.cells)
        for cell in cells:
            text = cell.strip()
            label, value = extract_label_value_arrow(text)
            if "complaint" in label.lower() and "sr" in label.lower() and value:
                data["serial"] = value
            elif label.lower() == "date" and value:
                data["date"] = value
            elif "received by" in label.lower():
                data["received_by"] = value
            elif "receipt" in label.lower() and "date" in label.lower():
                data["receipt_date"] = value
            elif "receipt" in label.lower() and "time" in label.lower():
                data["receipt_time"] = value
            elif "mode of receipt" in label.lower():
                data["mode_of_receipt"] = value
            elif "nature" in label.lower() and "complaint" in label.lower():
                data["complaint_nature"] = value
            elif "client" in label.lower() and "name" in label.lower():
                data["client_name"] = value
            elif "product" in label.lower() and "type" in label.lower():
                data["product_type"] = value
            elif "analysed" in label.lower():
                data["analysed_by"] = value
            elif "closed" in label.lower() and "by" in label.lower():
                data["closed_by"] = value
            elif "customer informed" in label.lower():
                data["customer_informed_vide"] = value
            elif "clientplatform" in label.lower() or "client platform" in label.lower():
                data["clientplatform_confirmation"] = value
    
    return data

def extract_F10(doc: Document) -> dict:
    """F/10 Customer Feedback Form."""
    data = {}
    if not doc.tables:
        return data
    t = doc.tables[0]
    
    for row in t.rows:
        cells = all_cells(row.cells)
        for cell in cells:
            text = cell.strip()
            label, value = extract_label_value_arrow(text)
            if "date" in label.lower() and value:
                data["date"] = value
            elif "year" in label.lower() and value:
                data["year"] = value
            elif "name" in label.lower() and "address" not in label.lower() and value:
                data["client_name"] = value
            elif "address" in label.lower() and value:
                data["address"] = value
            elif "project" in label.lower() and value:
                data["project_name"] = value
            elif "reviewed" in label.lower() and value:
                data["reviewed_by"] = value
            elif "corrective" in label.lower() and "ref" in label.lower():
                data["corrective_action_ref"] = value
            elif "distributor" in label.lower() and "sign" in label.lower():
                data["distributor_signature"] = value
    
    return data

def extract_F11(doc: Document) -> dict:
    """F/11 Production Plan — header fields + line items table."""
    data = {}
    if not doc.tables:
        return data
    t = doc.tables[0]
    
    # Extract header fields
    for row in t.rows:
        cells = all_cells(row.cells)
        for cell in cells:
            text = cell.strip()
            label, value = extract_label_value_arrow(text)
            if "sr. no" in label.lower() and value:
                data["serial"] = value
            elif label.lower() == "date" and value:
                data["date"] = value
            elif "month" in label.lower() and value:
                data["month"] = value
            elif "updated" in label.lower() and "by" in label.lower():
                data["updated_by"] = value
            elif "remarks" in label.lower() and value:
                data["remarks"] = value
    
    # Extract items array from data rows
    items = []
    # Find header row with product/batch columns
    header_row_idx = -1
    for ri, row in enumerate(t.rows):
        cells_text = [c.text.strip().lower() for c in row.cells]
        joined = " ".join(cells_text)
        if "product" in joined and ("batch" in joined or "plan" in joined or "qty" in joined):
            header_row_idx = ri
            break
    
    if header_row_idx >= 0:
        # Map header columns
        headers = [c.text.strip().lower() for c in t.rows[header_row_idx].cells]
        for row in t.rows[header_row_idx + 1:]:
            cells = all_cells(row.cells)
            if not any(c.strip() for c in cells):
                continue  # skip empty rows
            item = {}
            for ci, h in enumerate(headers):
                if ci < len(cells) and cells[ci].strip():
                    h_clean = h.strip()
                    if "product" in h_clean:
                        item["product"] = cells[ci].strip()
                    elif "batch" in h_clean:
                        item["batchNo"] = cells[ci].strip()
                    elif "plan" in h_clean and "date" in h_clean:
                        item["planDate"] = cells[ci].strip()
                    elif "plan" in h_clean and "completion" in h_clean:
                        item["planCompletion"] = cells[ci].strip()
                    elif "plan" in h_clean and "size" in h_clean:
                        item["planSize"] = cells[ci].strip()
                    elif "actual" in h_clean and "date" in h_clean:
                        item["actualDate"] = cells[ci].strip()
                    elif "actual" in h_clean and "qty" in h_clean:
                        item["actualQty"] = cells[ci].strip()
                    elif "yield" in h_clean:
                        item["yieldPercent"] = cells[ci].strip()
            if item:
                items.append(item)
    
    if items:
        data["items"] = items
    return data

def extract_F12(doc: Document) -> dict:
    """F/12 Disposal of Non-Conforming Products — header + table with 13 cols.
    Column layout (13 cells, some merged):
    0=Sr.No, 1=Date, 2=Stage, 3=Product, 4=Id.No, 5-6=Reason(merged), 7=Qty, 8=DisposalAction, 9-10=Re-Inspection(merged), 11=Qty.OK, 12=Signature
    """
    data = {}
    if not doc.tables:
        return data
    t = doc.tables[0]
    
    # Extract header fields (serial, month) from rows 0-1
    for row in t.rows[:3]:
        cells = all_cells(row.cells)
        for cell in cells:
            text = cell.strip()
            label, value = extract_label_value_arrow(text)
            if "sr. no" in label.lower() and value:
                data["serial"] = value
            elif "month" in label.lower() and value:
                data["month"] = value
    
    # Find header row (the one with "Sr. No" + "Date" + "Stage")
    header_row_idx = -1
    for ri, row in enumerate(t.rows):
        cells_text = [c.text.strip().lower() for c in row.cells]
        if cells_text[0].startswith("sr. no") and "date" in cells_text and "stage" in cells_text:
            header_row_idx = ri
            break
    
    if header_row_idx < 0:
        return data
    
    # Build column map from header row positions
    # Use the 13-cell layout: position → field name
    COL_MAP = {
        0: "srNo",
        1: "date", 
        2: "stage",
        3: "productName",
        4: "idNo",
        5: "reason",     # merged with 6
        6: "reason",     # duplicate (merged cell)
        7: "qty",
        8: "disposalAction",
        9: "reInspection",  # merged with 10
        10: "reInspection", # duplicate (merged cell)
        11: "qtyOk",
        12: "signature",
    }
    
    items = []
    for row in t.rows[header_row_idx + 1:]:
        cells = all_cells(row.cells)
        if not any(c.strip() for c in cells):
            continue  # skip empty rows
        
        item = {}
        for ci in range(min(len(cells), 13)):
            field = COL_MAP.get(ci)
            if field and cells[ci].strip():
                # Don't overwrite with duplicate merged cell value
                if field not in item:
                    item[field] = cells[ci].strip()
        
        if item.get("srNo") or item.get("date") or item.get("productName"):
            items.append(item)
    
    if items:
        data["items"] = items
    
    # Department and authorised_signature not in DOCX table — leave empty
    data.setdefault("department", "")
    data.setdefault("authorised_signature", "")
    
    return data

def extract_F13(doc: Document) -> dict:
    """F/13 Purchase Order."""
    data = {}
    if not doc.tables:
        return data
    t = doc.tables[0]
    
    for row in t.rows:
        cells = all_cells(row.cells)
        for cell in cells:
            text = cell.strip()
            label, value = extract_label_value_arrow(text)
            if "purchase order no" in label.lower() and value:
                data["serial"] = value
            elif label.lower() == "date" and value:
                data["date"] = value
            elif label.lower().startswith("to") and value:
                data["to"] = value
            elif "total" in label.lower() and "amount" in label.lower():
                data["total_amount"] = value
            elif "ordered" in label.lower() and "by" in label.lower():
                data["ordered_by"] = value
            elif "authorised" in label.lower() and "by" in label.lower():
                data["authorised_by"] = value
            elif "terms" in label.lower() and value:
                data["terms"] = value
    
    return data

def extract_F14(doc: Document) -> dict:
    """F/14 Indent and Incoming Inspection Record."""
    data = {}
    if not doc.tables:
        return data
    t = doc.tables[0]
    
    for row in t.rows:
        cells = all_cells(row.cells)
        for cell in cells:
            text = cell.strip()
            label, value = extract_label_value_arrow(text)
            if "f/14" in label.lower() and "rev" in label.lower() and value:
                data["serial"] = value
            elif "date" in label.lower() and value and "serial" not in data:
                pass  # dates are in table rows
    
    # The table has Date | Item Description | Qty | Name Of Supplier | Inspection Status | Inspected By
    items = []
    header_row_idx = -1
    for ri, row in enumerate(t.rows):
        cells_text = [c.text.strip().lower() for c in row.cells]
        joined = " ".join(cells_text)
        if "date" in joined and "item" in joined and "qty" in joined:
            header_row_idx = ri
            break
    
    if header_row_idx >= 0:
        for row in t.rows[header_row_idx + 1:]:
            cells = all_cells(row.cells)
            if not any(c.strip() for c in cells):
                continue
            unique = dedup_cells(row.cells)
            if len(unique) >= 3:
                items.append({
                    "date": unique[0],
                    "itemDescription": unique[1] if len(unique) > 1 else "",
                    "qty": unique[2] if len(unique) > 2 else "",
                    "supplier": unique[3] if len(unique) > 3 else "",
                    "inspectionStatus": unique[4] if len(unique) > 4 else "",
                    "inspectedBy": unique[5] if len(unique) > 5 else "",
                })
    
    if items:
        data["items"] = items
    
    # Extract signature fields
    for row in t.rows:
        cells = all_cells(row.cells)
        for cell in cells:
            text = cell.strip().lower()
            if "prepared" in text and "by" in text:
                _, val = extract_label_value_arrow(cell)
                data["prepared_by"] = val
            elif "checked" in text and "by" in text:
                _, val = extract_label_value_arrow(cell)
                data["checked_by"] = val
    
    return data

def extract_F15(doc: Document) -> dict:
    """F/15 Approved Vendor List — multi-row table."""
    data = {}
    if not doc.tables:
        return data
    t = doc.tables[0]
    
    # Row 0: title + serial
    for row in t.rows[:2]:
        cells = all_cells(row.cells)
        for cell in cells:
            text = cell.strip()
            label, value = extract_label_value_arrow(text)
            if "f/15" in label.lower() and "rev" in label.lower() and value:
                data["serial"] = value
    
    # Vendor list rows: Date of Approval | Name of Supplier | Scope of Supply | Approval Criteria | Remarks
    items = []
    header_row_idx = -1
    for ri, row in enumerate(t.rows):
        cells_text = [c.text.strip().lower() for c in row.cells]
        joined = " ".join(cells_text)
        if "date" in joined and "supplier" in joined and "scope" in joined:
            header_row_idx = ri
            break
    
    if header_row_idx >= 0:
        for row in t.rows[header_row_idx + 1:]:
            cells = all_cells(row.cells)
            if not any(c.strip() for c in cells):
                continue
            unique = dedup_cells(row.cells)
            if len(unique) >= 2:
                items.append({
                    "dateOfApproval": unique[0],
                    "supplierName": unique[1] if len(unique) > 1 else "",
                    "scopeOfSupply": unique[2] if len(unique) > 2 else "",
                    "approvalCriteria": unique[3] if len(unique) > 3 else "",
                    "remarks": unique[4] if len(unique) > 4 else "",
                })
    
    if items:
        data["items"] = items
    
    return data

def extract_F16(doc: Document) -> dict:
    """F/16 Vendor Rating/Registration Form."""
    data = {}
    if not doc.tables:
        return data
    t = doc.tables[0]
    
    for row in t.rows:
        cells = all_cells(row.cells)
        if len(cells) >= 2:
            label = cells[0].strip().lower()
            value = cells[1].strip() if len(cells) > 1 else ""
            if not value and len(cells) > 2:
                value = cells[2].strip()
            
            if "name" in label and "auth" not in label:
                data["name"] = value
            elif "address" in label:
                data["address"] = value
            elif "contact" in label and "person" in label:
                data["contact_person"] = value
            elif "tel" in label or "fax" in label:
                data["tel_fax"] = value
            elif "mobile" in label:
                data["mobile_no"] = value
            elif "residence" in label:
                data["residence_no"] = value
            elif "employee" in label and "strength" in label:
                data["employee_strength"] = value
            elif "sites" in label or "branch" in label:
                data["sites_branches"] = value
            elif "sister" in label:
                data["sister_concerns"] = value
            elif "past" in label and "experience" in label:
                data["past_experience"] = value
            elif "reference" in label:
                data["reference"] = value
            elif "association" in label and "year" in label:
                data["association_years"] = value
            elif "approval" in label and "reason" in label:
                data["approval_reason"] = value
            elif "vendor" in label and "auth" in label and "name" in label:
                data["vendor_auth_name"] = value
            elif "vendor" in label and "auth" in label and "design" in label:
                data["vendor_auth_designation"] = value
            elif "authorised" in label and "date" in label:
                data["authorised_date"] = value
            elif "authorised" in label and "by" in label:
                data["authorised_by"] = value
            elif "f/16" in label and "rev" in label:
                _, serial = extract_label_value_arrow(cells[0])
                if serial:
                    data["serial"] = serial
    
    return data

def extract_F17(doc: Document) -> dict:
    """F/17 QA Test Request Slip."""
    data = {}
    if not doc.tables:
        return data
    t = doc.tables[0]
    
    for row in t.rows:
        cells = all_cells(row.cells)
        for cell in cells:
            text = cell.strip()
            label, value = extract_label_value_arrow(text)
            ll = label.lower()
            if "sr. no" in ll and value:
                data["serial"] = value
            elif ll == "date" and value:
                data["date"] = value
            elif "product" in ll and "name" in ll:
                data["product_name"] = value
            elif "batch" in ll and "no" in ll:
                data["batch_no"] = value
            elif "sample" in ll and "desc" in ll:
                data["sample_description"] = value
            elif "sample" in ll and "qty" in ll:
                data["sample_qty"] = value
            elif "request" in ll and "from" in ll:
                data["request_from"] = value
            elif "request" in ll and "to" in ll:
                data["request_to"] = value
            elif "requested" in ll and "by" in ll:
                data["requested_by"] = value
            elif "test result" in ll and "ref" in ll:
                data["test_result_ref"] = value
            elif "signature" in ll and "approved" in ll:
                data["signature_approved"] = value
            elif "signature" in ll and "requested" in ll:
                data["signature_requested"] = value
    
    return data

def extract_F18(doc: Document) -> dict:
    """F/18 Product Re-Call Report — table with recall data."""
    data = {}
    if not doc.tables:
        return data
    t = doc.tables[0]
    
    # Row 0: title + serial
    for row in t.rows[:2]:
        cells = all_cells(row.cells)
        for cell in cells:
            text = cell.strip()
            label, value = extract_label_value_arrow(text)
            if "f/18" in label.lower() and "rev" in label.lower() and value:
                data["serial"] = value
    
    # Data row (R2 in our sample): Date | Name of Products | Reference Inward No | Qty Taken | Products Identified By | Released By | Requested By | Verified By | Verified On | Status | Entry Closed On | Entry Closed By
    for row in t.rows:
        cells = all_cells(row.cells)
        unique = dedup_cells(row.cells)
        if len(unique) >= 3 and any(k in unique[0].lower() for k in ["date", "/01/", "/02/"]):
            data["date"] = unique[0]
            data["authorised_by"] = unique[7] if len(unique) > 7 else ""
            data["department"] = unique[9] if len(unique) > 9 else ""
    
    return data

def extract_F19(doc: Document) -> dict:
    """F/19 Product Description Form — key-value rows."""
    data = {}
    if not doc.tables:
        return data
    t = doc.tables[0]
    
    for row in t.rows:
        cells = all_cells(row.cells)
        if len(cells) >= 2:
            label = cells[0].strip()
            value = cells[1].strip()
            if label and value:
                data[label] = value
    
    return data

def extract_F20(doc: Document) -> dict:
    """F/20 Management Review Meeting Agenda."""
    data = {}
    if not doc.tables:
        return data
    t = doc.tables[0]
    
    for row in t.rows:
        cells = all_cells(row.cells)
        for cell in cells:
            text = cell.strip()
            label, value = extract_label_value_arrow(text)
            if "meeting" in label.lower() and "date" in label.lower():
                data["meeting_date"] = value
            elif "meeting" in label.lower() and "time" in label.lower():
                data["meeting_time"] = value
            elif "meeting" in label.lower() and "place" in label.lower():
                data["meeting_place"] = value
            elif "approved" in label.lower() and "by" in label.lower():
                data["approved_by"] = value
            elif "date" in label.lower() and value:
                data["date"] = value
    
    return data

def extract_F21(doc: Document) -> dict:
    """F/21 Management Review Meeting Minutes."""
    data = {}
    
    # Try table extraction first
    if doc.tables:
        t = doc.tables[0]
        for row in t.rows:
            cells = all_cells(row.cells)
            for cell in cells:
                text = cell.strip()
                label, value = extract_label_value_arrow(text)
                if "meeting" in label.lower() and "date" in label.lower():
                    data["meeting_date"] = value
                elif "meeting" in label.lower() and "time" in label.lower():
                    data["meeting_time"] = value
                elif "meeting" in label.lower() and "place" in label.lower():
                    data["meeting_place"] = value
                elif "approved" in label.lower() and "by" in label.lower():
                    data["approved_by"] = value
                # Also try direct text parsing
                if "meeting date" in text.lower():
                    m = re.search(r'\d{2}/\d{2}/\d{4}', text)
                    if m:
                        data["meeting_date"] = m.group(0)
                if "meeting time" in text.lower():
                    m = re.search(r'\d{1,2}:\d{2}', text)
                    if m:
                        data["meeting_time"] = m.group(0)
                if "meeting place" in text.lower():
                    parts = text.split(":", 1)
                    if len(parts) > 1:
                        data["meeting_place"] = parts[1].strip()
    
    # Extract discussion points from paragraphs
    full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    if full_text:
        data["discussion_points"] = full_text[:2000]
    
    # If we got at least one field, return
    if data:
        return data
    
    # Fallback: extract any table cell content as data using regex patterns
    if doc.tables:
        for row in doc.tables[0].rows:
            for cell in row.cells:
                text = cell.text.strip()
                if not text:
                    continue
                # Extract date patterns
                m = re.search(r'meeting\s*date[:\s]*([0-9/]{8,10})', text, re.I)
                if m:
                    data["meeting_date"] = m.group(1).strip()
                m = re.search(r'meeting\s*time[:\s]*([0-9:apAPmM\s]{4,})', text, re.I)
                if m:
                    val = m.group(1).strip()
                    if val:
                        data["meeting_time"] = val
                m = re.search(r'meeting\s*place[:\s]*(\S+[^:]*)', text, re.I)
                if m:
                    val = m.group(1).strip()
                    if val:
                        data["meeting_place"] = val
                m = re.search(r'approved\s*by[:\s]*([^\n]{2,})', text, re.I)
                if m:
                    data["approved_by"] = m.group(1).strip()
    
    return data

def extract_F22(doc: Document) -> dict:
    """F/22 Corrective Action Report."""
    data = {}
    if not doc.tables:
        return data
    t = doc.tables[0]
    
    for row in t.rows:
        cells = all_cells(row.cells)
        for cell in cells:
            text = cell.strip()
            label, value = extract_label_value_arrow(text)
            ll = label.lower()
            if "sr. no" in ll and value:
                data["serial"] = value
            elif ll == "date" and value:
                data["date"] = value
            elif "department" in ll and value:
                data["department"] = value
            elif "prepared" in ll and "by" in ll:
                data["prepared_by"] = value
            elif "approved" in ll and "by" in ll:
                data["approved_by"] = value
            elif "responsible" in ll:
                data["responsible"] = value
            elif "target" in ll and "date" in ll:
                data["target_date"] = value
            elif "verification" in ll and "date" in ll:
                data["verification_date"] = value
            elif "verification" in ll:
                data["verification"] = value
            elif "inprocess" in ll and "specify" in ll:
                data["inprocess_specify"] = value
            elif "others" in ll and "specify" in ll:
                data["others_specify"] = value
    
    # Extract root_cause, nc_description, corrective_action from text
    for row in t.rows:
        cells = all_cells(row.cells)
        unique = dedup_cells(row.cells)
        for cell in unique:
            cl = cell.lower()
            if "root cause" in cl:
                idx = unique.index(cell)
                if idx + 1 < len(unique):
                    data["root_cause"] = unique[idx + 1]
            elif "non-conformit" in cl or "non conformit" in cl:
                idx = unique.index(cell)
                if idx + 1 < len(unique):
                    data["nc_description"] = unique[idx + 1]
            elif "corrective action" in cl:
                idx = unique.index(cell)
                if idx + 1 < len(unique):
                    data["corrective_action"] = unique[idx + 1]
    
    return data

def extract_F23(doc: Document) -> dict:
    """F/23 Master List of Records — multi-row table."""
    data = {}
    if not doc.tables:
        return data
    t = doc.tables[0]
    
    # Find serial
    for row in t.rows[:2]:
        cells = all_cells(row.cells)
        for cell in cells:
            _, value = extract_label_value_arrow(cell.strip())
            if "f/23" in value.lower():
                data["serial"] = value
    
    # Header: Record No. | Title Of Record | Format No. | Frequency | Method Of Filing | Access | Storage Place | Retention Period | Person Responsible
    items = []
    header_row_idx = -1
    for ri, row in enumerate(t.rows):
        cells_text = [c.text.strip().lower() for c in row.cells]
        joined = " ".join(cells_text)
        if "record no" in joined and "title" in joined:
            header_row_idx = ri
            break
    
    if header_row_idx >= 0:
        for row in t.rows[header_row_idx + 1:]:
            cells = all_cells(row.cells)
            if not any(c.strip() for c in cells):
                continue
            unique = dedup_cells(row.cells)
            if len(unique) >= 2:
                items.append({
                    "recordNo": unique[0],
                    "title": unique[1] if len(unique) > 1 else "",
                    "formatNo": unique[2] if len(unique) > 2 else "",
                    "frequency": unique[3] if len(unique) > 3 else "",
                    "methodOfFiling": unique[4] if len(unique) > 4 else "",
                    "access": unique[5] if len(unique) > 5 else "",
                    "storagePlace": unique[6] if len(unique) > 6 else "",
                    "retentionPeriod": unique[7] if len(unique) > 7 else "",
                    "personResponsible": unique[8] if len(unique) > 8 else "",
                })
    
    if items:
        data["items"] = items
    
    # Extract date, department, authorised_signature
    for row in t.rows:
        cells = all_cells(row.cells)
        for cell in cells:
            text = cell.strip()
            label, value = extract_label_value_arrow(text)
            if "date" in label.lower() and value and "date" not in data:
                data["date"] = value
            elif "department" in label.lower() and value:
                data["department"] = value
            elif "authorised" in label.lower() and "sign" in label.lower():
                data["authorised_signature"] = value
    
    return data

def extract_F24(doc: Document) -> dict:
    """F/24 Objectives & Targets."""
    data = {}
    if not doc.tables:
        return data
    t = doc.tables[0]
    
    for row in t.rows:
        cells = all_cells(row.cells)
        for cell in cells:
            text = cell.strip()
            label, value = extract_label_value_arrow(text)
            ll = label.lower()
            if "department" in ll and value:
                data["department"] = value
            elif "year" in ll and value:
                data["year"] = value
            elif "signature" in ll and value:
                data["signature"] = value
            elif "f/24" in ll and "rev" in ll and value:
                data["serial"] = value
    
    # Extract objectives items
    items = []
    for row in t.rows:
        cells = all_cells(row.cells)
        unique = dedup_cells(row.cells)
        if len(unique) >= 3 and not any(k in unique[0].lower() for k in ["department", "year", "signature", "f/24", "objectives"]):
            items.append({
                "objective": unique[0],
                "target": unique[1] if len(unique) > 1 else "",
                "status": unique[2] if len(unique) > 2 else "",
            })
    
    if items:
        data["items"] = items
    
    return data

def extract_F25(doc: Document) -> dict:
    """F/25 Audit Plan."""
    data = {}
    if not doc.tables:
        return data
    t = doc.tables[0]
    
    for row in t.rows:
        cells = all_cells(row.cells)
        for cell in cells:
            text = cell.strip()
            label, value = extract_label_value_arrow(text)
            ll = label.lower()
            if "sr. no" in ll and value:
                data["serial"] = value
            elif "date" in ll and value and "date" not in data:
                data["date"] = value
            elif "from" in ll and value:
                data["from"] = value
            elif ll == "to" and value:
                data["to"] = value
            elif "prepared" in ll and "by" in ll:
                data["prepared_by"] = value
            elif "approved" in ll and "by" in ll:
                data["approved_by"] = value
            elif "last audit" in ll and "plan no" in ll:
                data["last_audit_plan_no"] = value
            elif "last audit" in ll and "plan date" in ll:
                data["last_audit_plan_date"] = value
            elif "last audit" in ll and "month" in ll:
                data["last_audit_month"] = value
    
    return data

def extract_F28(doc: Document) -> dict:
    """F/28 Training Attendance Sheet."""
    data = {}
    if not doc.tables:
        return data
    t = doc.tables[0]
    
    for row in t.rows:
        cells = all_cells(row.cells)
        for cell in cells:
            text = cell.strip()
            label, value = extract_label_value_arrow(text)
            ll = label.lower()
            if "sr. no" in ll and value:
                data["serial"] = value
            elif "date" in ll and value and "date" not in data:
                data["date"] = value
            elif "training" in ll and "topic" in ll:
                data["training_topic"] = value
            elif "trainer" in ll and value:
                data["trainer"] = value
            elif "conducted" in ll and "by" in ll:
                data["conducted_by"] = value
    
    # Extract attendance items (employee names + signatures)
    items = []
    header_row_idx = -1
    for ri, row in enumerate(t.rows):
        cells_text = [c.text.strip().lower() for c in row.cells]
        joined = " ".join(cells_text)
        if "name" in joined and ("sign" in joined or "attendance" in joined):
            header_row_idx = ri
            break
    
    if header_row_idx >= 0:
        for row in t.rows[header_row_idx + 1:]:
            cells = all_cells(row.cells)
            if not any(c.strip() for c in cells):
                continue
            unique = dedup_cells(row.cells)
            if len(unique) >= 1 and unique[0]:
                items.append({
                    "name": unique[0],
                    "signature": unique[1] if len(unique) > 1 else "",
                })
    
    if items:
        data["items"] = items
    
    return data

def extract_F29(doc: Document) -> dict:
    """F/29 Employee Training & Competence Record Sheet."""
    data = {}
    if not doc.tables:
        return data
    t = doc.tables[0]
    
    for row in t.rows:
        cells = all_cells(row.cells)
        for cell in cells:
            text = cell.strip()
            label, value = extract_label_value_arrow(text)
            ll = label.lower()
            if "sr. no" in ll and value:
                data["serial"] = value
            elif "prepared" in ll and "by" in ll:
                data["prepared_by"] = value
            elif "authorised" in ll and "by" in ll:
                data["authorised_by"] = value
            elif "assessed" in ll and "by" in ll:
                data["assessed_by"] = value
            elif "assessed" in ll and "on" in ll:
                data["assessed_on"] = value
    
    return data

def extract_F30(doc: Document) -> dict:
    """F/30 Performance Appraisal Report — complex form with eval scores."""
    data = {}
    if not doc.tables:
        return data
    t = doc.tables[0]
    
    # Extract header fields
    for row in t.rows:
        cells = all_cells(row.cells)
        for cell in cells:
            text = cell.strip()
            label, value = extract_label_value_arrow(text)
            ll = label.lower()
            if "sr. no" in ll and value:
                data["serial"] = value
            elif "date" in ll and value and "date" not in data:
                data["date"] = value
            elif "employee" in ll and "name" in ll:
                data["employee_name"] = value
            elif "department" in ll and value and "department" not in data:
                data["department"] = value
            elif "designation" in ll and value:
                data["designation"] = value
            elif "working" in ll and "month" in ll:
                data["working_months"] = value
            elif "evaluated" in ll and "by" in ll:
                data["evaluated_by"] = value
            elif "evaluator" in ll and "name" in ll and "2" not in ll and "3" not in ll:
                data["evaluator_name"] = value
            elif "evaluator_name2" in ll:
                data["evaluator_name2"] = value
            elif "evaluator_name3" in ll:
                data["evaluator_name3"] = value
            elif "total" in ll and "mark" in ll:
                data["total_marking"] = value
            elif "training" in ll and "need" in ll:
                data["training_need"] = value
            elif "promotion" in ll:
                data["promotion"] = value
            elif "increment" in ll and "last" not in ll:
                data["increment"] = value
            elif "last" in ll and "increment" in ll:
                data["last_increment"] = value
            elif "responsibility" in ll:
                data["responsibility"] = value
            elif "suggestion" in ll:
                data["suggestions"] = value
            elif "authorities" in ll:
                data["authorities"] = value
    
    # Extract evaluation scores: eval_{section_num}_{item_idx}
    # The DOCX table has evaluation rows with score columns (1-4)
    # We need to map them to eval_1_0, eval_1_1, etc.
    # Sections are defined in the template (EVAL_SECTIONS)
    # For now, extract all radio-button-like cells with scores
    
    return data

def extract_F32(doc: Document) -> dict:
    """F/32 Research And Development Request Report."""
    data = {}
    if not doc.tables:
        return data
    t = doc.tables[0]
    
    for row in t.rows:
        cells = all_cells(row.cells)
        for cell in cells:
            text = cell.strip()
            label, value = extract_label_value_arrow(text)
            ll = label.lower()
            if "sr. no" in ll and value:
                data["serial"] = value
            elif "date" in ll and value:
                data["date"] = value
    
    # Extract KV pairs
    for row in t.rows:
        cells = all_cells(row.cells)
        if len(cells) >= 2:
            label = cells[0].strip()
            value = cells[1].strip()
            if label and value and "sr" not in label.lower():
                data[label.lower().replace(" ", "_")] = value
    
    return data

def extract_F34(doc: Document) -> dict:
    """F/34 Design Verification Report."""
    data = {}
    if not doc.tables:
        return data
    t = doc.tables[0]
    
    for row in t.rows:
        cells = all_cells(row.cells)
        for cell in cells:
            text = cell.strip()
            label, value = extract_label_value_arrow(text)
            if "sr. no" in label.lower() and value:
                data["serial"] = value
            elif "date" in label.lower() and value:
                data["date"] = value
    
    # Extract all KV pairs
    for row in t.rows:
        cells = all_cells(row.cells)
        if len(cells) >= 2:
            label = cells[0].strip()
            value = cells[1].strip()
            if label and value and "sr" not in label.lower():
                data[label.lower().replace(" ", "_")] = value
    
    return data

def extract_F35(doc: Document) -> dict:
    """F/35 Design and Development Monitoring Register — 11-column table.
    Expected items: {productName, specification, newSpecification, customerName, reasonOfDevelopment, startDate, targetDate, progress, remarks}
    DOCX columns: ProductName | Specification | NewSpecification | Customer | Reason | DevCompletion | ActualCompletion | ReasonRejection | ActionTaken | Status | DesignHeadSign
    """
    data = {}
    if not doc.tables:
        return data
    t = doc.tables[0]
    
    # Extract serial from header
    for row in t.rows[:2]:
        cells = all_cells(row.cells)
        for cell in cells:
            if "f/35" in cell.lower() and "rev" in cell.lower():
                m = re.search(r'F/35[-_]?\d*', cell)
                if m:
                    data["serial"] = m.group(0)
    
    # Find header row
    header_row_idx = -1
    for ri, row in enumerate(t.rows):
        cells_text = [c.text.strip().lower() for c in row.cells]
        if "product name" in cells_text and "specification" in cells_text:
            header_row_idx = ri
            break
    
    if header_row_idx < 0:
        return data
    
    # Column mapping (11 columns)
    COL_MAP = {
        0: "productName",
        1: "specification",
        2: "newSpecification",
        3: "customerName",
        4: "reasonOfDevelopment",
        5: "startDate",      # Development Completion Date
        6: "targetDate",     # Actual Completion Date
        7: "progress",       # Reason for Rejection
        8: "remarks",        # Action Taken
    }
    
    items = []
    for row in t.rows[header_row_idx + 1:]:
        cells = all_cells(row.cells)
        if not any(c.strip() for c in cells):
            continue
        
        item = {}
        for ci in range(min(len(cells), 11)):
            field = COL_MAP.get(ci)
            if field and cells[ci].strip():
                if field not in item:
                    item[field] = cells[ci].strip()
        
        if item.get("productName"):
            items.append(item)
    
    if items:
        data["items"] = items
    
    return data

def extract_F37(doc: Document) -> dict:
    """F/37 Experiment Data Sheet."""
    data = {}
    if not doc.tables:
        return data
    t = doc.tables[0]
    
    for row in t.rows:
        cells = all_cells(row.cells)
        for cell in cells:
            text = cell.strip()
            label, value = extract_label_value_arrow(text)
            if "sr. no" in label.lower() and value:
                data["serial"] = value
            elif "date" in label.lower() and value:
                data["date"] = value
    
    # Extract all KV pairs
    for row in t.rows:
        cells = all_cells(row.cells)
        if len(cells) >= 2:
            label = cells[0].strip()
            value = cells[1].strip()
            if label and value and "sr" not in label.lower():
                data[label.lower().replace(" ", "_")] = value
    
    return data

def extract_F40(doc: Document) -> dict:
    """F/40 Competence Matrix.
    Expected items: {srNo, designation, qualReq, qualAvail, expReq, expAvail, skillReq, skillAvail, training1-5}
    """
    data = {}
    if not doc.tables:
        return data
    t = doc.tables[0]
    
    # Extract serial from header (F/40 Rev No.)
    for row in t.rows[:3]:
        cells = all_cells(row.cells)
        for cell in cells:
            text = cell.strip()
            if "f/40" in text.lower() and "rev" in text.lower():
                m = re.search(r'F/40[-_]?\d*', text)
                if m:
                    data["serial"] = m.group(0)
            elif "reviewed" in text.lower() and "by" in text.lower():
                # Find the value (usually in a different cell)
                for c2 in cells:
                    c2t = c2.strip()
                    if c2t and "reviewed" not in c2t.lower() and "f/40" not in c2t.lower():
                        if "reviewed_by" not in data:
                            data["reviewed_by"] = c2t
                        else:
                            data["reviewed_on"] = c2t
    
    # Find header row (the one with "Sr. No" + "Designation")
    header_row_idx = -1
    for ri, row in enumerate(t.rows):
        cells_text = [c.text.strip().lower() for c in row.cells]
        if cells_text[0].startswith("sr") and "designation" in " ".join(cells_text):
            header_row_idx = ri
            break
    
    if header_row_idx < 0:
        return data
    
    # Column mapping based on DOCX structure:
    # 0=Sr.No, 1-3=Designation(merged), 4-5=QualReq(merged), 6-8=ExpReq(merged),
    # 9-12=SkillReq(merged), 13-31=Training columns (many)
    # But the actual data rows have specific positions. Use positional mapping:
    COL_MAP = {
        0: "srNo",
        1: "designation",  # merged 1-3
        4: "qualReq",     # merged 4-5
        6: "expReq",      # merged 6-8
        9: "skillReq",    # merged 9-12
        # Training columns start at 13
        13: "training1",
        14: "training2",
        15: "training3",
        16: "training4",
        17: "training5",
    }
    
    items = []
    for row in t.rows[header_row_idx + 1:]:
        cells = all_cells(row.cells)
        if not any(c.strip() for c in cells):
            continue
        
        item = {}
        for ci in range(min(len(cells), 20)):
            field = COL_MAP.get(ci)
            if field and cells[ci].strip():
                if field not in item:
                    item[field] = cells[ci].strip()
        
        if item.get("srNo") or item.get("designation"):
            items.append(item)
    
    if items:
        data["items"] = items
    
    return data

def extract_F41(doc: Document) -> dict:
    """F/41 Competence Gap Analysis."""
    data = {}
    if not doc.tables:
        return data
    t = doc.tables[0]
    
    for row in t.rows:
        cells = all_cells(row.cells)
        for cell in cells:
            text = cell.strip()
            label, value = extract_label_value_arrow(text)
            if "sr. no" in label.lower() and value:
                data["serial"] = value
            elif "date" in label.lower() and value:
                data["date"] = value
    
    # Extract gap analysis items
    items = []
    header_row_idx = -1
    for ri, row in enumerate(t.rows):
        cells_text = [c.text.strip().lower() for c in row.cells]
        joined = " ".join(cells_text)
        if "gap" in joined or ("competence" in joined and "need" in joined):
            header_row_idx = ri
            break
    
    if header_row_idx >= 0:
        for row in t.rows[header_row_idx + 1:]:
            cells = all_cells(row.cells)
            if not any(c.strip() for c in cells):
                continue
            unique = dedup_cells(row.cells)
            if len(unique) >= 2:
                items.append({
                    "name": unique[0],
                    "currentCompetence": unique[1] if len(unique) > 1 else "",
                    "gap": unique[2] if len(unique) > 2 else "",
                    "trainingNeed": unique[3] if len(unique) > 3 else "",
                })
    
    if items:
        data["items"] = items
    
    return data

def extract_F42(doc: Document) -> dict:
    """F/42 Annual Training Plan."""
    data = {}
    if not doc.tables:
        return data
    t = doc.tables[0]
    
    for row in t.rows:
        cells = all_cells(row.cells)
        for cell in cells:
            text = cell.strip()
            label, value = extract_label_value_arrow(text)
            if "sr. no" in label.lower() and value:
                data["serial"] = value
            elif "date" in label.lower() and value:
                data["date"] = value
    
    # Extract training plan items
    items = []
    header_row_idx = -1
    for ri, row in enumerate(t.rows):
        cells_text = [c.text.strip().lower() for c in row.cells]
        joined = " ".join(cells_text)
        if "training" in joined and ("month" in joined or "quarter" in joined or "topic" in joined):
            header_row_idx = ri
            break
    
    if header_row_idx >= 0:
        for row in t.rows[header_row_idx + 1:]:
            cells = all_cells(row.cells)
            if not any(c.strip() for c in cells):
                continue
            unique = dedup_cells(row.cells)
            if len(unique) >= 2:
                items.append({
                    "month": unique[0],
                    "topic": unique[1] if len(unique) > 1 else "",
                    "targetAudience": unique[2] if len(unique) > 2 else "",
                    "trainer": unique[3] if len(unique) > 3 else "",
                })
    
    if items:
        data["items"] = items
    
    return data

def extract_F43(doc: Document) -> dict:
    """F/43 Induction Training Record."""
    data = {}
    if not doc.tables:
        return data
    t = doc.tables[0]
    
    for row in t.rows:
        cells = all_cells(row.cells)
        for cell in cells:
            text = cell.strip()
            label, value = extract_label_value_arrow(text)
            ll = label.lower()
            if "sr. no" in ll and value:
                data["serial"] = value
            elif "date" in ll and value and "date" not in data:
                data["date"] = value
            elif "employee" in ll and "name" in ll:
                data["employee_name"] = value
            elif "department" in ll and value and "department" not in data:
                data["department"] = value
            elif "designation" in ll and value:
                data["designation"] = value
            elif "date" in ll and "joining" in ll:
                data["date_of_joining"] = value
            elif "qualification" in ll:
                data["qualification"] = value
            elif "topics" in ll and value:
                data["topics"] = value
            elif "effectiveness" in ll and value:
                data["effectiveness"] = value
            elif "issued" in ll and "by" in ll:
                data["issued_by"] = value
            elif "authorised" in ll and "sign" in ll:
                data["authorised_sign"] = value
            elif "inductee" in ll and "sign" in ll:
                data["inductee_sign"] = value
            elif "trainer" in ll and "sign" in ll:
                data["trainer_sign"] = value
            elif "sign" in ll and "date" in ll:
                data["sign_date"] = value
    
    return data

def extract_F44(doc: Document) -> dict:
    """F/44 Job Description — paragraph-based, no tables.
    Expected: serial, position, reports_to, approved_by, delegation, responsibilities
    """
    data = {}
    
    # Extract serial from first paragraph
    for p in doc.paragraphs[:5]:
        if "F/44" in p.text:
            m = re.search(r'F/44[-_]?\d*', p.text)
            if m:
                data["serial"] = m.group(0)
            break
    
    # Extract fields from paragraphs
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        ll = text.lower()
        if "job title" in ll and ":" in text:
            data["position"] = text.split(":", 1)[1].strip()
        elif "reports to" in ll and ":" in text:
            data["reports_to"] = text.split(":", 1)[1].strip()
        elif "approved by" in ll and ":" in text:
            data["approved_by"] = text.split(":", 1)[1].strip()
        elif "delegation" in ll and ":" in text:
            data["delegation"] = text.split(":", 1)[1].strip()
        elif "responsibilities" in ll.lower():
            # Collect all paragraphs after this as responsibilities
            pass
    
    # Collect responsibilities (paragraphs under "Key Responsibilities" section)
    resp_lines = []
    in_resp = False
    for p in doc.paragraphs:
        text = p.text.strip()
        if "key responsibilities" in text.lower():
            in_resp = True
            continue
        if in_resp:
            if text and not text.isupper() and len(text) > 10:
                resp_lines.append(text)
            # Stop at next major section
            if text.isupper() and "RESPONS" not in text.upper():
                break
    
    if resp_lines:
        data["responsibilities"] = "\n".join(resp_lines)
    
    return data

def extract_F45(doc: Document) -> dict:
    """F/45 Master List of Documents."""
    data = {}
    if not doc.tables:
        return data
    t = doc.tables[0]
    
    for row in t.rows[:2]:
        cells = all_cells(row.cells)
        for cell in cells:
            _, value = extract_label_value_arrow(cell.strip())
            if "f/45" in value.lower():
                data["serial"] = value
    
    # Header: Doc No. | Title | Rev No. | Date | Copy No. | Location
    items = []
    header_row_idx = -1
    for ri, row in enumerate(t.rows):
        cells_text = [c.text.strip().lower() for c in row.cells]
        joined = " ".join(cells_text)
        if "doc" in joined and "title" in joined and "rev" in joined:
            header_row_idx = ri
            break
    
    if header_row_idx >= 0:
        for row in t.rows[header_row_idx + 1:]:
            cells = all_cells(row.cells)
            if not any(c.strip() for c in cells):
                continue
            unique = dedup_cells(row.cells)
            if len(unique) >= 2:
                items.append({
                    "docNo": unique[0],
                    "title": unique[1] if len(unique) > 1 else "",
                    "revNo": unique[2] if len(unique) > 2 else "",
                    "date": unique[3] if len(unique) > 3 else "",
                    "copyNo": unique[4] if len(unique) > 4 else "",
                    "location": unique[5] if len(unique) > 5 else "",
                })
    
    if items:
        data["items"] = items
    
    return data

def extract_F46(doc: Document) -> dict:
    """F/46 Management of Change Plan."""
    data = {}
    if not doc.tables:
        return data
    t = doc.tables[0]
    
    for row in t.rows:
        cells = all_cells(row.cells)
        for cell in cells:
            text = cell.strip()
            label, value = extract_label_value_arrow(text)
            ll = label.lower()
            if "sr. no" in ll and value:
                data["serial"] = value
            elif "date" in ll and value:
                data["date"] = value
            elif "change" in ll and "description" in ll:
                data["change_description"] = value
            elif "reason" in ll:
                data["reason"] = value
            elif "approved" in ll and "by" in ll:
                data["approved_by"] = value
    
    # Extract KV pairs
    for row in t.rows:
        cells = all_cells(row.cells)
        if len(cells) >= 2:
            label = cells[0].strip()
            value = cells[1].strip()
            if label and value and "sr" not in label.lower():
                data[label.lower().replace(" ", "_")] = value
    
    return data

def extract_F47(doc: Document) -> dict:
    """F/47 Internal Audit Checklist."""
    data = {}
    if not doc.tables:
        return data
    t = doc.tables[0]
    
    for row in t.rows:
        cells = all_cells(row.cells)
        for cell in cells:
            text = cell.strip()
            label, value = extract_label_value_arrow(text)
            ll = label.lower()
            if "sr. no" in ll and value:
                data["serial"] = value
            elif "date" in ll and value:
                data["date"] = value
            elif "auditor" in ll and value:
                data["auditor"] = value
            elif "department" in ll and value:
                data["department"] = value
    
    # Extract checklist items
    items = []
    header_row_idx = -1
    for ri, row in enumerate(t.rows):
        cells_text = [c.text.strip().lower() for c in row.cells]
        joined = " ".join(cells_text)
        if "clause" in joined or "requirement" in joined or "audit" in joined and "find" in joined:
            header_row_idx = ri
            break
    
    if header_row_idx >= 0:
        for row in t.rows[header_row_idx + 1:]:
            cells = all_cells(row.cells)
            if not any(c.strip() for c in cells):
                continue
            unique = dedup_cells(row.cells)
            if len(unique) >= 2:
                items.append({
                    "clause": unique[0],
                    "requirement": unique[1] if len(unique) > 1 else "",
                    "finding": unique[2] if len(unique) > 2 else "",
                    "status": unique[3] if len(unique) > 3 else "",
                })
    
    if items:
        data["items"] = items
    
    return data

def extract_F48(doc: Document) -> dict:
    """F/48 Internal Audit Report — table with multi-line cells.
    Expected: serial, audit_type, date_of_audit, audit_team, audit_standard, audit_location, audit_scope, auditee, summary_report, audit_findings, followup_required, followup_date, auditor_signature
    """
    data = {}
    if not doc.tables:
        return data
    t = doc.tables[0]
    
    for row in t.rows:
        cells = all_cells(row.cells)
        for cell in cells:
            text = cell.strip()
            if not text:
                continue
            ll = text.lower()
            
            if "audit report no" in ll:
                m = re.search(r'F/48[-_]?\d*', text)
                if m:
                    data["serial"] = m.group(0)
            elif "type of audit" in ll:
                # Extract value after label
                for sep in [":", "\n"]:
                    if sep in text:
                        val = text.split(sep, 1)[1].strip()
                        if val and val not in data.get("audit_type", ""):
                            data["audit_type"] = val
                            break
            elif "date of audit" in ll:
                for sep in [":", "\n"]:
                    if sep in text:
                        val = text.split(sep, 1)[1].strip()
                        if val:
                            data["date_of_audit"] = val
                            break
            elif "audit team" in ll:
                for sep in [":", "\n"]:
                    if sep in text:
                        val = text.split(sep, 1)[1].strip()
                        if val:
                            data["audit_team"] = val
                            break
            elif "audit standard" in ll:
                for sep in [":", "\n"]:
                    if sep in text:
                        val = text.split(sep, 1)[1].strip()
                        if val:
                            data["audit_standard"] = val
                            break
            elif "audit location" in ll:
                for sep in [":", "\n"]:
                    if sep in text:
                        val = text.split(sep, 1)[1].strip()
                        if val:
                            data["audit_location"] = val
                            break
            elif "audit scope" in ll:
                for sep in [":", "\n"]:
                    if sep in text:
                        val = text.split(sep, 1)[1].strip()
                        if val:
                            data["audit_scope"] = val
                            break
            elif "auditee" in ll:
                for sep in [":", "\n"]:
                    if sep in text:
                        val = text.split(sep, 1)[1].strip()
                        if val:
                            data["auditee"] = val
                            break
            elif "summary report" in ll:
                for sep in [":", "\n"]:
                    if sep in text:
                        val = text.split(sep, 1)[1].strip()
                        if val:
                            data["summary_report"] = val
                            break
            elif "audit findings" in ll:
                for sep in [":", "\n"]:
                    if sep in text:
                        val = text.split(sep, 1)[1].strip()
                        if val:
                            data["audit_findings"] = val
                            break
            elif "follow-up audit required" in ll or "followup" in ll:
                if "yes" in ll:
                    data["followup_required"] = "Yes"
                elif "no" in ll:
                    data["followup_required"] = "No"
                # Extract date if present
                m = re.search(r'\d{2}/\d{2}/\d{4}', text)
                if m:
                    data["followup_date"] = m.group(0)
            elif "signature" in ll:
                m = re.search(r'\(([^)]+)\)', text)
                if m:
                    data["auditor_signature"] = m.group(1).strip()
    
    return data

def extract_F50(doc: Document) -> dict:
    """F/50 Customer Property Monitoring Register."""
    data = {}
    if not doc.tables:
        return data
    t = doc.tables[0]
    
    for row in t.rows[:2]:
        cells = all_cells(row.cells)
        for cell in cells:
            _, value = extract_label_value_arrow(cell.strip())
            if "f/50" in value.lower():
                data["serial"] = value
    
    # Extract items
    items = []
    header_row_idx = -1
    for ri, row in enumerate(t.rows):
        cells_text = [c.text.strip().lower() for c in row.cells]
        joined = " ".join(cells_text)
        if "date" in joined and "customer" in joined and "property" in joined:
            header_row_idx = ri
            break
    
    if header_row_idx >= 0:
        for row in t.rows[header_row_idx + 1:]:
            cells = all_cells(row.cells)
            if not any(c.strip() for c in cells):
                continue
            unique = dedup_cells(row.cells)
            if len(unique) >= 2:
                items.append({
                    "date": unique[0],
                    "customer": unique[1] if len(unique) > 1 else "",
                    "property": unique[2] if len(unique) > 2 else "",
                    "condition": unique[3] if len(unique) > 3 else "",
                })
    
    if items:
        data["items"] = items
    
    return data

# ── Extractor registry ──────────────────────────────────────────
EXTRACTORS = {
    "F/08": extract_F08,
    "F/09": extract_F09,
    "F/10": extract_F10,
    "F/11": extract_F11,
    "F/12": extract_F12,
    "F/13": extract_F13,
    "F/14": extract_F14,
    "F/15": extract_F15,
    "F/16": extract_F16,
    "F/17": extract_F17,
    "F/18": extract_F18,
    "F/19": extract_F19,
    "F/20": extract_F20,
    "F/21": extract_F21,
    "F/22": extract_F22,
    "F/23": extract_F23,
    "F/24": extract_F24,
    "F/25": extract_F25,
    "F/28": extract_F28,
    "F/29": extract_F29,
    "F/30": extract_F30,
    "F/32": extract_F32,
    "F/34": extract_F34,
    "F/35": extract_F35,
    "F/37": extract_F37,
    "F/40": extract_F40,
    "F/41": extract_F41,
    "F/42": extract_F42,
    "F/43": extract_F43,
    "F/44": extract_F44,
    "F/45": extract_F45,
    "F/46": extract_F46,
    "F/47": extract_F47,
    "F/48": extract_F48,
    "F/50": extract_F50,
}

# ── SQL generation ───────────────────────────────────────────────
def json_to_sql_safe(obj: dict) -> str:
    """Convert dict to JSON string safe for SQL."""
    return json.dumps(obj, ensure_ascii=False, indent=None, separators=(',', ':'))

def generate_update_sql(serial: str, form_data: dict) -> str:
    """Generate a single UPDATE statement."""
    # Escape single quotes in the JSON string
    json_str = json_to_sql_safe(form_data)
    escaped = json_str.replace("'", "''")
    return f"UPDATE public.records SET form_data = '{escaped}'::jsonb WHERE serial = '{serial}';"

# ── Main ETL pipeline ────────────────────────────────────────────
def main():
    dry_run = "--dry-run" in sys.argv
    form_filter = None
    if "--form" in sys.argv:
        idx = sys.argv.index("--form")
        if idx + 1 < len(sys.argv):
            form_filter = sys.argv[idx + 1]
    
    print("=" * 60)
    print("QBase DOCX ETL Backfill")
    print("=" * 60)
    
    # Scan files
    all_files = []
    for root, dirs, files in os.walk(BASE_DIR):
        for fn in files:
            serial = extract_serial(fn)
            if serial:
                form_code = serial.split("-")[0]
                if form_filter and form_code != form_filter:
                    continue
                all_files.append({
                    'path': os.path.join(root, fn),
                    'serial': serial,
                    'form_code': form_code,
                })
    
    all_files.sort(key=lambda x: (x['form_code'], x['serial']))
    print(f"Found {len(all_files)} DOCX/DOC files to process")
    
    # Process each file
    updates = []
    errors = []
    stats = {}
    
    for fi, file_info in enumerate(all_files):
        serial = file_info['serial']
        form_code = file_info['form_code']
        filepath = file_info['path']
        
        if form_code not in EXTRACTORS:
            errors.append(f"{serial}: No extractor for {form_code}")
            continue
        
        stats[form_code] = stats.get(form_code, 0) + 1
        
        # Open document
        doc = open_document(filepath)
        if not doc:
            errors.append(f"{serial}: Could not open {filepath}")
            continue
        
        # Extract data
        try:
            data = EXTRACTORS[form_code](doc)
            if not data:
                errors.append(f"{serial}: No data extracted")
                continue
            # Ensure serial is set
            if "serial" not in data:
                data["serial"] = serial
            
            updates.append({
                'serial': serial,
                'form_code': form_code,
                'data': data,
            })
            
            if (fi + 1) % 20 == 0:
                print(f"  Processed {fi + 1}/{len(all_files)}...")
                
        except Exception as e:
            errors.append(f"{serial}: {e}")
    
    # Generate SQL
    print(f"\n{'=' * 60}")
    print(f"Extraction complete: {len(updates)} records extracted, {len(errors)} errors")
    print(f"\nErrors:")
    for e in errors[:20]:
        print(f"  {e}")
    if len(errors) > 20:
        print(f"  ... and {len(errors) - 20} more")
    
    print(f"\nStats by form code:")
    for fc in sorted(stats.keys()):
        count = sum(1 for u in updates if u['form_code'] == fc)
        print(f"  {fc}: {count}/{stats[fc]} extracted")
    
    if dry_run:
        print("\n--dry-run: skipping SQL file generation")
        # Print sample for F/12
        for u in updates:
            if u['form_code'] == 'F/12':
                print(f"\nSample F/12 data:")
                print(json.dumps(u['data'], indent=2, ensure_ascii=False)[:2000])
                break
        return
    
    # Generate SQL file
    header = f"""-- ============================================================================
-- QBase — Canonical Records Backfill from Source DOCX Files
-- ============================================================================
-- Generated by scripts/backfill-from-docx.py
-- 
-- This migration UPDATEs form_data with canonical, template-matched JSONB
-- extracted directly from the original Word .docx/.doc source files.
--
-- HOW TO APPLY:
--   1. Log into Supabase Dashboard → SQL Editor
--   2. Paste this ENTIRE file → Run
--   3. All updates are matched by record serial
--
-- SAFETY: All updates preserve serial integrity. No INSERTs, no DELETEs.
-- ============================================================================
"""
    
    sql_lines = [header]
    sql_lines.append(f"-- Total records to update: {len(updates)}")
    sql_lines.append(f"-- Generated from {len(all_files)} source files")
    sql_lines.append("")
    sql_lines.append("BEGIN;")
    sql_lines.append("")
    
    # Group by form code
    current_form = None
    for u in sorted(updates, key=lambda x: (x['form_code'], x['serial'])):
        if u['form_code'] != current_form:
            current_form = u['form_code']
            sql_lines.append(f"-- ── {current_form} ──────────────────────────────────────────────")
        
        sql = generate_update_sql(u['serial'], u['data'])
        sql_lines.append(sql)
    
    sql_lines.append("")
    sql_lines.append("COMMIT;")
    sql_lines.append("")
    sql_lines.append("-- ══════════════════════════════════════════════════════════════════")
    sql_lines.append("-- VERIFICATION: Check record counts after update")
    sql_lines.append("-- ══════════════════════════════════════════════════════════════════")
    sql_lines.append(f"-- SELECT form_code, COUNT(*) FROM public.records GROUP BY form_code ORDER BY form_code;")
    
    sql_content = "\n".join(sql_lines)
    
    # Write to output file
    OUTPUT_FILE.parent.mkdir(parents=True, exist_ok=True)
    with open(OUTPUT_FILE, "w") as f:
        f.write(sql_content)
    
    print(f"\nSQL file written: {OUTPUT_FILE}")
    print(f"Size: {len(sql_content)} bytes")
    
    # Also write errors to a log file
    if errors:
        with open("/tmp/backfill_errors.log", "w") as f:
            f.write("\n".join(errors))
        print(f"Errors logged: /tmp/backfill_errors.log")

if __name__ == "__main__":
    main()