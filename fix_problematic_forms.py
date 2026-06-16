#!/usr/bin/env python3
"""
Quick fix for problematic forms: F/23, F/28, F/29, F/35, F/45, F/50
Updates their QBase records with properly extracted data.
"""
import json, os, urllib.request, re
from docx import Document
from pathlib import Path

RECORDS_DIR = Path('/home/Kepv/Downloads/03.Records')

with open('/mnt/ahmed/Projects/qbase/.env.local') as f:
    env = {}
    for line in f:
        line = line.strip()
        if line and not line.startswith('#') and '=' in line:
            k,v=line.split('=',1)
            env[k]=v

SUPABASE_URL = env['SUPABASE_URL']
SERVICE_KEY = env['SUPABASE_SERVICE_ROLE_KEY']

def update_record(serial, form_data):
    url = f"{SUPABASE_URL}/rest/v1/records?serial=eq.{serial}"
    req = urllib.request.Request(url, data=json.dumps({'form_data': form_data}).encode(), method='PATCH')
    req.add_header('apikey', SERVICE_KEY)
    req.add_header('Authorization', f'Bearer {SERVICE_KEY}')
    req.add_header('Content-Type', 'application/json')
    urllib.request.urlopen(req, timeout=10)

def find_file(form_code, serial):
    """Find DOCX file by serial."""
    for root, dirs, files in os.walk(RECORDS_DIR):
        for f in files:
            if f.endswith(('.docx', '.doc')) and serial.replace('/', '').replace('-', '') in f.replace('-', '').replace('/', '').replace('_', ''):
                if serial.split('-')[1] in f:
                    return os.path.join(root, f)
    return None

def clean_cells(row):
    cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
    clean = []
    for c in cells:
        if clean and clean[-1] == c: continue
        clean.append(c)
    return clean

# =====================================================================
# F/45 Master List of Documents
# =====================================================================
print("=== Fixing F/45 ===")
filepath = find_file('F/45', 'F/45-001')
if filepath:
    doc = Document(filepath)
    table = doc.tables[0]
    docs = []
    headers = None
    for ri, row in enumerate(table.rows):
        cells = clean_cells(row)
        if ri == 3:  # Header row
            headers = [c.lower().strip() for c in cells]
            continue
        if ri < 5: continue  # Skip meta rows
        if len(cells) >= 5:
            d = {}
            if len(cells) > 1: d['sr'] = cells[0]
            if len(cells) > 2: d['title'] = cells[1]
            if len(cells) > 3: d['document_number'] = cells[2]
            if len(cells) > 5: d['current_revision'] = cells[5]
            if len(cells) > 6: d['revision_date'] = cells[6]
            if len(cells) > 7: d['revision_details'] = cells[7]
            if d.get('title') and d['title'] != 'Title of Document' and d['title'] != 'Sr.':
                docs.append(d)

    form_data = {'documents': json.dumps(docs, ensure_ascii=False)} if docs else {}
    form_data['last_updated_by'] = 'Ahmed Khaled'
    form_data['update_date'] = '01/02/2026'
    update_record('F/45-001', form_data)
    print(f"  F/45-001: {len(docs)} documents extracted")

# =====================================================================
# F/50 Customer Property
# =====================================================================
print("=== Fixing F/50 ===")
for serial in ['F/50-001']:
    filepath = find_file('F/50', serial)
    if not filepath: continue
    doc = Document(filepath)
    table = doc.tables[0]
    props = []
    for ri, row in enumerate(table.rows):
        cells = clean_cells(row)
        if ri == 0: continue
        if ri == 1: continue  # Header
        if len(cells) >= 6:
            prop = {}
            if len(cells) > 0: prop['received_date'] = cells[0]
            if len(cells) > 1: prop['name_of_property'] = cells[1]
            if len(cells) > 2: prop['purpose'] = cells[2]
            if len(cells) > 3: prop['qty'] = cells[3]
            if len(cells) > 4: prop['customer'] = cells[4]
            if len(cells) > 5: prop['inspection_status'] = cells[5]
            if prop.get('name_of_property') and prop['name_of_property'] != 'Name Of Property':
                props.append(prop)

    form_data = {'properties': json.dumps(props, ensure_ascii=False)} if props else {}
    update_record(serial, form_data)
    print(f"  {serial}: {len(props)} properties extracted")

# =====================================================================
# F/35 Design Monitoring
# =====================================================================
print("=== Fixing F/35 ===")
for serial in ['F/35-001', 'F/35-002', 'F/35-003', 'F/35-004']:
    filepath = find_file('F/35', serial)
    if not filepath: continue
    doc = Document(filepath)
    table = doc.tables[0]
    projects = []
    for ri, row in enumerate(table.rows):
        cells = clean_cells(row)
        if ri == 0: continue
        if ri == 1: continue  # Header
        if len(cells) >= 3:
            proj = {}
            if len(cells) > 0: proj['product_name'] = cells[0]
            if len(cells) > 1: proj['specification'] = cells[1]
            if len(cells) > 2: proj['new_specification'] = cells[2]
            if len(cells) > 3: proj['customer'] = cells[3]
            if len(cells) > 4: proj['reason'] = cells[4]
            if len(cells) > 5: proj['completion_date'] = cells[5]
            if len(cells) > 6: proj['actual_completion'] = cells[6]
            if proj.get('product_name') and proj['product_name'] != 'Product Name':
                projects.append(proj)

    form_data = {'projects': json.dumps(projects, ensure_ascii=False)} if projects else {}
    update_record(serial, form_data)
    print(f"  {serial}: {len(projects)} projects extracted")

# =====================================================================
# F/28 Training Attendance
# =====================================================================
print("=== Fixing F/28 ===")
for root, dirs, files in os.walk(RECORDS_DIR / '05 HR & Training Records' / 'F-28 - Training Attendance'):
    for f in files:
        if not f.endswith(('.docx', '.doc')): continue
        m = re.search(r'F[/\-_]28[-\_](\d+)', f)
        if not m: continue
        serial = f"F/28-{int(m.group(1)):03d}"
        filepath = os.path.join(root, f)
        doc = Document(filepath)
        table = doc.tables[0]
        participants = []
        for ri, row in enumerate(table.rows):
            cells = clean_cells(row)
            if ri == 0: continue
            if len(cells) >= 4:
                p = {}
                if len(cells) > 0: p['sl_no'] = cells[0]
                if len(cells) > 1: p['name'] = cells[1]
                if len(cells) > 2: p['department'] = cells[2]
                if len(cells) > 3: p['id_no'] = cells[3]
                if len(cells) > 4: p['training_date'] = cells[4]
                if len(cells) > 5: p['signature'] = cells[5]
                if p.get('name') and p['name'] != 'Name Of The Participant' and p['name'] != 'Name':
                    participants.append(p)

        form_data = {'participants': json.dumps(participants, ensure_ascii=False)} if participants else {}
        update_record(serial, form_data)
        print(f"  {serial}: {len(participants)} participants")

# =====================================================================
# F/29 Training Record
# =====================================================================
print("=== Fixing F/29 ===")
for root, dirs, files in os.walk(RECORDS_DIR / '05 HR & Training Records' / 'F-29 - Training Record'):
    for f in files:
        if not f.endswith(('.docx', '.doc')): continue
        m = re.search(r'F[/\-_]29[-\_](\d+)', f)
        if not m: continue
        serial = f"F/29-{int(m.group(1)):03d}"
        filepath = os.path.join(root, f)
        doc = Document(filepath)
        table = doc.tables[0]
        employees = []

        # Header detection
        header_idx = None
        for ri, row in enumerate(table.rows):
            cells = clean_cells(row)
            if any('Name Of Employee' in c or 'Qualification' in c for c in cells):
                header_idx = ri
                break

        if header_idx is not None:
            for ri in range(header_idx + 1, len(table.rows)):
                row = table.rows[ri]
                cells = clean_cells(row)
                if len(cells) >= 5:
                    emp = {}
                    if len(cells) > 1: emp['name'] = cells[1]
                    if len(cells) > 2: emp['qualification_req'] = cells[2]
                    if len(cells) > 3: emp['qualification_avail'] = cells[3]
                    if len(cells) > 4: emp['experience_req'] = cells[4]
                    if len(cells) > 5: emp['experience_avail'] = cells[5]
                    if emp.get('name') and 'Employee' not in emp['name']:
                        employees.append(emp)

        form_data = {'employees': json.dumps(employees, ensure_ascii=False)} if employees else {}
        update_record(serial, form_data)
        print(f"  {serial}: {len(employees)} employees")

# =====================================================================
# F/23 Master List of Records
# =====================================================================
print("=== Fixing F/23 ===")
filepath = find_file('F/23', 'F/23-001')
if filepath:
    doc = Document(filepath)
    table = doc.tables[0]
    records = []
    headers = None
    for ri, row in enumerate(table.rows):
        cells = clean_cells(row)
        if ri == 0:
            headers = [c.lower().strip() for c in cells]
            continue
        if len(cells) >= 3:
            rec = {}
            for ci, h in enumerate(headers or []):
                if ci < len(cells):
                    rec[h] = cells[ci]
            if rec.get('record no') and rec['record no'] != 'Record No.' and rec['record no'].startswith('F/'):
                records.append(rec)

    form_data = {'records_list': json.dumps(records, ensure_ascii=False)} if records else {}
    update_record('F/23-001', form_data)
    print(f"  F/23-001: {len(records)} records extracted")

print("\n✅ All problematic forms fixed!")
import re