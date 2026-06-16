#!/usr/bin/env python3
"""
QMS → QBase Migration Script
Downloads all DOCX records from Google Drive, extracts data, and inserts into QBase Supabase.
"""
import json, os, sys, re, urllib.request, urllib.parse
from datetime import datetime, timezone
from docx import Document
from pathlib import Path

# ============================================================================
# CONFIG
# ============================================================================
DRIVE_ROOT = '1zA3ZqbFwxsa75DBl55oXybB2jyxvagYS'
TOKEN_PATH = Path(os.path.expanduser('~/.hermes/google_token.json'))
SAMPLES_DIR = Path('/tmp/qms_samples')
MIGRATION_LOG = Path('/tmp/qms_migration_log.json')

# Load Supabase credentials
def load_env():
    env = {}
    with open('/mnt/ahmed/Projects/qbase/.env.local') as f:
        for line in f:
            line = line.strip()
            if line and not line.startswith('#') and '=' in line:
                k, v = line.split('=', 1)
                env[k] = v
    return env

ENV = load_env()
SUPABASE_URL = ENV['SUPABASE_URL']
SERVICE_KEY = ENV['SUPABASE_SERVICE_ROLE_KEY']

# ============================================================================
# QBASE FORM SCHEMAS (from formSchemas.ts)
# ============================================================================
FORM_SCHEMAS = {
    'F/08': {
        'name': 'Order Form',
        'section': 1,
        'section_name': 'Sales & Customer Service',
        'fields': ['serial', 'date', 'client_name', 'mode_of_receipt', 'items',
                    'test_certificate_required', 'delivery_schedule', 'complies',
                    'order_status', 'remarks', 'reviewed_by', 'bill_no', 'despatch_date']
    },
    'F/09': {
        'name': 'Customer Complaint',
        'section': 1,
        'section_name': 'Sales & Customer Service',
        'fields': ['serial', 'date', 'receipt_date', 'receipt_time', 'received_by',
                    'mode_of_receipt', 'client_name', 'product_type', 'description',
                    'complaint_nature', 'corrective_action', 'result_of_action',
                    'actions_proposed', 'customer_informed_date', 'analysed_by', 'closed_by']
    },
    'F/10': {
        'name': 'Customer Feedback',
        'section': 1,
        'section_name': 'Sales & Customer Service',
        'fields': ['serial', 'date', 'year', 'client_name', 'address',
                    'rating_product_quality', 'rating_order_processing',
                    'rating_complaint_handling', 'rating_delivery', 'rating_price',
                    'project_name', 'comment_text', 'suggestions',
                    'distributor_signature', 'reviewed_by', 'action_proposed',
                    'corrective_action_ref', 'remarks']
    },
    'F/50': {
        'name': 'Customer Property',
        'section': 1,
        'section_name': 'Sales & Customer Service',
        'fields': ['serial', 'entries', 'received_date', 'name_of_property',
                    'purpose_for', 'received_qty', 'name_of_customer',
                    'inward_inspection', 'received_by', 'outward_date',
                    'outward_qty', 'balance_qty', 'damage_summary', 'outward_by']
    },
    'F/11': {
        'name': 'Production Plan',
        'section': 2,
        'section_name': 'Operations & Production',
        'fields': ['serial', 'month', 'year', 'date', 'projects', 'name',
                    'client', 'status', 'notes', 'prepared_by', 'approved_by']
    },
    'F/19': {
        'name': 'Product Description',
        'section': 2,
        'section_name': 'Operations & Production',
        'fields': ['serial', 'date', 'project_name', 'client_name', 'description',
                    'specifications', 'requirements', 'prepared_by', 'approved_by']
    },
    'F/12': {
        'name': 'Non-Conforming',
        'section': 3,
        'section_name': 'Quality & Audit',
        'fields': ['serial', 'date', 'project_name', 'nc_type', 'description',
                    'root_cause', 'corrective_action', 'preventive_action',
                    'status', 'closure_date', 'reported_by']
    },
    'F/17': {
        'name': 'QA Test Request',
        'section': 3,
        'section_name': 'Quality & Audit',
        'fields': ['serial', 'date', 'project_name', 'test_type', 'description',
                    'criteria', 'requested_by']
    },
    'F/18': {
        'name': 'Product Re-Call',
        'section': 3,
        'section_name': 'Quality & Audit',
        'fields': ['serial', 'date', 'product', 'reason', 'affected_items',
                    'resolution', 'authorized_by']
    },
    'F/22': {
        'name': 'Corrective Action',
        'section': 3,
        'section_name': 'Quality & Audit',
        'fields': ['serial', 'date', 'project_name', 'source', 'problem_description',
                    'root_cause', 'corrective_action', 'responsible', 'due_date',
                    'status', 'verification', 'initiated_by']
    },
    'F/25': {
        'name': 'Audit Plan',
        'section': 3,
        'section_name': 'Quality & Audit',
        'fields': ['serial', 'period', 'year', 'date', 'scope', 'audits',
                    'department', 'audit_date', 'auditor', 'status',
                    'prepared_by', 'approved_by']
    },
    'F/47': {
        'name': 'Audit Checklist',
        'section': 3,
        'section_name': 'Quality & Audit',
        'fields': ['serial', 'date', 'audit_ref', 'department', 'checklist_items',
                    'clause', 'requirement', 'compliant', 'evidence', 'auditor']
    },
    'F/48': {
        'name': 'Internal Audit Report',
        'section': 3,
        'section_name': 'Quality & Audit',
        'fields': ['serial', 'date', 'month', 'year', 'scope', 'findings',
                    'nc_count', 'observations', 'recommendations', 'auditor', 'reviewed_by']
    },
    'F/13': {
        'name': 'Purchase Order',
        'section': 4,
        'section_name': 'Procurement & Vendors',
        'fields': ['serial', 'date', 'supplier', 'items', 'description',
                    'quantity', 'unit', 'specifications', 'requested_by', 'approved_by']
    },
    'F/14': {
        'name': 'Incoming Inspection',
        'section': 4,
        'section_name': 'Procurement & Vendors',
        'fields': ['serial', 'date', 'po_ref', 'supplier', 'items_inspected',
                    'result', 'defects', 'inspector']
    },
    'F/15': {
        'name': 'Approved Vendor List',
        'section': 4,
        'section_name': 'Procurement & Vendors',
        'fields': ['serial', 'year', 'vendors', 'name', 'service',
                    'rating', 'status', 'prepared_by', 'approved_by']
    },
    'F/16': {
        'name': 'Supplier Registration Form',
        'section': 4,
        'section_name': 'Procurement & Vendors',
        'fields': ['serial', 'date', 'supplier_name', 'supplier_contact',
                    'service_type', 'qualifications', 'evaluation_result', 'evaluated_by']
    },
    'F/28': {
        'name': 'Training Attendance',
        'section': 5,
        'section_name': 'HR & Training',
        'fields': ['serial', 'date', 'course_name', 'trainer', 'project',
                    'attendees', 'name', 'id', 'department', 'attended',
                    'trainer_signature', 'manager_signature']
    },
    'F/29': {
        'name': 'Training Record',
        'section': 5,
        'section_name': 'HR & Training',
        'fields': ['serial', 'employee_name', 'employee_id', 'department',
                    'course_name', 'training_date', 'trainer', 'result',
                    'score', 'comments', 'recorded_by']
    },
    'F/30': {
        'name': 'Performance Appraisal',
        'section': 5,
        'section_name': 'HR & Training',
        'fields': ['serial', 'employee_name', 'employee_id', 'department',
                    'period', 'criteria', 'criterion', 'score', 'comments',
                    'overall_score', 'recommendations', 'evaluator', 'employee_signature']
    },
    'F/40': {
        'name': 'Competence Matrix',
        'section': 5,
        'section_name': 'HR & Training',
        'fields': ['serial', 'period', 'matrix', 'name', 'role',
                    'skill', 'level', 'prepared_by']
    },
    'F/41': {
        'name': 'Gap Analysis',
        'section': 5,
        'section_name': 'HR & Training',
        'fields': ['serial', 'date', 'matrix_ref', 'gaps',
                    'training_needed', 'prepared_by']
    },
    'F/42': {
        'name': 'Annual Training Plan',
        'section': 5,
        'section_name': 'HR & Training',
        'fields': ['serial', 'year', 'objectives', 'plan', 'course',
                    'target', 'quarter', 'method', 'prepared_by', 'approved_by']
    },
    'F/43': {
        'name': 'Induction Training Record',
        'section': 5,
        'section_name': 'HR & Training',
        'fields': ['serial', 'date', 'employee_name', 'employee_id',
                    'department', 'project', 'qualification', 'trainer',
                    'issued_by', 'performance', 'topics_covered',
                    'trainer_signature', 'manager_signature']
    },
    'F/44': {
        'name': 'Job Description',
        'section': 5,
        'section_name': 'HR & Training',
        'fields': ['serial', 'date', 'job_title', 'department',
                    'responsibilities', 'qualifications_required',
                    'reporting_to', 'prepared_by']
    },
    'F/32': {
        'name': 'R&D Request',
        'section': 6,
        'section_name': 'R&D & Design',
        'fields': ['serial', 'date', 'title', 'objective', 'methodology',
                    'expected_outcome', 'requested_by', 'approved_by']
    },
    'F/34': {
        'name': 'Design Verification',
        'section': 6,
        'section_name': 'R&D & Design',
        'fields': ['serial', 'date', 'project', 'design_ref',
                    'requirements_met', 'findings', 'verified_by']
    },
    'F/35': {
        'name': 'Design Monitoring',
        'section': 6,
        'section_name': 'R&D & Design',
        'fields': ['serial', 'month', 'year', 'project', 'progress',
                    'issues', 'next_steps', 'monitored_by']
    },
    'F/37': {
        'name': 'Experiment Data',
        'section': 6,
        'section_name': 'R&D & Design',
        'fields': ['serial', 'date', 'experiment_title', 'hypothesis',
                    'method', 'results', 'conclusion', 'recorded_by']
    },
    'F/20': {
        'name': 'Review Agenda',
        'section': 7,
        'section_name': 'Management & Documentation',
        'fields': ['serial', 'date', 'location', 'chairperson',
                    'agenda_items', 'item', 'presenter', 'duration', 'prepared_by']
    },
    'F/21': {
        'name': 'Review Minutes',
        'section': 7,
        'section_name': 'Management & Documentation',
        'fields': ['serial', 'meeting_date', 'chairperson', 'attendees',
                    'discussion', 'decisions', 'action_items', 'action',
                    'responsible', 'deadline', 'status', 'minutes_by', 'approved_by']
    },
    'F/23': {
        'name': 'Master List of Records',
        'section': 7,
        'section_name': 'Management & Documentation',
        'fields': ['serial', 'date', 'records', 'form_code',
                    'record_serial', 'description', 'date_created',
                    'storage', 'maintained_by']
    },
    'F/24': {
        'name': 'Objectives & Targets',
        'section': 7,
        'section_name': 'Management & Documentation',
        'fields': ['serial', 'quarter', 'year', 'objectives', 'objective',
                    'target', 'actual', 'status', 'prepared_by', 'reviewed_by']
    },
    'F/45': {
        'name': 'Master List of Documents',
        'section': 7,
        'section_name': 'Management & Documentation',
        'fields': ['serial', 'date', 'documents', 'doc_id', 'title',
                    'version', 'status', 'date_created', 'maintained_by']
    },
    'F/46': {
        'name': 'Change Management',
        'section': 7,
        'section_name': 'Management & Documentation',
        'fields': ['serial', 'date', 'change_type', 'description',
                    'reason', 'impact', 'approved', 'approved_by']
    },
}

# Section mapping for folder names
SECTION_FOLDERS = {
    1: '01- Sales Records',
    2: '02- Operations Records',
    3: '03- Quality Record',
    4: '04 Procurement & Vendors Records',
    5: '05 HR & Training Records',
    6: '06 R&D & Design Records',
    7: '07 Management & Documentation Records',
}

# ============================================================================
# GOOGLE DRIVE API
# ============================================================================
def get_access_token():
    with open(TOKEN_PATH) as f:
        data = json.load(f)
    # Refresh if needed
    post_data = urllib.parse.urlencode({
        'client_id': data['client_id'],
        'client_secret': data['client_secret'],
        'refresh_token': data['refresh_token'],
        'grant_type': 'refresh_token',
    }).encode()
    req = urllib.request.Request('https://oauth2.googleapis.com/token', data=post_data)
    resp = urllib.request.urlopen(req)
    result = json.loads(resp.read())
    data['access_token'] = result['access_token']
    with open(TOKEN_PATH, 'w') as f:
        json.dump(data, f, indent=2)
    return result['access_token']

def drive_api(path, token):
    url = f"https://www.googleapis.com{path}"
    req = urllib.request.Request(url)
    req.add_header('Authorization', f'Bearer {token}')
    resp = urllib.request.urlopen(req)
    return json.loads(resp.read())

def download_file(file_id, filepath, token):
    meta = drive_api(f"/drive/v3/files/{file_id}?fields=mimeType,name", token)
    mime = meta.get('mimeType', '')
    
    dirname = os.path.dirname(filepath)
    basename = os.path.basename(filepath).replace('/', '-').replace('\\', '-')
    filepath = os.path.join(dirname, basename)
    os.makedirs(dirname, exist_ok=True)
    
    if mime == 'application/vnd.google-apps.document':
        url = f"https://www.googleapis.com/drive/v3/files/{file_id}/export?mimeType=application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif mime == 'application/vnd.google-apps.spreadsheet':
        url = f"https://www.googleapis.com/drive/v3/files/{file_id}/export?mimeType=text/csv"
    else:
        url = f"https://www.googleapis.com/drive/v3/files/{file_id}?alt=media"
    
    req = urllib.request.Request(url)
    req.add_header('Authorization', f'Bearer {token}')
    resp = urllib.request.urlopen(req)
    with open(filepath, 'wb') as f:
        f.write(resp.read())
    return (filepath, os.path.getsize(filepath))

# ============================================================================
# DOCX EXTRACTION — Generic table-based extractor
# ============================================================================
def extract_table_data(docx_path):
    """Extract all data from a DOCX with a single main table.
    Returns a dict of label->value pairs and any table rows."""
    try:
        doc = Document(docx_path)
    except Exception as e:
        return {'_error': str(e), '_path': docx_path}
    
    result = {}
    table_data = []
    
    for ti, table in enumerate(doc.tables):
        rows_data = []
        for row in table.rows:
            cells = [cell.text.strip().replace('\n', ' | ') for cell in row.cells]
            rows_data.append(cells)
        table_data.append(rows_data)
    
    result['_tables'] = table_data
    result['_num_tables'] = len(doc.tables)
    result['_num_paragraphs'] = len([p for p in doc.paragraphs if p.text.strip()])
    
    # Extract paragraphs as well
    result['_paragraphs'] = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    
    return result

def extract_serial_and_date(table_data):
    """Try to find serial number and date from first rows of the table."""
    serial = ''
    date = ''
    for row in table_data:
        for cell in row:
            # Look for serial pattern like F/08-001, F/30-029 etc
            m = re.search(r'(F[/\-]\d+[-]\d+)', cell)
            if m:
                serial = m.group(1).replace('\\', '/')
            # Look for date
            m2 = re.search(r'(\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4})', cell)
            if m2 and not date:
                date = m2.group(1)
    return serial, date

# ============================================================================
# SUPABASE API
# ============================================================================
def supabase_insert(record):
    """Insert a record into QBase Supabase."""
    url = f"{SUPABASE_URL}/rest/v1/records"
    data = json.dumps(record).encode()
    
    req = urllib.request.Request(url, data=data, method='POST')
    req.add_header('apikey', SERVICE_KEY)
    req.add_header('Authorization', f'Bearer {SERVICE_KEY}')
    req.add_header('Content-Type', 'application/json')
    req.add_header('Prefer', 'return=representation')
    
    try:
        resp = urllib.request.urlopen(req)
        return json.loads(resp.read())
    except urllib.error.HTTPError as e:
        error_body = e.read().decode()
        return {'_error': f'HTTP {e.code}: {error_body[:200]}'}

def supabase_get_existing():
    """Get all existing records from QBase."""
    url = f"{SUPABASE_URL}/rest/v1/records?select=id,form_code,serial,status"
    req = urllib.request.Request(url)
    req.add_header('apikey', SERVICE_KEY)
    req.add_header('Authorization', f'Bearer {SERVICE_KEY}')
    req.add_header('Range', '0-9999')
    
    resp = urllib.request.urlopen(req)
    return json.loads(resp.read())

# ============================================================================
# FORM-SPECIFIC EXTRACTORS
# Each extractor takes the raw table data and returns a form_data dict
# ============================================================================

def extract_f08(tables):
    """F/08 Order Form - single table with header row, then data rows."""
    if not tables:
        return {}
    rows = tables[0]
    data = {}
    
    # Row 0: form name + serial
    # Row 1: serial + date
    if len(rows) > 1:
        for cell in rows[1]:
            m = re.search(r'F/08[-]\d+', cell)
            if m:
                data['serial'] = m.group(0)
            m2 = re.search(r'\d{1,2}/\d{1,2}/\d{2,4}', cell)
            if m2:
                data['date'] = m2.group(0)
    
    # Row 2: Customer
    if len(rows) > 2 and len(rows[2]) > 3:
        data['client_name'] = rows[2][3] if len(rows[2]) > 3 else ''
    
    # Row 3: Mode of Receipt
    if len(rows) > 3 and len(rows[3]) > 3:
        data['mode_of_receipt'] = rows[3][3] if len(rows[3]) > 3 else ''
    
    # Rows 5-9: Product items
    items = []
    for i in range(5, min(10, len(rows))):
        row = rows[i]
        if len(row) > 1:
            product = row[1] if len(row) > 1 else ''
            spec = row[4] if len(row) > 4 else ''
            qty = row[11] if len(row) > 11 else ''
            if product.strip():
                items.append(f"{product}: {spec} (Qty: {qty})")
    data['items'] = '\n'.join(items) if items else ''
    
    # Row 10: Test Certificate
    if len(rows) > 10:
        for cell in rows[10]:
            if cell.strip() in ('Yes', 'No'):
                data['test_certificate_required'] = cell.strip()
    
    # Row 11: Delivery Schedule
    if len(rows) > 11 and len(rows[11]) > 7:
        data['delivery_schedule'] = rows[11][7] if len(rows[11]) > 7 else ''
    
    # Row 13: Order status
    if len(rows) > 13:
        for cell in rows[13]:
            if cell.strip() in ('Accepted', 'Rejected', 'On Hold'):
                data['order_status'] = cell.strip()
    
    # Row 14: Remarks and Reviewed By
    if len(rows) > 14:
        for cell in rows[14]:
            if cell.startswith('Remarks'):
                data['remarks'] = cell.split(':', 1)[1].strip() if ':' in cell else cell.replace('Remarks', '').strip()
            if cell.startswith('Reviewed By'):
                data['reviewed_by'] = cell.split(':', 1)[1].strip() if ':' in cell else cell.replace('Reviewed By', '').strip()
    
    # Row 15: Bill No
    if len(rows) > 15 and len(rows[15]) > 3:
        data['bill_no'] = rows[15][3] if len(rows[15]) > 3 else ''
    
    # Row 16: Despatch Date
    if len(rows) > 16 and len(rows[16]) > 3:
        data['despatch_date'] = rows[16][3] if len(rows[16]) > 3 else ''
    
    return data

# Generic extractor - extract by scanning all rows for label: value patterns
def extract_generic(tables, form_code):
    """Generic extractor that finds label-value pairs in table cells."""
    if not tables:
        return {}
    
    data = {}
    rows = tables[0]
    
    # First, find serial from filename or table content
    for row in rows:
        for cell in row:
            m = re.search(rf'{re.escape(form_code)}[-]\d+', cell)
            if m:
                data['serial'] = m.group(0)
    
    # Scan for label-value patterns
    label_map = {
        'serial': ['sr. no', 'serial no', 'serial number', 'ref. no', 'ref no'],
        'date': ['date'],
        'client_name': ['customer', 'client name', 'name of customer', 'name'],
        'mode_of_receipt': ['mode of receipt', 'mode'],
        'description': ['description', 'description of', 'details'],
        'remarks': ['remarks', 'remark', 'comments'],
        'reviewed_by': ['reviewed by', 'reviewed', 'approved by'],
    }
    
    for ri, row in enumerate(rows):
        for ci, cell in enumerate(row):
            cell_lower = cell.lower().strip()
            if ':' in cell:
                label, _, value = cell.partition(':')
                label_lower = label.lower().strip()
                value = value.strip()
                
                for key, patterns in label_map.items():
                    if key not in data and any(p in label_lower for p in patterns):
                        if value:
                            data[key] = value
    
    return data

# ============================================================================
# MAIN MIGRATION LOGIC
# ============================================================================

def map_form_code(folder_name):
    """Map a Drive folder name to a QBase form code."""
    folder_map = {
        'F/08 - Order Form': 'F/08',
        'F/09 - Customer Complaint': 'F/09',
        'F/10 - Customer Feedback': 'F/10',
        'F/50 - Customer Property': 'F/50',
        'F/11 - Production Plan': 'F/11',
        'F/19 - Product Description': 'F/19',
        'F/12 - Non-Conforming': 'F/12',
        'F/17 - QA Test Request': 'F/17',
        'F/18 - Product Re-Call': 'F/18',
        'F/22 - Corrective Action': 'F/22',
        'F/25 - Audit Plan': 'F/25',
        'F/47 - Audit Checklist': 'F/47',
        'F/48 - Internal Audit Report': 'F/48',
        'F-16 Supplier Registration Form': 'F/16',
        'F/13 - Purchase Order': 'F/13',
        'F/14 - Incoming Inspection': 'F/14',
        'F/15 - Approved Vendor List': 'F/15',
        'F/28 - Training Attendance': 'F/28',
        'F/29 - Training Record': 'F/29',
        'F/30 - Performance Appraisal': 'F/30',
        'F/40 - Competence Matrix': 'F/40',
        'F/41 - Gap Analysis': 'F/41',
        'F/42 - Annual Training Plan': 'F/42',
        'F/43 - Induction Training Record': 'F/43',
        'F/44 - Job Description': 'F/44',
        'F/32 - R&D Request': 'F/32',
        'F/34 - Design Verification': 'F/34',
        'F/35 - Design Monitoring': 'F/35',
        'F/37 - Experiment Data': 'F/37',
        'F/20 - Review Agenda': 'F/20',
        'F/21 - Review Minutes': 'F/21',
        'F/23 - Master List of Records': 'F/23',
        'F/24 - Objectives & Targets': 'F/24',
        'F/45 - Master List of Docs': 'F/45',
        'F/46 - Change Management': 'F/46',
    }
    return folder_map.get(folder_name)

if __name__ == '__main__':
    cmd = sys.argv[1] if len(sys.argv) > 1 else 'analyze'
    
    if cmd == 'analyze':
        print("Analyzing downloaded samples...")
        SAMPLES_DIR.mkdir(parents=True, exist_ok=True)
        files = sorted(SAMPLES_DIR.glob('*'))
        print(f"Found {len(files)} sample files")
        
        for f in files:
            data = extract_table_data(str(f))
            name = f.name
            print(f"\n{'='*60}")
            print(f"📄 {name}")
            print(f"   Tables: {data.get('_num_tables', 0)}, Paragraphs: {data.get('_num_paragraphs', 0)}")
            if data.get('_tables'):
                for ti, table in enumerate(data['_tables']):
                    print(f"   Table {ti}: {len(table)} rows x {len(table[0]) if table else 0} cols")
                    for ri, row in enumerate(table[:3]):
                        cells_preview = [c[:40] for c in row[:5]]
                        print(f"      Row {ri}: {cells_preview}")
                    if len(table) > 3:
                        print(f"      ... ({len(table)-3} more rows)")
    
    elif cmd == 'migrate':
        print("Starting migration...")
        # Full migration logic will be added after analysis
        print("Migration not yet implemented - run 'analyze' first")
    
    elif cmd == 'count':
        # Count files on Drive vs existing records in Supabase
        token = get_access_token()
        existing = supabase_get_existing()
        existing_serials = {r['serial'] for r in existing if r.get('serial')}
        
        print(f"Existing QBase records: {len(existing)}")
        print(f"Existing serials: {sorted(existing_serials)}")