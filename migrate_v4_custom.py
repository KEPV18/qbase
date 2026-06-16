#!/usr/bin/env python3
"""
QMS → QBase Migration — Custom Extractors v4
Per-form extractors for proper data extraction.
Reads local DOCX files, extracts with custom logic per form type.
"""
import json, os, sys, re, urllib.request
from docx import Document
from pathlib import Path
from collections import Counter

RECORDS_DIR = Path('/home/Kepv/Downloads/03.Records')

def load_env():
    env = {}
    with open('/mnt/ahmed/Projects/qbase/.env.local') as f:
        for line in f:
            line = line.strip()
            if line and not line.startswith('#') and '=' in line:
                k,v=line.split('=',1)
                env[k]=v
    return env

ENV = load_env()
SUPABASE_URL = ENV['SUPABASE_URL']
SERVICE_KEY = ENV['SUPABASE_SERVICE_ROLE_KEY']

# Form config
FORM_CONFIG = {
    'F/08': ('Order Form', 1, 'Sales & Customer Service'),
    'F/09': ('Customer Complaint', 1, 'Sales & Customer Service'),
    'F/10': ('Customer Feedback', 1, 'Sales & Customer Service'),
    'F/11': ('Production Plan', 2, 'Operations & Production'),
    'F/12': ('Non-Conforming', 3, 'Quality & Audit'),
    'F/13': ('Purchase Order', 4, 'Procurement & Vendors'),
    'F/14': ('Incoming Inspection', 4, 'Procurement & Vendors'),
    'F/15': ('Approved Vendor List', 4, 'Procurement & Vendors'),
    'F/16': ('Supplier Registration Form', 4, 'Procurement & Vendors'),
    'F/17': ('QA Test Request', 3, 'Quality & Audit'),
    'F/18': ('Product Re-Call', 3, 'Quality & Audit'),
    'F/19': ('Product Description', 2, 'Operations & Production'),
    'F/20': ('Review Agenda', 7, 'Management & Documentation'),
    'F/21': ('Review Minutes', 7, 'Management & Documentation'),
    'F/22': ('Corrective Action', 3, 'Quality & Audit'),
    'F/23': ('Master List of Records', 7, 'Management & Documentation'),
    'F/24': ('Objectives & Targets', 7, 'Management & Documentation'),
    'F/25': ('Audit Plan', 3, 'Quality & Audit'),
    'F/28': ('Training Attendance', 5, 'HR & Training'),
    'F/29': ('Training Record', 5, 'HR & Training'),
    'F/30': ('Performance Appraisal', 5, 'HR & Training'),
    'F/32': ('R&D Request', 6, 'R&D & Design'),
    'F/34': ('Design Verification', 6, 'R&D & Design'),
    'F/35': ('Design Monitoring', 6, 'R&D & Design'),
    'F/37': ('Experiment Data', 6, 'R&D & Design'),
    'F/40': ('Competence Matrix', 5, 'HR & Training'),
    'F/41': ('Gap Analysis', 5, 'HR & Training'),
    'F/42': ('Annual Training Plan', 5, 'HR & Training'),
    'F/43': ('Induction Training Record', 5, 'HR & Training'),
    'F/44': ('Job Description', 5, 'HR & Training'),
    'F/45': ('Master List of Docs', 7, 'Management & Documentation'),
    'F/46': ('Change Management', 7, 'Management & Documentation'),
    'F/47': ('Audit Checklist', 3, 'Quality & Audit'),
    'F/48': ('Internal Audit Report', 3, 'Quality & Audit'),
    'F/50': ('Customer Property', 1, 'Sales & Customer Service'),
}

# Map folder patterns
FORM_FOLDER_MAP = {
    'F-08': 'F/08', 'F-09': 'F/09', 'F-10': 'F/10', 'F-11': 'F/11', 'F-12': 'F/12',
    'F-13': 'F/13', 'F-14': 'F/14', 'F-15': 'F/15', 'F-16': 'F/16', 'F-17': 'F/17',
    'F-18': 'F/18', 'F-19': 'F/19', 'F-20': 'F/20', 'F-21': 'F/21', 'F-22': 'F/22',
    'F-23': 'F/23', 'F-24': 'F/24', 'F-25': 'F/25', 'F-28': 'F/28', 'F-29': 'F/29',
    'F-30': 'F/30', 'F-32': 'F/32', 'F-34': 'F/34', 'F-35': 'F/35', 'F-37': 'F/37',
    'F-40': 'F/40', 'F-41': 'F/41', 'F-42': 'F/42', 'F-43': 'F/43', 'F-44': 'F/44',
    'F-45': 'F/45', 'F-46': 'F/46', 'F-47': 'F/47', 'F-48': 'F/48', 'F-50': 'F/50',
}

# =============================================================================
# CUSTOM EXTRACTORS
# =============================================================================

def clean_cells(row):
    """Remove duplicate consecutive cells."""
    cells = [cell.text.strip().replace('\n',' ') for cell in row.cells]
    clean = []
    for c in cells:
        if clean and clean[-1] == c: continue
        clean.append(c)
    return clean

def extract_generic(filepath):
    """Generic fallback extractor."""
    doc = Document(filepath)
    data = {'_type': 'generic'}
    for table in doc.tables:
        for row in table.rows:
            cells = clean_cells(row)
            for i in range(len(cells)-1):
                label = cells[i].lower().strip()
                value = cells[i+1].strip()
                if value and label and len(label) < 80 and len(value) < 500:
                    # Check for label:value pattern
                    if ':' in cells[i]:
                        parts = cells[i].split(':', 1)
                        l = parts[0].strip().lower()
                        v = parts[1].strip() if len(parts) > 1 else ''
                        if v: data[l] = v
                    elif len(label) > 3 and len(value) > 0:
                        data[label] = value
    return data

def extract_f08(filepath):
    """F/08 Order Form"""
    doc = Document(filepath)
    data = {'_type': 'f08'}
    table = doc.tables[0] if doc.tables else None
    if not table: return data

    for row in table.rows:
        cells = clean_cells(row)
        row_text = '|'.join(cells)

        # Date
        m = re.search(r'Date\s*[🡪:\t\s]+(\d{2}/\d{2}/\d{4})', row_text)
        if m: data['date'] = m.group(1)

        # Customer
        if 'Customer' in row_text and not data.get('client_name'):
            for c in cells:
                if 'Customer' not in c and len(c) > 3:
                    data['client_name'] = c; break

        # Mode of Receipt
        if 'Receipt' in row_text and not data.get('mode_of_receipt'):
            for c in cells:
                if 'Mode' not in c and 'Receipt' not in c and len(c) > 5:
                    data['mode_of_receipt'] = c; break

        # Items header
        if 'Product Name' in row_text and 'Specifications' in row_text:
            data['_items_header'] = True
            continue

        # Items rows
        if data.get('_items_header'):
            if any(k in row_text for k in ['Video Detection', 'Vocal AI', 'Tennis', 'Omniaz', 'ETH AI', 'Widget']):
                product = [c for c in cells if 'Service' in c or 'Analytics' in c or 'Annotation' in c or 'Testing' in c or 'Widget' in c]
                specs = [c for c in cells if 'labeling' in c or 'Conversational' in c or 'Match tagging' in c or 'Data annotation' in c or 'validation' in c or 'optimization' in c or 'evaluation' in c or 'Spec' in c]
                qty = [c for c in cells if 'client volume' in c.lower() or c.strip().isdigit()]
                if product:
                    item = {'product_name': product[0]}
                    if specs: item['specifications'] = specs[0]
                    if qty: item['qty'] = qty[0]
                    elif 'As per client volume' in row_text: item['qty'] = 'As per client volume'
                    if '_items' not in data: data['_items'] = []
                    data['_items'].append(item)

        # Test certificate
        if 'Requirement' in row_text and 'Test Certificate' in row_text:
            for c in cells:
                if c.strip() in ['No', 'Yes', 'N/A']:
                    data['test_certificate_required'] = 'Yes' if c == 'Yes' else 'No'

        # Delivery schedule
        if 'Delivery' in row_text and 'Schedule' in row_text:
            for c in cells:
                if 'Delivery' not in c and 'Schedule' not in c and len(c) > 10:
                    data['delivery_schedule'] = c; break

        # Statutory
        if 'Statutory' in row_text or ('Regulatory' in row_text and 'Requirements' in row_text):
            for c in cells:
                if len(c) > 10 and 'Statutory' not in c:
                    data['complies'] = c; break

        # Order status
        if 'Accepted' in row_text:
            data['order_status'] = 'Accepted'
        if 'Order' in row_text and 'Rejected' in row_text:
            data['order_status'] = 'Rejected'

        # Bill No
        if 'Bill No' in row_text:
            for c in cells:
                if 'Bill' not in c and 'No' not in c and c.strip():
                    data['bill_no'] = c; break

        # Despatch
        if 'Despatch' in row_text:
            for c in cells:
                if 'Despatch' not in c and c.strip():
                    data['despatch_date'] = c; break

        # Remarks
        if 'Remarks' in row_text:
            for c in cells:
                if 'Remarks' not in c and len(c) > 10:
                    data['remarks'] = c; break

        # Reviewed by
        if 'Reviewed' in row_text and 'Representative' in row_text:
            for c in cells:
                if 'Reviewed' not in c and 'Representative' not in c and len(c) > 3:
                    data['reviewed_by'] = c; break

    if '_items_header' in data: del data['_items_header']
    return data

def extract_f10(filepath):
    """F/10 Customer Feedback"""
    doc = Document(filepath)
    data = {'_type': 'f10'}
    table = doc.tables[0] if doc.tables else None
    if not table: return data

    for row in table.rows:
        cells = clean_cells(row)
        row_text = '|'.join(cells)

        # Serial
        m = re.search(r'F[/\-](\d{2})[-\s]*0*(\d+)', row_text)
        if m and not data.get('serial'): data['serial'] = f"F/{m.group(1)}-{int(m.group(2)):03d}"

        # Date
        m = re.search(r'(\d{2}/\d{2}/\d{4})', row_text)
        if m and not data.get('date'): data['date'] = m.group(1)

        # Client
        if 'Customer' in row_text or 'Client' in row_text:
            for c in cells:
                if 'Customer' not in c and 'Client' not in c and len(c) > 3 and 'Feedback' not in c:
                    data['client_name'] = c; break

        # Feedback type
        if 'Type' in row_text and 'Feedback' in row_text:
            for c in cells:
                if 'Complaint' in c or 'Suggestion' in c or 'Appreciation' in c:
                    data['feedback_type'] = c; break

        # Description
        if 'Description' in row_text or 'Details' in row_text:
            for c in cells:
                if 'Description' not in c and len(c) > 10:
                    data['description'] = c; break

        # Action taken
        if 'Action' in row_text and 'Taken' in row_text:
            for c in cells:
                if 'Action' not in c and len(c) > 5:
                    data['action_taken'] = c; break

        # Responsible person
        if 'Responsible' in row_text or 'Person' in row_text:
            for c in cells:
                if 'Responsible' not in c and len(c) > 3:
                    data['responsible_person'] = c; break

        # Status
        if 'Status' in row_text:
            for c in cells:
                if c in ['Open', 'Closed', 'In Progress', 'Resolved', 'Pending']:
                    data['status'] = c; break

    return data

def extract_f19(filepath):
    """F/19 Product Description"""
    doc = Document(filepath)
    data = {'_type': 'f19'}
    table = doc.tables[0] if doc.tables else None
    if not table: return data

    for row in table.rows:
        cells = clean_cells(row)
        row_text = '|'.join(cells)

        # Parameters with 🡪 or :
        for c in cells:
            for sep in ['🡪', ':', '\t']:
                if sep in c:
                    parts = c.split(sep, 1)
                    if len(parts) == 2:
                        label = parts[0].strip().lower()
                        value = parts[1].strip()
                        if value and len(label) > 3:
                            key = label.replace(' ', '_')
                            data[key] = value

        # Row-based label:value
        if len(cells) >= 2:
            for i in range(len(cells)-1):
                label = cells[i].strip().lower()
                value = cells[i+1].strip()
                if value and 'parameter' not in label and len(label) > 3:
                    key = label.replace(' ', '_').replace('(', '').replace(')', '')
                    if key not in data:
                        data[key] = value

    return data

def extract_f20(filepath):
    """F/20 Review Agenda — paragraph-based"""
    doc = Document(filepath)
    data = {'_type': 'f20'}
    all_text = '\n'.join([p.text.strip() for p in doc.paragraphs if p.text.strip()])

    # Date
    m = re.search(r'Meeting\s*Date\s*[:\s]*(\d{2}/\d{2}/\d{4})', all_text, re.IGNORECASE)
    if m: data['date'] = m.group(1)

    # Venue
    m = re.search(r'Venue\s*[:\s]*(.+)', all_text, re.IGNORECASE)
    if m: data['venue'] = m.group(1).strip()

    # Agenda items
    agenda_items = []
    lines = all_text.split('\n')
    for line in lines:
        if re.match(r'^\d+\)[\.\s]', line) or re.match(r'^\d+\.', line):
            agenda_items.append(line.strip())
    if agenda_items:
        data['agenda_items'] = '\n'.join(agenda_items)

    data['_raw_text'] = all_text[:2000]
    return data

def extract_f21(filepath):
    """F/21 Review Minutes — paragraph-based"""
    doc = Document(filepath)
    data = {'_type': 'f21'}
    all_text = '\n'.join([p.text.strip() for p in doc.paragraphs if p.text.strip()])

    # Date
    m = re.search(r'Meeting\s*Date\s*[:\s]*(\d{2}/\d{2}/\d{4})', all_text, re.IGNORECASE)
    if m: data['date'] = m.group(1)

    # Attendees
    m = re.search(r'Attendees\s*[:\s]*(.+?)(?:\n|Agenda|Minutes)', all_text, re.IGNORECASE | re.DOTALL)
    if m: data['attendees'] = m.group(1).strip()[:500]

    # Decisions
    decisions = []
    lines = all_text.split('\n')
    for i, line in enumerate(lines):
        if 'decision' in line.lower() or 'agreed' in line.lower() or 'approved' in line.lower():
            decisions.append(line.strip())
    if decisions:
        data['decisions'] = '\n'.join(decisions[:20])

    data['_raw_text'] = all_text[:2000]
    return data

def extract_f23(filepath):
    """F/23 Master List of Records — table with records list"""
    doc = Document(filepath)
    data = {'_type': 'f23'}
    table = doc.tables[0] if doc.tables else None
    if not table: return data

    records = []
    headers = None
    for ri, row in enumerate(table.rows):
        cells = clean_cells(row)
        if ri == 0:
            headers = [c.lower().strip() for c in cells]
            continue
        if len(cells) >= 3 and any(cells):
            record = {}
            for ci, h in enumerate(headers or []):
                if ci < len(cells):
                    record[h] = cells[ci]
            if record.get('record no') or record.get('title of record'):
                records.append(record)

    if records:
        data['records_list'] = json.dumps(records, ensure_ascii=False)

    return data

def extract_f28(filepath):
    """F/28 Training Attendance — table with participants"""
    doc = Document(filepath)
    data = {'_type': 'f28'}
    table = doc.tables[0] if doc.tables else None
    if not table: return data

    participants = []
    headers = None
    for ri, row in enumerate(table.rows):
        cells = clean_cells(row)
        if ri == 0:
            headers = [c.lower().strip() for c in cells]
            continue
        if len(cells) >= 4:
            participant = {}
            for ci, h in enumerate(headers or []):
                if ci < len(cells):
                    participant[h] = cells[ci]
            if participant.get('name of the participant') or participant.get('name'):
                participants.append(participant)

    if participants:
        data['participants'] = json.dumps(participants, ensure_ascii=False)

    # Find training topic/date from header area
    all_text = '\n'.join(cells[0] for row in table.rows[:3] for cell in row.cells if cell.text.strip())
    m = re.search(r'(\d{2}/\d{2}/\d{4})', all_text)
    if m: data['training_date'] = m.group(1)

    return data

def extract_f29(filepath):
    """F/29 Training Record — table with employee data"""
    doc = Document(filepath)
    data = {'_type': 'f29'}
    table = doc.tables[0] if doc.tables else None
    if not table: return data

    employees = []
    headers = None
    for ri, row in enumerate(table.rows):
        cells = clean_cells(row)
        if ri == 0:
            headers = [c.lower().strip() for c in cells]
            continue
        if len(cells) >= 5:
            emp = {}
            for ci, h in enumerate(headers or []):
                if ci < len(cells):
                    emp[h] = cells[ci]
            if emp.get('name of employee') or emp.get('name'):
                employees.append(emp)

    if employees:
        data['employees'] = json.dumps(employees, ensure_ascii=False)

    # Find assessment date
    all_text = '\n'.join([p.text.strip() for p in doc.paragraphs if p.text.strip()])
    m = re.search(r'Assessment\s*[:\s]*(\w+\s*\d{4})', all_text, re.IGNORECASE)
    if m: data['assessment_date'] = m.group(1)

    return data

def extract_f35(filepath):
    """F/35 Design Monitoring — table with projects"""
    doc = Document(filepath)
    data = {'_type': 'f35'}
    table = doc.tables[0] if doc.tables else None
    if not table: return data

    projects = []
    headers = None
    for ri, row in enumerate(table.rows):
        cells = clean_cells(row)
        if ri == 0:
            headers = [c.lower().strip() for c in cells]
            continue
        if len(cells) >= 3:
            proj = {}
            for ci, h in enumerate(headers or []):
                if ci < len(cells):
                    proj[h] = cells[ci]
            if proj.get('product name') or proj.get('name of the customer'):
                projects.append(proj)

    if projects:
        data['projects'] = json.dumps(projects, ensure_ascii=False)

    return data

def extract_f44(filepath):
    """F/44 Job Description — paragraph-based"""
    doc = Document(filepath)
    data = {'_type': 'f44'}
    all_text = '\n'.join([p.text.strip() for p in doc.paragraphs if p.text.strip()])

    # Job title
    m = re.search(r'Job\s*Title\s*[:\s]*(.+)', all_text, re.IGNORECASE)
    if m: data['job_title'] = m.group(1).strip()

    # Department
    m = re.search(r'Department\s*[:\s]*(.+)', all_text, re.IGNORECASE)
    if m: data['department'] = m.group(1).strip()

    # Responsibilities
    m = re.search(r'Responsibilities\s*[:\s]*(.+?)(?:Qualifications|Skills|Reporting|$)', all_text, re.IGNORECASE | re.DOTALL)
    if m: data['responsibilities'] = m.group(1).strip()[:1000]

    # Qualifications
    m = re.search(r'Qualifications\s*[:\s]*(.+?)(?:Skills|Experience|$)', all_text, re.IGNORECASE | re.DOTALL)
    if m: data['qualifications'] = m.group(1).strip()[:500]

    data['_raw_text'] = all_text[:2000]
    return data

def extract_f45(filepath):
    """F/45 Master List of Docs — table with documents"""
    doc = Document(filepath)
    data = {'_type': 'f45'}
    table = doc.tables[0] if doc.tables else None
    if not table: return data

    docs = []
    headers = None
    for ri, row in enumerate(table.rows):
        cells = clean_cells(row)
        if ri == 0:
            headers = [c.lower().strip() for c in cells]
            continue
        if len(cells) >= 4:
            d = {}
            for ci, h in enumerate(headers or []):
                if ci < len(cells):
                    d[h] = cells[ci]
            if d.get('title of document') or d.get('document number'):
                docs.append(d)

    if docs:
        data['documents'] = json.dumps(docs, ensure_ascii=False)

    # Last updated
    all_text = '\n'.join(cells[0] for row in table.rows[:5] for cell in row.cells if cell.text.strip())
    m = re.search(r'Last\s*Updated\s*[:\s]*(.+)', all_text, re.IGNORECASE)
    if m: data['last_updated_by'] = m.group(1).strip()

    m = re.search(r'Update\s*Date\s*[:\s]*(\d{2}/\d{2}/\d{4})', all_text, re.IGNORECASE)
    if m: data['update_date'] = m.group(1)

    return data

def extract_f47(filepath):
    """F/47 Audit Checklist — table with checklist items"""
    doc = Document(filepath)
    data = {'_type': 'f47'}
    table = doc.tables[0] if doc.tables else None
    if not table: return data

    items = []
    for ri, row in enumerate(table.rows):
        cells = clean_cells(row)
        if len(cells) >= 3:
            # Check for requirement + status pattern
            req = [c for c in cells if len(c) > 10 and not c.strip() in ['OBS', 'OK', 'NC', 'NA']]
            status = [c for c in cells if c.strip() in ['OBS', 'OK', 'NC', 'NA', 'Compliant', 'Non-Compliant']]
            if req and status:
                items.append({'requirement': req[0], 'status': status[0]})
            elif req and len(cells) >= 2:
                items.append({'requirement': req[0], 'status': cells[-1]})

    if items:
        data['checklist_items'] = json.dumps(items, ensure_ascii=False)

    return data

def extract_f48(filepath):
    """F/48 Internal Audit Report — mixed table + paragraphs"""
    doc = Document(filepath)
    data = {'_type': 'f48'}

    # Table first
    table = doc.tables[0] if doc.tables else None
    if table:
        for row in table.rows:
            cells = clean_cells(row)
            row_text = '|'.join(cells)

            if 'DATE OF AUDIT' in row_text or 'Date of Audit' in row_text:
                m = re.search(r'(\d{2}/\d{2}/\d{4})', row_text)
                if m: data['date'] = m.group(1)

            if 'AUDIT REPORT NO' in row_text or 'Report No' in row_text:
                m = re.search(r'F[/\-](\d{2})[-\s]*0*(\d+)', row_text)
                if m: data['serial'] = f"F/{m.group(1)}-{int(m.group(2)):03d}"

            if 'TYPE OF AUDIT' in row_text or 'Type' in row_text:
                for c in cells:
                    if 'Internal' in c or 'External' in c:
                        data['audit_type'] = c; break

            if 'AUDIT TEAM' in row_text:
                for c in cells:
                    if 'Auditor' in c or 'Ahmed' in c or 'Team' in c:
                        data['audit_team'] = c; break

            if 'AUDIT SCOPE' in row_text:
                for c in cells:
                    if 'Scope' not in c and len(c) > 10:
                        data['audit_scope'] = c; break

    # Paragraphs for findings
    all_text = '\n'.join([p.text.strip() for p in doc.paragraphs if p.text.strip()])

    m = re.search(r'FINDINGS?\s*[:\s]*(.+?)(?:CONCLUSION|RECOMMENDATION|$)', all_text, re.IGNORECASE | re.DOTALL)
    if m: data['findings'] = m.group(1).strip()[:1000]

    m = re.search(r'CONCLUSION\s*[:\s]*(.+?)(?:SIGNATURE|$)', all_text, re.IGNORECASE | re.DOTALL)
    if m: data['conclusion'] = m.group(1).strip()[:500]

    return data

def extract_f50(filepath):
    """F/50 Customer Property — table with property items"""
    doc = Document(filepath)
    data = {'_type': 'f50'}
    table = doc.tables[0] if doc.tables else None
    if not table: return data

    properties = []
    headers = None
    for ri, row in enumerate(table.rows):
        cells = clean_cells(row)
        if ri == 0:
            headers = [c.lower().strip() for c in cells]
            continue
        if len(cells) >= 5:
            prop = {}
            for ci, h in enumerate(headers or []):
                if ci < len(cells):
                    prop[h] = cells[ci]
            if prop.get('name of property') or prop.get('received date'):
                properties.append(prop)

    if properties:
        data['properties'] = json.dumps(properties, ensure_ascii=False)

    return data

# Dispatch table
EXTRACTORS = {
    'F/08': extract_f08,
    'F/10': extract_f10,
    'F/19': extract_f19,
    'F/20': extract_f20,
    'F/21': extract_f21,
    'F/23': extract_f23,
    'F/28': extract_f28,
    'F/29': extract_f29,
    'F/35': extract_f35,
    'F/44': extract_f44,
    'F/45': extract_f45,
    'F/47': extract_f47,
    'F/48': extract_f48,
    'F/50': extract_f50,
}

def extract_form(form_code, filepath):
    """Route to correct extractor."""
    extractor = EXTRACTORS.get(form_code, extract_generic)
    return extractor(filepath)

# =============================================================================
# SCAN & MIGRATE
# =============================================================================

def main():
    cmd = sys.argv[1] if len(sys.argv) > 1 else 'analyze'

    # Get existing records
    url = f"{SUPABASE_URL}/rest/v1/records?select=serial,form_code,status,form_data"
    req = urllib.request.Request(url)
    req.add_header('apikey', SERVICE_KEY)
    req.add_header('Authorization', f'Bearer {SERVICE_KEY}')
    req.add_header('Range', '0-99999')
    existing = json.loads(urllib.request.urlopen(req, timeout=30).read())
    existing_by_serial = {r['serial']: r for r in existing if r.get('serial')}
    print(f"📊 Existing QBase records: {len(existing)} ({len(existing_by_serial)} unique)")

    # Scan local files
    all_files = []
    for section_dir in sorted(RECORDS_DIR.iterdir()):
        if not section_dir.is_dir(): continue
        if 'form' in section_dir.name.lower(): continue
        for form_dir in sorted(section_dir.iterdir()):
            if not form_dir.is_dir(): continue
            # Determine form code
            fc = None
            for prefix, code in FORM_FOLDER_MAP.items():
                if prefix in form_dir.name:
                    fc = code; break
            if not fc: continue
            for ext in ['*.docx', '*.doc']:
                for f in sorted(form_dir.rglob(ext)):
                    if 'Archive' in str(f) or 'archive' in str(f): continue
                    all_files.append((f, fc))

    print(f"📄 Found {len(all_files)} DOCX files")
    print()

    total = 0
    updated = 0
    skipped = 0
    errors = []

    for i, (fpath, form_code) in enumerate(all_files):
        total += 1
        if i % 20 == 0:
            print(f"  [{i:3d}/{len(all_files)}] Processing {form_code}...")
            sys.stdout.flush()

        # Extract serial from filename
        m = re.search(r'[Ff][_/\-](\d{2})[\-_](\d+)', fpath.name, re.IGNORECASE)
        if not m:
            errors.append(f"{form_code}/{fpath.name}: no serial in filename")
            continue
        serial = f"F/{m.group(1)}-{int(m.group(2)):03d}"

        # Skip if not in QBase
        if serial not in existing_by_serial:
            skipped += 1
            continue

        try:
            extracted = extract_form(form_code, str(fpath))
            n_fields = len([k for k in extracted.keys() if not k.startswith('_')])

            if cmd == 'migrate':
                # Build form_data
                form_data = {}
                for k, v in extracted.items():
                    if k.startswith('_'): continue
                    if isinstance(v, (list, dict)):
                        form_data[k] = json.dumps(v, ensure_ascii=False)
                    elif isinstance(v, str):
                        form_data[k] = v
                    else:
                        form_data[k] = str(v) if v else ''
                if '_raw_text' in extracted:
                    form_data['_raw_text'] = extracted['_raw_text'][:2000]
                if '_items' in extracted:
                    form_data['items'] = json.dumps(extracted['_items'], ensure_ascii=False)

                # Update in Supabase
                update_url = f"{SUPABASE_URL}/rest/v1/records?serial=eq.{serial}"
                update_data = {'form_data': form_data}
                update_req = urllib.request.Request(update_url, data=json.dumps(update_data).encode(), method='PATCH')
                update_req.add_header('apikey', SERVICE_KEY)
                update_req.add_header('Authorization', f'Bearer {SERVICE_KEY}')
                update_req.add_header('Content-Type', 'application/json')
                try:
                    urllib.request.urlopen(update_req, timeout=10)
                    updated += 1
                except urllib.error.HTTPError as e:
                    errors.append(f"{serial}: HTTP {e.code} {e.read().decode()[:100]}")

        except Exception as e:
            errors.append(f"{serial}: {type(e).__name__}: {str(e)[:80]}")

    print(f"\n{'='*60}")
    print(f"📊 CUSTOM EXTRACTOR RESULTS ({cmd.upper()} mode)")
    print(f"   Total files: {total}")
    print(f"   Skipped (not in QBase): {skipped}")
    if cmd == 'migrate':
        print(f"   Updated: {updated}")
    print(f"   Errors: {len(errors)}")
    if errors:
        for e in errors[:20]:
            print(f"     ⚠️ {e}")

if __name__ == '__main__':
    main()