#!/usr/bin/env python3
"""
Verify ALL records in QBase against their source DOCX files.
Reports discrepancies field-by-field.
"""
import json, os, re, urllib.request
from pathlib import Path
from docx import Document
from collections import defaultdict

RECORDS_DIR = Path('/home/Kepv/Downloads/03.Records')

def load_env():
    env = {}
    with open('/mnt/ahmed/Projects/qbase/.env.local') as f:
        for line in f:
            line = line.strip()
            if line and not line.startswith('#') and '=' in line:
                k, v = line.split('=', 1)
                env[k] = v
    return env

env = load_env()
BASE_URL = env['SUPABASE_URL']
HEADERS = {
    'apikey': env['SUPABASE_SERVICE_ROLE_KEY'],
    'Authorization': 'Bearer ' + env['SUPABASE_SERVICE_ROLE_KEY'],
    'Content-Type': 'application/json',
    'Prefer': 'return=representation'
}

def supabase_get(path):
    url = f"{BASE_URL}/rest/v1/{path}"
    req = urllib.request.Request(url, headers=HEADERS)
    return json.loads(urllib.request.urlopen(req, timeout=30).read())

# ── Find DOCX file for a serial ─────────────────────────────────────
def find_docx_file(serial):
    """Find the DOCX file path for a given serial like F/08-001."""
    # Try various naming patterns
    patterns = [
        serial.replace('/', '_'),          # F_08-001
        serial.replace('/', '-'),          # F-08-001
        serial,                             # F/08-001 (unlikely in filename)
    ]
    
    for root, dirs, files in os.walk(RECORDS_DIR):
        for file in files:
            if not file.endswith('.docx'):
                continue
            for pat in patterns:
                if pat in file:
                    return Path(root) / file
    return None

# ── Extract data from DOCX ──────────────────────────────────────────
def clean_cell(text):
    return text.strip().replace('\n', ' ').replace('\r', '').replace('\t', ' ')

def extract_f08(doc):
    """Extract F/08 Order Form data."""
    table = doc.tables[0]
    data = {}
    
    for row in table.rows:
        cells = [clean_cell(cell.text) for cell in row.cells]
        if len(cells) < 2:
            continue
        
        label = cells[0].lower()
        value = cells[-1] if len(cells) > 1 else ''
        
        if 'customer' in label and 'name' in label:
            data['client_name'] = value
        elif 'mode of receipt' in label:
            data['mode_of_receipt'] = value
        elif 'statutory' in label and 'regulatory' in label:
            data['statutory_and_regulatory_requirements_if_any'] = value
        elif 'test certificate' in label:
            data['test_certificate_required'] = value
        elif 'delivery schedule' in label:
            data['delivery_schedule'] = value
        elif 'order' in label and ('status' in label or 'decision' in label):
            data['order_status'] = value
        elif 'remarks' in label:
            data['remarks'] = value
        elif 'reviewed by' in label:
            data['reviewed_by'] = value
        elif 'bill no' in label:
            data['bill_no'] = value
        elif 'despatch date' in label:
            data['despatch_date'] = value
    
    return data

def extract_generic(doc):
    """Generic extractor - grab all labeled fields from first table."""
    if not doc.tables:
        return {}
    table = doc.tables[0]
    data = {}
    
    for row in table.rows:
        cells = [clean_cell(cell.text) for cell in row.cells]
        if len(cells) >= 2:
            label = cells[0]
            value = cells[-1]
            if label and value and label != value:
                # Create a key from label
                key = re.sub(r'[^a-zA-Z0-9_]', '_', label.lower())
                key = re.sub(r'_+', '_', key).strip('_')
                if len(key) > 3:
                    data[key] = value
    return data

# ── Compare QBase vs DOCX ───────────────────────────────────────────
def compare_record(qbase_record):
    serial = qbase_record['serial']
    form_code = qbase_record['form_code']
    qbase_data = qbase_record.get('form_data', {})
    
    # Find DOCX
    docx_path = find_docx_file(serial)
    if not docx_path:
        return {'serial': serial, 'error': 'DOCX file not found', 'docx_path': None}
    
    # Extract from DOCX
    try:
        doc = Document(str(docx_path))
        if form_code == 'F/08':
            docx_data = extract_f08(doc)
        else:
            docx_data = extract_generic(doc)
    except Exception as e:
        return {'serial': serial, 'error': f'Extraction failed: {e}', 'docx_path': str(docx_path)}
    
    # Compare key fields
    discrepancies = []
    all_keys = set(qbase_data.keys()) | set(docx_data.keys())
    
    for key in all_keys:
        qval = str(qbase_data.get(key, '')).strip()
        dval = str(docx_data.get(key, '')).strip()
        
        # Skip empty vs empty
        if not qval and not dval:
            continue
        
        # Skip internal fields
        if key.startswith('_') or key in ['serial', 'form_code', 'id', 'created_at', 'updated_at', 'status']:
            continue
        
        if qval != dval:
            discrepancies.append({
                'field': key,
                'qbase': qval[:100] if qval else '(empty)',
                'docx': dval[:100] if dval else '(empty)',
            })
    
    return {
        'serial': serial,
        'form_code': form_code,
        'docx_path': str(docx_path),
        'docx_keys': list(docx_data.keys()),
        'qbase_keys': list(qbase_data.keys()),
        'discrepancies': discrepancies,
        'docx_data': docx_data,
        'error': None,
    }

# ── Main ────────────────────────────────────────────────────────────
if __name__ == '__main__':
    print("🔍 Fetching all records from QBase...")
    records = supabase_get('records?select=serial,form_code,form_data&limit=300')
    print(f"   Found {len(records)} records")
    
    results = []
    missing_docx = []
    with_discrepancies = []
    
    for i, rec in enumerate(records):
        serial = rec['serial']
        print(f"  [{i+1}/{len(records)}] Checking {serial}...", end=' ')
        
        result = compare_record(rec)
        
        if result['error']:
            if 'not found' in result['error']:
                missing_docx.append(serial)
                print("❌ DOCX missing")
            else:
                print(f"❌ {result['error']}")
        elif result['discrepancies']:
            with_discrepancies.append(result)
            print(f"⚠️ {len(result['discrepancies'])} fields differ")
        else:
            print("✅ OK")
        
        results.append(result)
    
    # Summary
    print(f"\n{'='*70}")
    print(f"📊 VERIFICATION SUMMARY")
    print(f"{'='*70}")
    print(f"Total records checked: {len(records)}")
    print(f"Missing DOCX files: {len(missing_docx)}")
    print(f"Records with discrepancies: {len(with_discrepancies)}")
    print(f"Perfect matches: {len(records) - len(missing_docx) - len(with_discrepancies)}")
    
    if missing_docx:
        print(f"\n📁 Missing DOCX files ({len(missing_docx)}):")
        for s in missing_docx:
            print(f"   - {s}")
    
    if with_discrepancies:
        print(f"\n⚠️ Records with discrepancies ({len(with_discrepancies)}):")
        for r in with_discrepancies:
            print(f"\n   {r['serial']} ({r['form_code']}):")
            for d in r['discrepancies'][:5]:  # Show first 5
                print(f"      Field: {d['field']}")
                print(f"        QBase: {d['qbase']}")
                print(f"        DOCX:  {d['docx']}")
            if len(r['discrepancies']) > 5:
                print(f"      ... and {len(r['discrepancies']) - 5} more")
    
    # Save full report
    report_path = '/tmp/qbase_verification_report.json'
    with open(report_path, 'w') as f:
        json.dump({
            'summary': {
                'total': len(records),
                'missing_docx': len(missing_docx),
                'with_discrepancies': len(with_discrepancies),
                'perfect': len(records) - len(missing_docx) - len(with_discrepancies),
            },
            'missing_docx': missing_docx,
            'with_discrepancies': with_discrepancies,
            'all_results': results,
        }, f, indent=2, default=str)
    
    print(f"\n📄 Full report saved to: {report_path}")
