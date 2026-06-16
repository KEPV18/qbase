#!/usr/bin/env python3
"""
Comprehensive QBase data quality fix.
Re-extracts ALL records from DOCX with per-form custom extractors.
Updates QBase with complete, accurate data.
"""
import json, os, re, urllib.request
from pathlib import Path
from docx import Document
from collections import defaultdict

RECORDS_DIR = Path('/home/Kepv/Downloads/03.Records')

def load_env():
    env = {}
    with open('/mnt/ahmed/Projects/qbase/.env.local') as f:
        for line in f:
            line = line.strip()
            if line and not line.startswith('#') and '=' in line:
                k, v = line.split('=', 1)
                env[k] = v
    return env

env = load_env()
BASE_URL = env['SUPABASE_URL']
HEADERS = {
    'apikey': env['SUPABASE_SERVICE_ROLE_KEY'],
    'Authorization': 'Bearer ' + env['SUPABASE_SERVICE_ROLE_KEY'],
    'Content-Type': 'application/json',
    'Prefer': 'return=representation'
}

def supabase_get(path):
    url = f"{BASE_URL}/rest/v1/{path}"
    req = urllib.request.Request(url, headers=HEADERS)
    return json.loads(urllib.request.urlopen(req, timeout=30).read())

def supabase_patch(path, data):
    url = f"{BASE_URL}/rest/v1/{path}"
    req = urllib.request.Request(url, data=json.dumps(data).encode(), headers=HEADERS, method='PATCH')
    return urllib.request.urlopen(req, timeout=30).read()

def clean_cell(text):
    return text.strip().replace('\n', ' ').replace('\r', '').replace('\t', ' ')

def find_docx_file(serial):
    """Find DOCX file for serial."""
    patterns = [
        serial.replace('/', '_'),
        serial.replace('/', '-'),
    ]
    for root, dirs, files in os.walk(RECORDS_DIR):
        for file in files:
            if not file.endswith('.docx'):
                continue
            for pat in patterns:
                if pat in file:
                    return Path(root) / file
    return None

# ============================================================================
# CUSTOM EXTRACTORS PER FORM
# ============================================================================

def extract_f08(doc):
    """F/08 Order Form — robust extractor."""
    table = doc.tables[0]
    data = {}
    
    for ri, row in enumerate(table.rows):
        cells = [clean_cell(c.text) for c in row.cells]
        if len(cells) < 2:
            continue
        
        label = cells[0].lower()
        value = cells[-1]
        
        if 'customer' in label:
            data['client_name'] = value
        elif 'mode of receipt' in label:
            data['mode_of_receipt'] = value
        elif 'statutory' in label and 'regulatory' in label:
            data['Statutory And Regulatory Requirements, If Any'] = value
            data['statutory_requirements'] = value
        elif 'test certificate' in label:
            data['test_certificate_required'] = value
        elif 'delivery schedule' in label:
            data['delivery_schedule'] = value
        elif label == 'order' or 'order status' in label:
            data['order_status'] = value
        elif 'remarks' in label:
            data['remarks'] = value
        elif 'reviewed by' in label:
            data['reviewed_by'] = value
        elif 'bill no' in label:
            data['bill_no'] = value
        elif 'despatch date' in label:
            data['despatch_date'] = value
        elif 'date' in label and len(cells) >= 6:
            # This might be the main date row
            if 'serial' in cells[0].lower() or 'sr.' in cells[0].lower():
                # Find date value
                for c in cells[1:]:
                    if re.match(r'\d{1,2}[/-]\d{1,2}[/-]\d{2,4}', c):
                        data['date'] = c
                        break
    
    # Extract items table
    items = []
    in_items = False
    for row in table.rows:
        cells = [clean_cell(c.text) for c in row.cells]
        if not cells:
            continue
        
        first = cells[0].lower()
        if 'product name' in first or 'specifications' in first:
            in_items = True
            continue
        
        if in_items and len(cells) >= 3:
            # Check if this is a data row (has content, not headers)
            if cells[0] and not any(x in cells[0].lower() for x in ['product', 'spec', 'qty', 'sr.']):
                try:
                    sn = int(cells[0])
                    items.append({
                        'product_name': cells[1] if len(cells) > 1 else '',
                        'specifications': cells[2] if len(cells) > 2 else '',
                        'qty': cells[3] if len(cells) > 3 else '',
                    })
                except:
                    pass
    
    if items:
        data['items'] = items
    
    return data

def extract_f30(doc):
    """F/30 Performance Appraisal — table-based form."""
    table = doc.tables[0]
    data = {}
    
    for row in table.rows:
        cells = [clean_cell(c.text) for c in row.cells]
        if len(cells) < 2:
            continue
        
        label = cells[0].lower().replace(':', '').strip()
        value = cells[-1].strip()
        
        if not value or value == label:
            continue
        
        if 'employee name' in label:
            data['employee_name'] = value
        elif 'date' in label and 'appraisal' in label:
            data['date'] = value
        elif 'department' in label:
            data['department'] = value
        elif 'position' in label:
            data['position'] = value
        elif 'review period' in label:
            data['review_period'] = value
        elif 'evaluated by' in label:
            data['evaluated_by'] = value
        elif 'total marking' in label or 'overall score' in label:
            data['total_marking'] = value
        elif 'promotion' in label or 'increment' in label:
            data['promotion_increment'] = value
        elif 'suggestions' in label:
            data['suggestions'] = value
    
    return data

def extract_f43(doc):
    """F/43 Induction Training — multi-row table."""
    table = doc.tables[0]
    data = {}
    
    # Header info
    for row in table.rows[:5]:
        cells = [clean_cell(c.text) for c in row.cells]
        if len(cells) >= 2:
            label = cells[0].lower()
            value = cells[-1]
            if 'employee name' in label:
                data['employee_name'] = value
            elif 'date of joining' in label:
                data['date_of_joining'] = value
            elif 'department' in label:
                data['department'] = value
            elif 'designation' in label:
                data['designation'] = value
    
    return data

def extract_generic_improved(doc, form_code):
    """Improved generic extractor — skips headers better."""
    if not doc.tables:
        return {}
    table = doc.tables[0]
    data = {}
    
    # Common layout keywords to skip
    skip_keywords = {'sr.', 's/n', 'sl no', 'serial', 'signature', 'date', 'name',
                     'department', 'remarks', 'no', 'id', 'f/' + form_code.split('/')[1],
                     'rev no', 'page', 'form', 'doc no'}
    
    for row in table.rows:
        cells = [clean_cell(c.text) for c in row.cells]
        if len(cells) < 2:
            continue
        
        label = cells[0].strip().lower()
        value = cells[-1].strip()
        
        # Skip if value equals label (header row)
        if value.lower() == label:
            continue
        
        # Skip if label is a generic header
        if any(kw in label for kw in skip_keywords) and len(label) < 20:
            continue
        
        # Skip if value looks like a header
        if any(kw in value.lower() for kw in skip_keywords) and len(value) < 15:
            continue
        
        # Skip empty values
        if not value:
            continue
        
        # Create key
        key = re.sub(r'[^a-zA-Z0-9_]', '_', label).strip('_')
        key = re.sub(r'_+', '_', key)
        if len(key) > 3:
            data[key] = value
    
    return data

def extract_raw_text(doc):
    """Extract all text from document as fallback."""
    texts = []
    for para in doc.paragraphs:
        if para.text.strip():
            texts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            row_text = ' | '.join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                texts.append(row_text)
    return '\n'.join(texts)

# ============================================================================
# MAIN
# ============================================================================

def main():
    print("🔍 Fetching all records from QBase...")
    records = supabase_get('records?select=id,serial,form_code,form_data&limit=300')
    print(f"   Found {len(records)} records\n")
    
    fixed = 0
    failed = 0
    skipped = 0
    
    for i, rec in enumerate(records):
        serial = rec['serial']
        form_code = rec['form_code']
        
        print(f"[{i+1}/{len(records)}] {serial} ({form_code})...", end=' ')
        
        # Find DOCX
        docx_path = find_docx_file(serial)
        if not docx_path:
            print("❌ DOCX not found")
            skipped += 1
            continue
        
        try:
            doc = Document(str(docx_path))
            
            # Extract with appropriate extractor
            if form_code == 'F/08':
                new_data = extract_f08(doc)
            elif form_code == 'F/30':
                new_data = extract_f30(doc)
            elif form_code == 'F/43':
                new_data = extract_f43(doc)
            else:
                new_data = extract_generic_improved(doc, form_code)
            
            # Always add raw text as fallback
            raw_text = extract_raw_text(doc)
            new_data['_raw_text'] = raw_text[:5000]  # Limit size
            
            # Merge with existing data (preserve fields we couldn't extract)
            existing = rec.get('form_data', {})
            merged = {**existing, **new_data}
            
            # Update QBase
            supabase_patch(f"records?id=eq.{rec['id']}", {'form_data': merged})
            print(f"✅ Updated ({len(new_data)} fields)")
            fixed += 1
            
        except Exception as e:
            print(f"❌ Error: {e}")
            failed += 1
    
    print(f"\n{'='*60}")
    print(f"📊 RESULTS: {fixed} fixed, {failed} failed, {skipped} skipped")
    print(f"{'='*60}")

if __name__ == '__main__':
    main()
