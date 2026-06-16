#!/usr/bin/env python3
"""
Fix ALL records in QBase using the proper custom extractors from migrate_v4_custom.py.
This script imports extractors from migrate_v4_custom and replaces form_data completely.
"""
import json, os, re, sys
import urllib.request
from pathlib import Path
from docx import Document

# Add path to import from migrate_v4_custom
sys.path.insert(0, '/mnt/ahmed/Projects/qbase')
from migrate_v4_custom import (
    clean_cells, extract_f08, extract_f10, extract_f19,
    extract_generic, FORM_CONFIG, FORM_FOLDER_MAP, RECORDS_DIR
)

def load_env():
    env = {}
    with open('/mnt/ahmed/Projects/qbase/.env.local') as f:
        for line in f:
            line = line.strip()
            if line and not line.startswith('#') and '=' in line:
                k, v = line.split('=', 1)
                env[k] = v
    return env

env = load_env()
BASE_URL = env['SUPABASE_URL']
HEADERS = {
    'apikey': env['SUPABASE_SERVICE_ROLE_KEY'],
    'Authorization': 'Bearer ' + env['SUPABASE_SERVICE_ROLE_KEY'],
    'Content-Type': 'application/json',
    'Prefer': 'return=representation'
}

def supabase_get(path):
    url = f"{BASE_URL}/rest/v1/{path}"
    req = urllib.request.Request(url, headers=HEADERS)
    return json.loads(urllib.request.urlopen(req, timeout=30).read())

def supabase_patch(path, data):
    url = f"{BASE_URL}/rest/v1/{path}"
    req = urllib.request.Request(url, data=json.dumps(data).encode(), headers=HEADERS, method='PATCH')
    return urllib.request.urlopen(req, timeout=30).read()

def find_docx_file(serial):
    """Find DOCX file for serial."""
    patterns = [serial.replace('/', '_'), serial.replace('/', '-')]
    for root, dirs, files in os.walk(RECORDS_DIR):
        for file in files:
            if not file.endswith('.docx'):
                continue
            for pat in patterns:
                if pat in file:
                    return Path(root) / file
    return None

def extract_with_proper_extractor(filepath, form_code):
    """Use the right extractor from migrate_v4_custom."""
    extractors = {
        'F/08': extract_f08,
        'F/10': extract_f10,
        'F/19': extract_f19,
    }
    
    if form_code in extractors:
        return extractors[form_code](filepath)
    else:
        # Use generic but better than before
        return extract_generic(filepath)

def main():
    print("🔍 Fetching all records from QBase...")
    records = supabase_get('records?select=id,serial,form_code,form_data&limit=300')
    print(f"   Found {len(records)} records\n")
    
    fixed = 0
    skipped = 0
    failed = 0
    
    for i, rec in enumerate(records):
        serial = rec['serial']
        form_code = rec['form_code']
        
        print(f"[{i+1}/{len(records)}] {serial}...", end=' ')
        
        docx_path = find_docx_file(serial)
        if not docx_path:
            print("❌ DOCX not found")
            skipped += 1
            continue
        
        try:
            # Extract with proper extractor
            new_data = extract_with_proper_extractor(str(docx_path), form_code)
            
            # Clean up: remove _type and internal keys
            for k in list(new_data.keys()):
                if k.startswith('_') or k in ['_type', '_items_header']:
                    del new_data[k]
            
            # Also add raw text for reference
            try:
                doc = Document(str(docx_path))
                raw_texts = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        raw_texts.append(para.text.strip())
                for table in doc.tables:
                    for row in table.rows:
                        row_text = ' | '.join(c.text.strip() for c in row.cells if c.text.strip())
                        if row_text:
                            raw_texts.append(row_text)
                new_data['_raw_text'] = '\n'.join(raw_texts)[:3000]
            except:
                pass
            
            # Update QBase
            supabase_patch(f"records?id=eq.{rec['id']}", {'form_data': new_data})
            print(f"✅ Updated ({len(new_data)} fields)")
            fixed += 1
            
        except Exception as e:
            print(f"❌ Error: {e}")
            failed += 1
    
    print(f"\n{'='*60}")
    print(f"📊 RESULTS: {fixed} fixed, {failed} failed, {skipped} skipped")
    print(f"{'='*60}")

if __name__ == '__main__':
    main()
