#!/usr/bin/env python3
"""
QMS -> QBase Migration Script v3
Reads DOCX files from local directory, extracts data, inserts into Supabase.
Source: /home/Kepv/Downloads/03.Records/

Usage:
  python3 migrate_v3.py analyze    # Count & preview only
  python3 migrate_v3.py migrate    # Actually insert into Supabase
"""
import json, os, sys, re, urllib.request, urllib.parse
from docx import Document
from pathlib import Path
from collections import Counter

# ============================================================================
# CONFIG
# ============================================================================
RECORDS_DIR = Path('/home/Kepv/Downloads/03.Records')

def load_env():
    env = {}
    with open('/mnt/ahmed/Projects/qbase/.env.local') as f:
        for line in f:
            line = line.strip()
            if line and not line.startswith('#') and '=' in line:
                k, v = line.split('=', 1)
                env[k] = v
    return env

ENV = load_env()
SUPABASE_URL = ENV['SUPABASE_URL']
SERVICE_KEY = ENV['SUPABASE_SERVICE_ROLE_KEY']

FORM_MAP = {
    'F-08 - Order Form':                ('F/08', 1, 'Sales & Customer Service', 'Order Form'),
    'F-09 - Customer Complaint':        ('F/09', 1, 'Sales & Customer Service', 'Customer Complaint'),
    'F-10 - Customer Feedback':         ('F/10', 1, 'Sales & Customer Service', 'Customer Feedback'),
    'F-50 - Customer Property':         ('F/50', 1, 'Sales & Customer Service', 'Customer Property'),
    'F-11 - Production Plan':           ('F/11', 2, 'Operations & Production', 'Production Plan'),
    'F-19 - Product Description':       ('F/19', 2, 'Operations & Production', 'Product Description'),
    'F-12 - Non-Conforming':            ('F/12', 3, 'Quality & Audit', 'Non-Conforming'),
    'F-17 - QA Test Request':           ('F/17', 3, 'Quality & Audit', 'QA Test Request'),
    'F-18 - Product Re-Call':           ('F/18', 3, 'Quality & Audit', 'Product Re-Call'),
    'F-22 - Corrective Action':         ('F/22', 3, 'Quality & Audit', 'Corrective Action'),
    'F-25 - Audit Plan':                ('F/25', 3, 'Quality & Audit', 'Audit Plan'),
    'F-47 - Audit Checklist':           ('F/47', 3, 'Quality & Audit', 'Audit Checklist'),
    'F-48 - Internal Audit Report':     ('F/48', 3, 'Quality & Audit', 'Internal Audit Report'),
    'F-13 - Purchase Order':            ('F/13', 4, 'Procurement & Vendors', 'Purchase Order'),
    'F-14 - Incoming Inspection':       ('F/14', 4, 'Procurement & Vendors', 'Incoming Inspection'),
    'F-15 - Approved Vendor List':      ('F/15', 4, 'Procurement & Vendors', 'Approved Vendor List'),
    'F-16 Supplier Registration Form':  ('F/16', 4, 'Procurement & Vendors', 'Supplier Registration Form'),
    'F-28 - Training Attendance':       ('F/28', 5, 'HR & Training', 'Training Attendance'),
    'F-29 - Training Record':           ('F/29', 5, 'HR & Training', 'Training Record'),
    'F-30 - Performance Appraisal':     ('F/30', 5, 'HR & Training', 'Performance Appraisal'),
    'F-40 - Competence Matrix':         ('F/40', 5, 'HR & Training', 'Competence Matrix'),
    'F-41 - Gap Analysis':              ('F/41', 5, 'HR & Training', 'Gap Analysis'),
    'F-42 - Annual Training Plan':      ('F/42', 5, 'HR & Training', 'Annual Training Plan'),
    'F-43 - Induction Training Record': ('F/43', 5, 'HR & Training', 'Induction Training Record'),
    'F-44 - Job Description':           ('F/44', 5, 'HR & Training', 'Job Description'),
    'F-32 - R&D Request':               ('F/32', 6, 'R&D & Design', 'R&D Request'),
    'F-34 - Design Verification':       ('F/34', 6, 'R&D & Design', 'Design Verification'),
    'F-35 - Design Monitoring':         ('F/35', 6, 'R&D & Design', 'Design Monitoring'),
    'F-37 - Experiment Data':           ('F/37', 6, 'R&D & Design', 'Experiment Data'),
    'F-20 - Review Agenda':             ('F/20', 7, 'Management & Documentation', 'Review Agenda'),
    'F-21 - Review Minutes':            ('F/21', 7, 'Management & Documentation', 'Review Minutes'),
    'F-23 - Master List of Records':    ('F/23', 7, 'Management & Documentation', 'Master List of Records'),
    'F-24 - Objectives & Targets':      ('F/24', 7, 'Management & Documentation', 'Objectives & Targets'),
    'F-45 - Master List of Docs':        ('F/45', 7, 'Management & Documentation', 'Master List of Docs'),
    'F-46 - Change Management':          ('F/46', 7, 'Management & Documentation', 'Change Management'),
}

# ============================================================================
# DOCX PARSER
# ============================================================================
def extract_serial_from_filename(filename):
    """Extract serial from filename like 'F-08-001.docx', 'F_19-001.docx', 'F-30-029.doc'"""
    # Try pattern with dash: F-XX-NNN or F_XX-NNN
    m = re.search(r'[Ff][_/\-](\d{2})[\-_](\d+)', filename, re.IGNORECASE)
    if m:
        return f"F/{m.group(1)}-{m.group(2)}"
    # Try pattern with slash: F/XX-NNN
    m = re.search(r'(F[/\-]\d{2}[-]\d+)', filename, re.IGNORECASE)
    if m:
        s = m.group(1).replace('\\', '/')
        m2 = re.match(r'F(\d{2})[-/](\d+)', s)
        if m2:
            return f"F/{m2.group(1)}-{m2.group(2)}"
    return None

def parse_qms_docx(filepath):
    doc = Document(filepath)
    all_paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    if not doc.tables:
        return _extract_paragraphs(all_paras)
    all_rows = []
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace('\n', ' | ').replace('\r', '') for cell in row.cells]
            all_rows.append(cells)
    return _extract_rows(all_rows, all_paras)

def _extract_paragraphs(paragraphs):
    data = {'_type': 'paragraphs'}
    full_text = '\n'.join(paragraphs)
    m = re.search(r'(F[/\-]\d{2}[-]\d+)', full_text)
    if m:
        s = m.group(1).replace('\\', '/')
        m2 = re.match(r'F(\d{2})[-/](\d+)', s)
        if m2: data['serial'] = f"F/{m2.group(1)}-{m2.group(2)}"
    for p in paragraphs:
        for sep in ['\U0001f822', '\t']:  # 🡪 arrow + tab
            if sep in p:
                label, _, value = p.partition(sep)
                label = label.strip().lower()
                value = value.strip()
                if value: _map_label(label, value, data)
    data['_raw_text'] = full_text[:3000]
    return data

def _extract_rows(rows, paragraphs=None):
    data = {'_type': 'table'}
    all_cells = []
    for row in rows: all_cells.extend([c for c in row if c])
    full_text = ' | '.join(all_cells)

    m = re.search(r'(F[/\-]\d{2}[-]\d+)', full_text)
    if m:
        s = m.group(1).replace('\\', '/')
        m2 = re.match(r'F(\d{2})[-/](\d+)', s)
        if m2: data['serial'] = f"F/{m2.group(1)}-{m2.group(2)}"
        else: data['serial'] = s

    for row in rows:
        for cell in row:
            for sep in ['\U0001f822', '\t']:
                if sep in cell:
                    parts = cell.split(sep)
                    if len(parts) >= 2:
                        label = parts[0].strip().lower()
                        value = sep.join(parts[1:]).strip()
                        if value and label: _map_label(label, value, data)

    header_idx = None
    for i, row in enumerate(rows):
        row_text = ' '.join(c.lower() for c in row)
        if any(kw in row_text for kw in ['sr. no', 'sr no', 'sr | no', 'topic no']):
            header_idx = i; break

    if header_idx is not None:
        headers = [c.lower().strip().replace('|', '').replace('\n', ' ').strip() for c in rows[header_idx]]
        for row in rows[header_idx + 1:]:
            if len(row) >= len(headers) and any(c.strip() for c in row):
                item = {}
                for ci, h in enumerate(headers):
                    if ci < len(row) and h: item[h] = row[ci].strip()
                if item and any(v for v in item.values() if v):
                    if '_items' not in data: data['_items'] = []
                    data['_items'].append(item)

    for row in rows:
        if len(row) < 2: continue
        for cell in row:
            if ':' in cell and '\U0001f822' not in cell and '\t' not in cell:
                parts = cell.split(':', 1)
                label = parts[0].strip().lower()
                value = parts[1].strip() if len(parts) > 1 else ''
                if value: _map_label(label, value, data)

    data['_raw_text'] = full_text[:3000]
    data['_num_rows'] = len(rows)
    return data

def _map_label(label, value, data):
    label = re.sub(r'\s+', ' ', label.replace('|', ' ').replace('\n', ' ')).strip()
    field_map = {
        'serial': ['serial no','serial','sr no','sr. no','ref no','ref. no','complaint sr. no','ref. no.'],
        'date': ['date','date of','dated'],
        'client_name': ['customer','client name','name of customer','name of the customer','client'],
        'mode_of_receipt': ['mode of receipt'],
        'description': ['description','details','desc','experiment title','title of training'],
        'remarks': ['remarks','remark','comments','comment'],
        'reviewed_by': ['reviewed by','reviewed'],
        'prepared_by': ['prepared by'],
        'employee_name': ['name of employee','employee name','name of the participant','name & designation','name of employees'],
        'department': ['department','dept'],
        'project': ['project','project name','name of product','product name'],
        'supplier': ['supplier','name of supplier'],
        'address': ['address'],
        'year': ['year'],
        'designation': ['designation','job title'],
        'qualification': ['qualification','qualifications'],
        'employee_id': ['id no','id no.','employee id'],
        'trainer': ['trainer','incharge'],
        'change_type': ['type of proposed change'],
    }
    for field, patterns in field_map.items():
        if field not in data:
            for p in patterns:
                if p in label and len(p) >= 3:
                    data[field] = value; return

# ============================================================================
# SUPABASE
# ============================================================================
def supabase_insert(records):
    clean_records = []
    for r in records:
        clean = {
            'form_code': r['form_code'],
            'serial': r['serial'],
            'form_name': r['form_name'],
            'status': 'draft',
            'section': r['section'],
            'section_name': r['section_name'],
            'form_data': r['form_data'],
        }
        clean_records.append(clean)

    url = f"{SUPABASE_URL}/rest/v1/records"
    data_str = json.dumps(clean_records, ensure_ascii=False).encode('utf-8')
    req = urllib.request.Request(url, data=data_str, method='POST')
    req.add_header('apikey', SERVICE_KEY)
    req.add_header('Authorization', f'Bearer {SERVICE_KEY}')
    req.add_header('Content-Type', 'application/json')
    req.add_header('Prefer', 'return=representation')
    try:
        resp = urllib.request.urlopen(req, timeout=30)
        return json.loads(resp.read())
    except urllib.error.HTTPError as e:
        return {'_error': f'HTTP {e.code}: {e.read().decode()[:500]}'}

def supabase_get_existing():
    url = f"{SUPABASE_URL}/rest/v1/records?select=id,form_code,serial,status"
    req = urllib.request.Request(url)
    req.add_header('apikey', SERVICE_KEY)
    req.add_header('Authorization', f'Bearer {SERVICE_KEY}')
    req.add_header('Range', '0-99999')
    return json.loads(urllib.request.urlopen(req, timeout=30).read())

def build_record(form_code, config, extracted, filename_serial):
    serial = filename_serial or extracted.get('serial', '')
    if not serial: return None

    form_data = {}
    for k, v in extracted.items():
        if k.startswith('_') or k == 'serial': continue
        if isinstance(v, (list, dict)):
            form_data[k] = json.dumps(v, ensure_ascii=False)
        elif isinstance(v, str):
            form_data[k] = v
        else:
            form_data[k] = str(v) if v else ''

    if '_items' in extracted and extracted['_items']:
        form_data['items'] = json.dumps(extracted['_items'], ensure_ascii=False)
    if '_raw_text' in extracted:
        form_data['_raw_text'] = extracted['_raw_text'][:2000]

    return {
        'form_code': form_code,
        'serial': serial,
        'form_name': config[3],
        'status': 'draft',
        'section': config[1],
        'section_name': config[2],
        'form_data': form_data,
    }

# ============================================================================
# MAIN
# ============================================================================
def main():
    cmd = sys.argv[1] if len(sys.argv) > 1 else 'analyze'

    existing = supabase_get_existing()
    existing_serials = {r['serial'] for r in existing if r.get('serial')}
    print(f"📊 Existing QBase records: {len(existing)} ({len(existing_serials)} unique serials)")

    # Scan local files
    all_files = []
    for section_dir in sorted(RECORDS_DIR.iterdir()):
        if not section_dir.is_dir() or 'form' in section_dir.name.lower():
            continue
        for form_dir in sorted(section_dir.iterdir()):
            if not form_dir.is_dir(): continue
            config = FORM_MAP.get(form_dir.name)
            if not config:
                continue
            for ext in ['*.docx', '*.doc']:
                for f in sorted(form_dir.rglob(ext)):
                    if 'Archive' in str(f) or 'archive' in str(f):
                        continue
                    all_files.append((f, config))

    print(f"📄 Found {len(all_files)} DOCX files to process\n")

    total_extracted = 0
    total_skipped = 0
    total_inserted = 0
    total_errors = []
    migration_log = []
    batch = []
    BATCH_SIZE = 10

    for i, (fpath, config) in enumerate(all_files):
        filename = fpath.name
        form_code = config[0]
        filename_serial = extract_serial_from_filename(filename)

        # Skip if already in QBase
        if filename_serial and filename_serial in existing_serials:
            total_skipped += 1
            continue

        try:
            extracted = parse_qms_docx(str(fpath))
            # Use filename serial (most reliable — avoids cross-reference confusion)
            if filename_serial:
                extracted['serial'] = filename_serial

            if not filename_serial and not extracted.get('serial'):
                total_errors.append(f"{form_code}/{filename}: no serial found")
                continue

            total_extracted += 1
            record = build_record(form_code, config, extracted, filename_serial)
            if not record:
                total_errors.append(f"{form_code}/{filename}: could not build record")
                continue

            n_fields = len([k for k in extracted.keys() if not k.startswith('_')])
            print(f"  [{i+1:3d}/{len(all_files)}] {record['serial']:12} → {n_fields:2d} fields  ({form_code})")
            sys.stdout.flush()

            if cmd == 'migrate':
                batch.append(record)
                if len(batch) >= BATCH_SIZE:
                    result = supabase_insert(batch)
                    if isinstance(result, dict) and '_error' in result:
                        total_errors.append(f"BATCH: {result['_error'][:200]}")
                        batch = []
                    else:
                        total_inserted += len(batch)
                        batch = []

            migration_log.append({
                'form_code': form_code,
                'serial': record['serial'],
                'file': filename,
                'fields': n_fields,
                'status': 'inserted' if cmd == 'migrate' else 'would_insert'
            })

        except Exception as e:
            total_errors.append(f"{form_code}/{filename}: {type(e).__name__}: {str(e)[:80]}")

    # Flush remaining batch
    if cmd == 'migrate' and batch:
        result = supabase_insert(batch)
        if isinstance(result, dict) and '_error' in result:
            total_errors.append(f"FINAL BATCH: {result['_error'][:200]}")
        else:
            total_inserted += len(batch)

    # Summary
    print(f"\n{'='*60}")
    print(f"📊 MIGRATION SUMMARY ({cmd.upper()} mode)")
    print(f"   Total DOCX files: {len(all_files)}")
    print(f"   Extracted: {total_extracted}")
    print(f"   Skipped (already in QBase): {total_skipped}")
    if cmd == 'migrate':
        print(f"   Inserted: {total_inserted}")
    print(f"   Errors: {len(total_errors)}")
    if total_errors:
        for e in total_errors[:20]:
            print(f"     ⚠️ {e}")

    # Save log
    with open('/tmp/qms_migration_log_v3.json', 'w') as f:
        json.dump(migration_log, f, indent=2, ensure_ascii=False)
    print(f"\n💾 Log saved to /tmp/qms_migration_log_v3.json")

if __name__ == '__main__':
    main()