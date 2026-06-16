#!/usr/bin/env python3
"""
Complete data quality fix for ALL QBase records.
Uses smart label:value parsing to properly extract data from all DOCX files.
"""
import json, os, re, urllib.request
from pathlib import Path
from docx import Document
from collections import defaultdict

RECORDS_DIR = Path('/home/Kepv/Downloads/03.Records')

def load_env():
    env = {}
    with open('/mnt/ahmed/Projects/qbase/.env.local') as f:
        for line in f:
            line = line.strip()
            if line and not line.startswith('#') and '=' in line:
                k, v = line.split('=', 1)
                env[k] = v
    return env

env = load_env()
BASE_URL = env['SUPABASE_URL']
HEADERS = {
    'apikey': env['SUPABASE_SERVICE_ROLE_KEY'],
    'Authorization': 'Bearer ' + env['SUPABASE_SERVICE_ROLE_KEY'],
    'Content-Type': 'application/json',
    'Prefer': 'return=representation'
}

def supabase_get(path):
    url = f"{BASE_URL}/rest/v1/{path}"
    req = urllib.request.Request(url, headers=HEADERS)
    return json.loads(urllib.request.urlopen(req, timeout=30).read())

def supabase_patch(path, data):
    url = f"{BASE_URL}/rest/v1/{path}"
    req = urllib.request.Request(url, data=json.dumps(data).encode(), headers=HEADERS, method='PATCH')
    return urllib.request.urlopen(req, timeout=30).read()

def find_docx_file(serial):
    patterns = [serial.replace('/', '_'), serial.replace('/', '-')]
    for root, dirs, files in os.walk(RECORDS_DIR):
        for file in files:
            if not file.endswith('.docx'):
                continue
            for pat in patterns:
                if pat in file:
                    return Path(root) / file
    return None

def clean_cell(text):
    return text.strip().replace('\n', ' ').replace('\r', '').replace('\t', ' ')

# ============================================================================
# SMART EXTRACTOR - uses label:value patterns
# ============================================================================

def extract_smart(doc, form_code):
    """Smart extractor that properly handles label:value pairs in DOCX tables."""
    data = {}
    
    # Skip words that indicate table headers, not data
    skip_values = {'yes', 'no', 'n/a', 'na', 'select', 'none', 'null'}
    skip_labels = {'sr.', 's/n', 'sl no', 'serial no.', 'signature', 'date', 'name', 
                   'department', 'remarks', 'no.', 'id no.', 'id', 'f/' + form_code.split('/')[1].lstrip('0'),
                   'page', 'form', 'rev no', 'doc no', 'qty.', 'rate', 'amount'}
    
    for table in doc.tables:
        for row in table.rows:
            cells = [clean_cell(c.text) for c in row.cells]
            if len(cells) < 2:
                continue
            
            # Strategy 1: Look for cells with 🡪 or : separator
            for cell in cells:
                # Split on 🡪 first ( Arabic-style arrow)
                if '🡪' in cell:
                    parts = cell.split('🡪', 1)
                    if len(parts) == 2:
                        label = parts[0].strip().lower()
                        value = parts[1].strip()
                        if value and len(label) > 2:
                            key = re.sub(r'[^a-zA-Z0-9_]', '_', label).strip('_')
                            key = re.sub(r'_+', '_', key)
                            if len(key) > 2:
                                data[key] = value
                
                # Split on : 
                elif ':' in cell and '🡪' not in cell:
                    parts = cell.split(':', 1)
                    if len(parts) == 2:
                        label = parts[0].strip().lower()
                        value = parts[1].strip()
                        if value and len(label) > 2 and len(value) > 0:
                            # Don't use if value is just a header word
                            if value.lower() not in skip_values:
                                key = re.sub(r'[^a-zA-Z0-9_]', '_', label).strip('_')
                                key = re.sub(r'_+', '_', key)
                                if len(key) > 2:
                                    data[key] = value
            
            # Strategy 2: Adjacent cells (label in first, value in second)
            # Only use if first cell looks like a label and second like a value
            for i in range(len(cells) - 1):
                label = cells[i].strip().lower()
                value = cells[i + 1].strip()
                
                # Skip if label is too short or looks like a value
                if len(label) < 3 or len(value) == 0:
                    continue
                
                # Skip if label is a generic header
                if any(kw in label for kw in skip_labels) and len(label) < 20:
                    continue
                
                # Skip if value looks like a header
                if any(kw in value.lower() for kw in skip_labels) and len(value) < 15:
                    continue
                
                # Skip if value equals label (header row)
                if value.lower() == label:
                    continue
                
                # Skip if value is just a yes/no/na
                if value.lower() in skip_values:
                    continue
                
                key = re.sub(r'[^a-zA-Z0-9_]', '_', label).strip('_')
                key = re.sub(r'_+', '_', key)
                if len(key) > 2 and key not in data:
                    data[key] = value
    
    return data

# ============================================================================
# F/08 CUSTOM EXTRACTOR
# ============================================================================

def extract_f08(doc):
    """F/08 Order Form - complete extractor."""
    data = {}
    if not doc.tables:
        return data
    
    table = doc.tables[0]
    items = []
    in_items = False
    
    for row in table.rows:
        cells = [clean_cell(c.text) for c in row.cells]
        row_text = '|'.join(cells)
        
        # Date
        m = re.search(r'(\d{2}/\d{2}/\d{4})', row_text)
        if m and 'date' not in data:
            data['date'] = m.group(1)
        
        # Customer
        if 'Customer' in row_text and 'client_name' not in data:
            for c in cells[1:]:
                if len(c) > 3 and 'Customer' not in c:
                    data['client_name'] = c
                    break
        
        # Mode of receipt
        if 'Mode' in row_text and 'Receipt' in row_text and 'mode_of_receipt' not in data:
            for c in cells[1:]:
                if len(c) > 5:
                    data['mode_of_receipt'] = c
                    break
        
        # Items header
        if 'Product Name' in row_text and 'Specifications' in row_text:
            in_items = True
            continue
        
        # Items rows
        if in_items and len(cells) >= 3:
            product = cells[1] if len(cells) > 1 else ''
            specs = cells[2] if len(cells) > 2 else ''
            qty = cells[3] if len(cells) > 3 else ''
            
            if product and len(product) > 3 and 'Product' not in product:
                items.append({
                    'product_name': product,
                    'specifications': specs,
                    'qty': qty or 'As per client volume'
                })
        
        # Test certificate
        if 'Test Certificate' in row_text and 'test_certificate_required' not in data:
            for c in cells:
                if c.strip() in ['No', 'Yes']:
                    data['test_certificate_required'] = c
                    break
        
        # Delivery schedule
        if 'Delivery' in row_text and 'Schedule' in row_text and 'delivery_schedule' not in data:
            for c in cells[1:]:
                if len(c) > 10:
                    data['delivery_schedule'] = c
                    break
        
        # Statutory
        if ('Statutory' in row_text or 'Regulatory' in row_text) and 'statutory_requirements' not in data:
            for c in cells[1:]:
                if len(c) > 10:
                    data['statutory_requirements'] = c
                    data['Statutory And Regulatory Requirements, If Any'] = c
                    break
        
        # Order status
        if 'Order' in row_text and ('Accepted' in row_text or 'Rejected' in row_text):
            for c in cells:
                if c in ['Accepted', 'Rejected']:
                    data['order_status'] = c
                    break
        
        # Remarks
        if 'Remarks' in row_text and 'remarks' not in data:
            for c in cells[1:]:
                if len(c) > 10:
                    data['remarks'] = c
                    break
        
        # Reviewed by
        if 'Reviewed' in row_text and 'reviewed_by' not in data:
            for c in cells[1:]:
                if len(c) > 3 and 'Reviewed' not in c:
                    data['reviewed_by'] = c
                    break
    
    if items:
        data['items'] = items
    
    return data

# ============================================================================
# MAIN
# ============================================================================

def main():
    print("🔍 Fetching all records from QBase...")
    records = supabase_get('records?select=id,serial,form_code,form_data&limit=300')
    print(f"   Found {len(records)} records\n")
    
    fixed = 0
    skipped = 0
    failed = 0
    
    for i, rec in enumerate(records):
        serial = rec['serial']
        form_code = rec['form_code']
        
        print(f"[{i+1}/{len(records)}] {serial}...", end=' ')
        
        docx_path = find_docx_file(serial)
        if not docx_path:
            print("❌ DOCX not found")
            skipped += 1
            continue
        
        try:
            doc = Document(str(docx_path))
            
            # Use appropriate extractor
            if form_code == 'F/08':
                new_data = extract_f08(doc)
            else:
                new_data = extract_smart(doc, form_code)
            
            # Add raw text
            try:
                raw_texts = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        raw_texts.append(para.text.strip())
                for table in doc.tables:
                    for row in table.rows:
                        row_text = ' | '.join(c.text.strip() for c in row.cells if c.text.strip())
                        if row_text:
                            raw_texts.append(row_text)
                new_data['_raw_text'] = '\n'.join(raw_texts)[:3000]
            except:
                pass
            
            # Update QBase
            supabase_patch(f"records?id=eq.{rec['id']}", {'form_data': new_data})
            print(f"✅ Updated ({len(new_data)} fields)")
            fixed += 1
            
        except Exception as e:
            print(f"❌ Error: {e}")
            failed += 1
    
    print(f"\n{'='*60}")
    print(f"📊 RESULTS: {fixed} fixed, {failed} failed, {skipped} skipped")
    print(f"{'='*60}")

if __name__ == '__main__':
    main()
