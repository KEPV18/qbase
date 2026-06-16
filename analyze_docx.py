#!/usr/bin/env python3
"""Analyze DOCX structure to understand QMS record format."""
import json, sys
from docx import Document

filepath = sys.argv[1] if len(sys.argv) > 1 else '/tmp/qms_samples/F08-001.docx'
doc = Document(filepath)

print("=== FULL TABLE ANALYSIS ===")
for ti, table in enumerate(doc.tables):
    print(f"\n--- Table {ti}: {len(table.rows)} rows x {len(table.columns)} cols ---")
    for ri, row in enumerate(table.rows):
        cells = []
        for ci, cell in enumerate(row.cells):
            text = cell.text.strip().replace('\n', ' | ')
            if text:
                cells.append(f"[{ci}] {text[:100]}")
        if cells:
            print(f"  Row {ri}: {'; '.join(cells)}")
        if ri > 25:
            print(f"  ... ({len(table.rows) - ri - 1} more rows)")
            break

# Also extract any labels and values we can identify
print("\n\n=== LABEL-VALUE PAIRS ===")
for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        prev_label = None
        for ci, cell in enumerate(row.cells):
            text = cell.text.strip()
            if text and ':' in text:
                parts = text.split(':', 1)
                if len(parts) == 2 and parts[1].strip():
                    print(f"  T{ti}.R{ri}.C{ci}: label='{parts[0].strip()}' value='{parts[1].strip()[:60]}'")