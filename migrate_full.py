#!/usr/bin/env python3
"""
QMS → QBase Migration Script (Full)
Downloads ALL DOCX records from Google Drive, extracts data, inserts into QBase Supabase.
Usage:
  python3 migrate_full.py analyze   # Analyze only, don't insert
  python3 migrate_full.py migrate   # Full migration
  python3 migrate_full.py validate  # Validate all records exist
"""
import json, os, sys, re, urllib.request, urllib.parse, time
from datetime import datetime, timezone
from docx import Document
from pathlib import Path

# ============================================================================
# CONFIG
# ============================================================================
DRIVE_ROOT = '1zA3ZqbFwxsa75DBl55oXybB2jyxvagYS'
TOKEN_PATH=Path.home()/'.hermes/google_token.json'
SAMPLES_DIR=Path('/tmp/qms_all_records')
MIGRATION_LOG=Path('/tmp/qms_migration_log.json')
BATCH_SIZE=10  # Insert records in batches of 10

# Load Supabase credentials
def load_env():
    env={}
    with open('/mnt/ahmed/Projects/qbase/.env.local') as f:
        for line in f:
            line=line.strip()
            if line and not line.startswith('#') and '=' in line:
                k,v=line.split('=',1)
                env[k]=v
    return env

ENV=load_env()
SUPABASE_URL=ENV['SUPABASE_URL']
SERVICE_KEY=ENV['SUPABASE_SERVICE_ROLE_KEY']

# ============================================================================
# FORM CONFIG — maps folder names to form codes and sections
# ============================================================================
FORM_CONFIG = {
    'F/08 - Order Form':                  {'code':'F/08','section':1,'name':'Order Form'},
    'F/09 - Customer Complaint':           {'code':'F/09','section':1,'name':'Customer Complaint'},
    'F/10 - Customer Feedback':            {'code':'F/10','section':1,'name':'Customer Feedback'},
    'F/50 - Customer Property':            {'code':'F/50','section':1,'name':'Customer Property'},
    'F/11 - Production Plan':              {'code':'F/11','section':2,'name':'Production Plan'},
    'F/19 - Product Description':          {'code':'F/19','section':2,'name':'Product Description'},
    'F/12 - Non-Conforming':               {'code':'F/12','section':3,'name':'Non-Conforming'},
    'F/17 - QA Test Request':              {'code':'F/17','section':3,'name':'QA Test Request'},
    'F/18 - Product Re-Call':              {'code':'F/18','section':3,'name':'Product Re-Call'},
    'F/22 - Corrective Action':            {'code':'F/22','section':3,'name':'Corrective Action'},
    'F/25 - Audit Plan':                   {'code':'F/25','section':3,'name':'Audit Plan'},
    'F/47 - Audit Checklist':              {'code':'F/47','section':3,'name':'Audit Checklist'},
    'F/48 - Internal Audit Report':        {'code':'F/48','section':3,'name':'Internal Audit Report'},
    'F-16 Supplier Registration Form':     {'code':'F/16','section':4,'name':'Supplier Registration Form'},
    'F/13 - Purchase Order':               {'code':'F/13','section':4,'name':'Purchase Order'},
    'F/14 - Incoming Inspection':          {'code':'F/14','section':4,'name':'Incoming Inspection'},
    'F/15 - Approved Vendor List':          {'code':'F/15','section':4,'name':'Approved Vendor List'},
    'F/28 - Training Attendance':           {'code':'F/28','section':5,'name':'Training Attendance'},
    'F/29 - Training Record':               {'code':'F/29','section':5,'name':'Training Record'},
    'F/30 - Performance Appraisal':         {'code':'F/30','section':5,'name':'Performance Appraisal'},
    'F/40 - Competence Matrix':             {'code':'F/40','section':5,'name':'Competence Matrix'},
    'F/41 - Gap Analysis':                  {'code':'F/41','section':5,'name':'Gap Analysis'},
    'F/42 - Annual Training Plan':          {'code':'F/42','section':5,'name':'Annual Training Plan'},
    'F/43 - Induction Training Record':     {'code':'F/43','section':5,'name':'Induction Training Record'},
    'F/44 - Job Description':               {'code':'F/44','section':5,'name':'Job Description'},
    'F/32 - R&D Request':                   {'code':'F/32','section':6,'name':'R&D Request'},
    'F/34 - Design Verification':           {'code':'F/34','section':6,'name':'Design Verification'},
    'F/35 - Design Monitoring':             {'code':'F/35','section':6,'name':'Design Monitoring'},
    'F/37 - Experiment Data':               {'code':'F/37','section':6,'name':'Experiment Data'},
    'F/20 - Review Agenda':                 {'code':'F/20','section':7,'name':'Review Agenda'},
    'F/21 - Review Minutes':                {'code':'F/21','section':7,'name':'Review Minutes'},
    'F/23 - Master List of Records':        {'code':'F/23','section':7,'name':'Master List of Records'},
    'F/24 - Objectives & Targets':           {'code':'F/24','section':7,'name':'Objectives & Targets'},
    'F/45 - Master List of Docs':           {'code':'F/45','section':7,'name':'Master List of Docs'},
    'F/46 - Change Management':             {'code':'F/46','section':7,'name':'Change Management'},
}

SECTION_NAMES = {
    1:'Sales & Customer Service',
    2:'Operations & Production',
    3:'Quality & Audit',
    4:'Procurement & Vendors',
    5:'HR & Training',
    6:'R&D & Design',
    7:'Management & Documentation',
}

# ============================================================================
# GOOGLE DRIVE API
# ============================================================================
def get_access_token():
    with open(TOKEN_PATH) as f: data=json.load(f)
    post_data=urllib.parse.urlencode({
        'client_id':data['client_id'],'client_secret':data['client_secret'],
        'refresh_token':data['refresh_token'],'grant_type':'refresh_token',
    }).encode()
    req=urllib.request.Request('https://oauth2.googleapis.com/token',data=post_data)
    resp=urllib.request.urlopen(req)
    result=json.loads(resp.read())
    data['access_token']=result['access_token']
    with open(TOKEN_PATH,'w') as f: json.dump(data,f,indent=2)
    return result['access_token']

def drive_list(folder_id, token, page_token=None):
    q=f"'{folder_id}' in parents and trashed=false"
    url=f"https://www.googleapis.com/drive/v3/files?q={urllib.parse.quote(q)}&fields=nextPageToken,files(id,name,mimeType)&pageSize=200"
    if page_token: url+=f"&pageToken={page_token}"
    req=urllib.request.Request(url)
    req.add_header('Authorization',f'Bearer {token}')
    resp=urllib.request.urlopen(req)
    return json.loads(resp.read())

def drive_list_all(folder_id, token):
    """List all files in a folder, handling pagination."""
    all_files=[]
    page_token=None
    while True:
        result=drive_list(folder_id, token, page_token)
        all_files.extend(result.get('files',[]))
        page_token=result.get('nextPageToken')
        if not page_token: break
    return all_files

def download_docx(file_id, token, outpath):
    """Download a file from Drive, exporting Google Docs as DOCX."""
    meta=drive_api(f"/drive/v3/files/{file_id}?fields=mimeType,name", token)
    mime=meta.get('mimeType','')
    os.makedirs(os.path.dirname(outpath), exist_ok=True)
    
    if mime=='application/vnd.google-apps.document':
        url=f"https://www.googleapis.com/drive/v3/files/{file_id}/export?mimeType=application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif mime.startswith('application/vnd.google-apps.spreadsheet'):
        return None  # Skip spreadsheets (not DOCX)
    else:
        url=f"https://www.googleapis.com/drive/v3/files/{file_id}?alt=media"
    
    req=urllib.request.Request(url)
    req.add_header('Authorization',f'Bearer {token}')
    try:
        resp=urllib.request.urlopen(req, timeout=30)
        with open(outpath,'wb') as f: f.write(resp.read())
        return outpath
    except Exception as e:
        print(f"  ⚠️ Download failed: {e}")
        return None

def drive_api(path, token):
    url=f"https://www.googleapis.com{path}"
    req=urllib.request.Request(url)
    req.add_header('Authorization',f'Bearer {token}')
    resp=urllib.request.urlopen(req)
    return json.loads(resp.read())

# ============================================================================
# DOCX EXTRACTION
# ============================================================================
def parse_qms_docx(filepath):
    """Parse any QMS DOCX and extract all data into a flat dict."""
    doc=Document(filepath)
    
    # Extract paragraphs
    all_paras=[p.text.strip() for p in doc.paragraphs if p.text.strip()]
    
    if not doc.tables:
        return _extract_paragraphs(all_paras)
    
    # Merge all table rows
    all_rows=[]
    for table in doc.tables:
        for row in table.rows:
            cells=[cell.text.strip().replace('\n',' | ').replace('\r','') for cell in row.cells]
            all_rows.append(cells)
    
    return _extract_rows(all_rows, all_paras)


def _extract_paragraphs(paragraphs):
    """For forms like F/44 that use paragraphs instead of tables."""
    data={'_type':'paragraphs'}
    full_text='\n'.join(paragraphs)
    
    # Serial
    m=re.search(r'(F[/\-]\d+[-]\d+)',full_text)
    if m: data['serial']=m.group(1).replace('\\','/').replace('-','/')
    # Normalize
    if 'serial' in data:
        s=data['serial']
        m2=re.match(r'F(\d{2})[-/](\d+)',s)
        if m2: data['serial']=f"F/{m2.group(1)}-{m2.group(2)}"
    
    # Date
    m=re.search(r'Date[:\s]*(\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4})',full_text)
    if m: data['date']=m.group(1)
    
    # Key-value pairs
    for p in paragraphs:
        for sep in ['🡪','\t']:
            if sep in p:
                label,_,value=p.partition(sep)
                label=label.strip().lower()
                value=value.strip()
                if value: _map_label(label,value,data)
    
    # Also capture full text as raw_data for forms with no clear structure
    data['_raw_text']=full_text[:3000]
    return data


def _extract_rows(rows, paragraphs=None):
    """Extract fields from table rows using pattern matching."""
    data={'_type':'table'}
    
    # Join all cell text
    all_cells=[]
    for row in rows:
        all_cells.extend([c for c in row if c])
    full_text=' | '.join(all_cells)
    
    # 1. Serial
    m=re.search(r'(F[/\-]\d+[-]\d+)',full_text)
    if m:
        s=m.group(1).replace('\\','/').replace('-','/')
        m2=re.match(r'F(\d{2})[-/](\d+)',s)
        if m2: data['serial']=f"F/{m2.group(1)}-{m2.group(2)}"
        else: data['serial']=s
    
    # 2. Extract tab-separated label🡪value pairs
    for row in rows:
        for cell in row:
            for sep in ['🡪','\t']:
                if sep in cell:
                    parts=cell.split(sep)
                    if len(parts)>=2:
                        label=parts[0].strip().lower()
                        value=sep.join(parts[1:]).strip()
                        # Clean label
                        label=re.sub(r'\s+', ' ', label).strip()
                        if value and label:
                            _map_label(label, value, data)
    
    # 3. Extract table data (items with column headers)
    header_idx=None
    for i,row in enumerate(rows):
        row_text=' '.join(c.lower() for c in row)
        if 'sr. no' in row_text or 'sr no' in row_text or 'sr | no' in row_text or 'topic no' in row_text:
            header_idx=i
            break
    
    if header_idx is not None:
        headers=[c.lower().strip().replace('|','').replace('\n',' ').strip() for c in rows[header_idx]]
        for row in rows[header_idx+1:]:
            if len(row)>=len(headers) and any(c.strip() for c in row):
                item={}
                for ci,h in enumerate(headers):
                    if ci<len(row) and h:
                        item[h]=row[ci].strip()
                if item and any(v for v in item.values() if v):
                    if '_items' not in data: data['_items']=[]
                    data['_items'].append(item)
    
    # 4. For simple label-value cells (e.g. F/16, F/32)
    if '_items' not in data and len(rows)>0:
        for row in rows:
            # Skip rows that look like headers or titles
            if len(row)<2: continue
            for cell in row:
                for sep in [':', '🡪', '\t']:
                    if sep in cell:
                        parts=cell.split(sep,1)
                        label=parts[0].strip().lower()
                        value=parts[1].strip() if len(parts)>1 else ''
                        if value: _map_label(label, value, data)
    
    # 5. Capture raw data
    data['_raw_text']=full_text[:3000]
    data['_num_rows']=len(rows)
    
    return data


def _map_label(label, value, data):
    """Map common label patterns to QBase field names."""
    label=label.replace('|',' ').replace('\n',' ').strip()
    label_clean=re.sub(r'\s+',' ',label)
    
    field_map={
        'serial':['serial no','serial','sr no','sr. no','ref no','ref. no','complaint sr. no','ref. no.','project number','project no'],
        'date':['date','date of','dated','date of joining'],
        'client_name':['customer','client name','name of customer','name of the customer','name :','client','name of client'],
        'mode_of_receipt':['mode of receipt','mode','receipt'],
        'description':['description','details','desc','object and variables','experiment title','title'],
        'remarks':['remarks','remark','comments','comment'],
        'reviewed_by':['reviewed by','reviewed'],
        'prepared_by':['prepared by','prepared'],
        'employee_name':['name of employee','employee name','name','name of the participant','name & designation','name of employees'],
        'department':['department','dept'],
        'project':['project','project name','name of product','product name'],
        'supplier':['supplier','name of supplier','to','to,'],
        'product_type':['product type','product','product name'],
        'quantity':['qty','quantity','qty.'],
        'specifications':['specifications','specification','specs'],
        'address':['address'],
        'year':['year'],
        'month':['month'],
        'period':['period','quarter'],
        'designation':['designation','job title'],
        'qualification':['qualification','qualifications','qualifications required'],
        'experience':['experience'],
        'from':['from'],
        'to':['to','to head'],
        'designation':['designation'],
        'employee_id':['id no','id no.','employee id'],
        'training_date':['training date'],
        'course_name':['title of training','course name','topic','training'],
        'score':['score'],
        'result':['result'],
        'trainer':['trainer','incharge','trainer name'],
        'reason':['reason','reason for','reason of development'],
        'impact':['impact'],
        'scope':['scope'],
        'objective':['objective','objectives'],
        'methodology':['methodology','method'],
        'findings':['findings','output observed'],
        'conclusion':['conclusion'],
        'request_for':['request for'],
        'name_of_customer':['name of the customer','name of customer'],
        'change_type':['type of proposed change','change type'],
        'delivery_schedule':['delivery schedule','delivery'],
        'order_status':['order status','status'],
        'test_certificate_required':['test certificate required'],
    }
    
    for field, patterns in field_map.items():
        if field not in data:
            for p in patterns:
                if p in label_clean and len(p)>=3:
                    data[field]=value
                    return

# ============================================================================
# SUPABASE API
# ============================================================================
def supabase_insert(records):
    """Insert records into QBase Supabase. Returns list of inserted records."""
    url=f"{SUPABASE_URL}/rest/v1/records"
    data=json.dumps(records).encode()
    
    req=urllib.request.Request(url,data=data,method='POST')
    req.add_header('apikey',SERVICE_KEY)
    req.add_header('Authorization',f'Bearer {SERVICE_KEY}')
    req.add_header('Content-Type','application/json')
    req.add_header('Prefer','return=representation')
    
    try:
        resp=urllib.request.urlopen(req)
        return json.loads(resp.read())
    except urllib.error.HTTPError as e:
        error_body=e.read().decode()
        return {'_error':f'HTTP {e.code}: {error_body[:300]}'}

def supabase_get_existing():
    """Get all existing records from QBase."""
    url=f"{SUPABASE_URL}/rest/v1/records?select=id,form_code,serial,status"
    req=urllib.request.Request(url)
    req.add_header('apikey',SERVICE_KEY)
    req.add_header('Authorization',f'Bearer {SERVICE_KEY}')
    req.add_header('Range','0-99999')
    
    resp=urllib.request.urlopen(req)
    return json.loads(resp.read())

def build_record(form_code, form_config, extracted):
    """Build a QBase record dict from extracted data."""
    serial=extracted.get('serial','')
    if not serial:
        return None
    
    # Build form_data — everything except _ private keys and serial
    form_data={}
    for k,v in extracted.items():
        if k.startswith('_') or k=='serial': continue
        if isinstance(v,list): v=json.dumps(v,ensure_ascii=False)
        elif isinstance(v,dict): v=json.dumps(v,ensure_ascii=False)
        form_data[k]=v
    
    # Items are stored as JSON string
    if '_items' in extracted:
        items=extracted['_items']
        if items:
            form_data['items']=json.dumps(items,ensure_ascii=False)
    
    # Raw text for reference (truncated)
    if '_raw_text' in extracted:
        form_data['_raw_text']=extracted['_raw_text'][:2000]
    
    record={
        'form_code':form_code,
        'serial':serial,
        'form_name':form_config['name'],
        'status':'migrated',
        'section':form_config['section'],
        'section_name':SECTION_NAMES.get(form_config['section'],''),
        'form_data':form_data,
    }
    
    return record

# ============================================================================
# MAIN
# ============================================================================
def main():
    cmd=sys.argv[1] if len(sys.argv)>1 else 'analyze'
    token=get_access_token()
    
    # Get existing records to avoid duplicates
    existing=supabase_get_existing()
    existing_serials={r['serial'] for r in existing if r.get('serial')}
    print(f"📊 Existing QBase records: {len(existing)}")
    print(f"   Existing serials: {sorted(existing_serials)[:10]}{'...' if len(existing_serials)>10 else ''}")
    
    # Get all form folders from Drive
    root_contents=drive_api(f"/drive/v3/files?q='{DRIVE_ROOT}'+in+parents+and+trashed=false&fields=files(id,name,mimeType)&pageSize=200",token)
    section_folders=[f for f in root_contents.get('files',[]) if f['mimeType']=='application/vnd.google-apps.folder' and 'Records' in f.get('name','')]
    
    total_files=0
    total_extracted=0
    total_inserted=0
    total_errors=[]
    migration_log=[]
    
    for section_folder in section_folders:
        section_name=section_folder['name']
        print(f"\n📂 {section_name}")
        
        # Get form folders
        section_contents=drive_api(f"/drive/v3/files?q='{section_folder['id']}'+in+parents+and+trashed=false&fields=files(id,name,mimeType)&pageSize=200",token)
        form_folders=[f for f in section_contents.get('files',[]) if f['mimeType']=='application/vnd.google-apps.folder']
        
        for form_folder in form_folders:
            folder_name=form_folder['name']
            config=FORM_CONFIG.get(folder_name)
            if not config:
                print(f"  ⚠️ Unknown folder: {folder_name}")
                continue
            
            form_code=config['code']
            print(f"  📁 {folder_name} ({form_code})")
            
            # Get all files in this form folder
            files=drive_list_all(form_folder['id'], token)
            docx_files=[f for f in files if f.get('mimeType','') in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document','application/msword','application/vnd.google-apps.document']]
            
            print(f"     Found {len(docx_files)} documents")
            total_files+=len(docx_files)
            
            for file in docx_files:
                file_name=file['name']
                safe_name=file_name.replace('/','-').replace('\\','-')
                outpath=str(SAMPLES_DIR/form_code/safe_name)
                
                # Skip if already exists in QBase
                # Extract serial from filename
                m=re.search(r'(F[/\-]\d+[-]\d+)',file_name)
                file_serial=m.group(1).replace('\\','/').replace('-','/') if m else None
                if file_serial:
                    m2=re.match(r'F(\d{2})[-/](\d+)',file_serial)
                    if m2: file_serial=f"F/{m2.group(1)}-{m2.group(2)}"
                
                if file_serial and file_serial in existing_serials:
                    print(f"     ⏭️  {file_name} — already in QBase as {file_serial}")
                    total_extracted+=1
                    continue
                
                # Download
                token=get_access_token()  # Refresh token before each download
                dl_path=download_docx(file['id'], token, outpath)
                if not dl_path:
                    total_errors.append(f"{form_code}/{file_name}: download failed")
                    continue
                
                # Parse
                try:
                    extracted=parse_qms_docx(dl_path)
                    extracted_serial=extracted.get('serial','')
                    if not extracted_serial and file_serial:
                        extracted_serial=file_serial
                        extracted['serial']=file_serial
                    
                    if not extracted_serial:
                        total_errors.append(f"{form_code}/{file_name}: no serial found")
                        continue
                    
                    total_extracted+=1
                    
                    # Build record
                    record=build_record(form_code, config, extracted)
                    if not record:
                        total_errors.append(f"{form_code}/{file_name}: could not build record")
                        continue
                    
                    # Determine status based on extracted data (default = migrated for historical records)
                    if 'date' in extracted:
                        record['date']=extracted['date']
                    
                    if cmd=='migrate':
                        # Insert into Supabase
                        result=supabase_insert([record])
                        if isinstance(result, dict) and '_error' in result:
                            total_errors.append(f"{form_code}/{file_serial}: insert error — {result['_error'][:100]}")
                        else:
                            total_inserted+=1
                            print(f"     ✅ {file_name} → {extracted_serial} inserted")
                    else:
                        print(f"     📝 {file_name} → {extracted_serial} ({len([k for k in extracted.keys() if not k.startswith('_')])} fields)")
                    
                    migration_log.append({
                        'form_code':form_code,
                        'serial':extracted_serial,
                        'file':file_name,
                        'status':'would_insert' if cmd=='analyze' else 'inserted',
                        'fields':len([k for k in extracted.keys() if not k.startswith('_')]),
                    })
                    
                except Exception as e:
                    total_errors.append(f"{form_code}/{file_name}: parse error — {str(e)[:80]}")
                    print(f"     ❌ {file_name}: {str(e)[:80]}")
                
                time.sleep(0.1)  # Rate limit
    
    # Summary
    print(f"\n{'='*60}")
    print(f"📊 MIGRATION SUMMARY")
    print(f"   Total files on Drive: {total_files}")
    print(f"   Successfully extracted: {total_extracted}")
    if cmd=='migrate':
        print(f"   Inserted into QBase: {total_inserted}")
    print(f"   Already in QBase (skipped): {len(existing_serials)}")
    print(f"   Errors: {len(total_errors)}")
    if total_errors:
        print(f"\n   Error details:")
        for e in total_errors[:20]:
            print(f"     {e}")
    
    # Save migration log
    with open(MIGRATION_LOG,'w') as f:
        json.dump(migration_log,f,indent=2,ensure_ascii=False)
    print(f"\n💾 Migration log saved to {MIGRATION_LOG}")

if __name__=='__main__':
    main()